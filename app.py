"""
Meeting-Protokoll App
=====================
Streamlit-App für den kompletten Protokoll-Workflow:
1. Audio hochladen
2. Transkribieren (OpenAI Whisper)
3. Protokoll generieren (Mistral AI - EU-basiert)
4. PDF erstellen
5. Download oder E-Mail-Versand

Starten mit: streamlit run app.py
"""

import os
import io
import re
import json
import secrets
import tempfile
import smtplib
import math
from datetime import datetime, timedelta
from pathlib import Path
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

import streamlit as st
import requests  # Für Mistral API (HTTP)
from openai import OpenAI  # Für Whisper Transkription
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from dotenv import load_dotenv

try:
    import resend as resend_lib
    RESEND_AVAILABLE = True
except ImportError:
    RESEND_AVAILABLE = False
import stripe  # Für Zahlungsabwicklung

# ffmpeg für Audio-Splitting
import subprocess
import shutil
import platform
import urllib.request

# ============================================================================
# KONFIGURATION - Hier Domain/Firma anpassen wenn neue Domain bereit
# ============================================================================
APP_DOMAIN = "minu-ai.ch"  # Domain
APP_URL = f"https://{APP_DOMAIN}"
COMPANY_NAME = "MINU-AI"  # Firmenname für Impressum
COMPANY_URL = f"https://{APP_DOMAIN}"
SUPPORT_EMAIL = f"support@{APP_DOMAIN}"  # support@minu-ai.ch
CONTACT_EMAIL = f"info@{APP_DOMAIN}"  # info@minu-ai.ch
# Analytics - Umami (cookieless, DSGVO-konform)
ANALYTICS_SCRIPT = "https://stats.spekt.ch/script.js"
ANALYTICS_ID = "a0b144e3-49da-4d39-bf78-9b1fb28b06d9"

# Erlaubte Länder für Stripe (gemäss AGB §5)
# CH, UK, EWR (IS, LI, NO) und EU-27
ALLOWED_COUNTRIES = [
    'CH',  # Schweiz
    'GB',  # Vereinigtes Königreich
    'IS', 'LI', 'NO',  # EWR (nicht EU)
    # EU-27 Länder
    'AT', 'BE', 'BG', 'HR', 'CY', 'CZ', 'DK', 'EE', 'FI', 'FR',
    'DE', 'GR', 'HU', 'IE', 'IT', 'LV', 'LT', 'LU', 'MT', 'NL',
    'PL', 'PT', 'RO', 'SK', 'SI', 'ES', 'SE'
]
# ============================================================================

# ============================================================================
# GeoIP-basierte Spracherkennung und Länderprüfung
# ============================================================================

def detect_country_from_ip():
    """
    Erkennt das Land des Besuchers anhand seiner IP-Adresse.
    Gibt den 2-stelligen Ländercode zurück (z.B. 'CH', 'DE', 'US').
    Bei Fehler wird None zurückgegeben.
    """
    try:
        response = requests.get(
            "http://ip-api.com/json/?fields=status,countryCode",
            timeout=3
        )
        if response.status_code == 200:
            data = response.json()
            if data.get("status") == "success":
                return data.get("countryCode", "").upper()
    except Exception:
        pass
    return None

def is_country_allowed(country_code: str = None) -> tuple:
    """
    Prüft ob das Land des Besuchers für den Kauf erlaubt ist (gemäss AGB §5).

    Args:
        country_code: Optional - 2-stelliger Ländercode. Falls None, wird IP-basiert ermittelt.

    Returns:
        tuple: (is_allowed: bool, country_code: str or None)
    """
    if country_code is None:
        country_code = detect_country_from_ip()

    if country_code is None:
        # Bei Fehler erlauben wir den Kauf (Fallback)
        return (True, None)

    is_allowed = country_code in ALLOWED_COUNTRIES
    return (is_allowed, country_code)

def get_country_name(country_code: str, lang: str = "en") -> str:
    """Gibt den Ländernamen für einen Ländercode zurück."""
    country_names = {
        "de": {
            "US": "USA", "CN": "China", "RU": "Russland", "IN": "Indien", "BR": "Brasilien",
            "KP": "Nordkorea", "IR": "Iran", "SY": "Syrien", "CU": "Kuba", "BY": "Belarus"
        },
        "en": {
            "US": "USA", "CN": "China", "RU": "Russia", "IN": "India", "BR": "Brazil",
            "KP": "North Korea", "IR": "Iran", "SY": "Syria", "CU": "Cuba", "BY": "Belarus"
        },
        "fr": {
            "US": "États-Unis", "CN": "Chine", "RU": "Russie", "IN": "Inde", "BR": "Brésil",
            "KP": "Corée du Nord", "IR": "Iran", "SY": "Syrie", "CU": "Cuba", "BY": "Biélorussie"
        },
        "it": {
            "US": "USA", "CN": "Cina", "RU": "Russia", "IN": "India", "BR": "Brasile",
            "KP": "Corea del Nord", "IR": "Iran", "SY": "Siria", "CU": "Cuba", "BY": "Bielorussia"
        }
    }
    names = country_names.get(lang, country_names["en"])
    return names.get(country_code, country_code)

def detect_language_from_ip():
    """
    Erkennt das Land des Besuchers anhand seiner IP-Adresse
    und gibt die passende Sprache zurück.

    Verwendet ip-api.com (kostenlos, 45 Anfragen/Minute)

    Sprachzuordnung:
    - DE, AT, LI (+ CH deutsch) → Deutsch
    - FR, BE, LU (+ CH französisch) → Französisch
    - IT (+ CH italienisch) → Italienisch
    - Alle anderen → Englisch
    """
    try:
        # IP-API aufrufen (kostenlos, keine API-Key nötig)
        response = requests.get(
            "http://ip-api.com/json/?fields=status,countryCode",
            timeout=3
        )

        if response.status_code == 200:
            data = response.json()

            if data.get("status") == "success":
                country_code = data.get("countryCode", "").upper()

                # Deutschsprachige Länder
                if country_code in ["DE", "AT", "LI"]:
                    return "de"

                # Französischsprachige Länder
                elif country_code in ["FR", "BE", "LU", "MC"]:
                    return "fr"

                # Italienischsprachig
                elif country_code in ["IT", "SM", "VA"]:
                    return "it"

                # Schweiz - Standard auf Deutsch (grösster Anteil)
                elif country_code == "CH":
                    return "de"

                # Alle anderen → Englisch
                else:
                    return "en"

    except Exception:
        # Bei Fehler (Timeout, Netzwerk, etc.) → Englisch als Fallback
        pass

    return "en"

# ============================================================================
# Mehrsprachigkeit - Übersetzungen
# ============================================================================

TRANSLATIONS = {
    "en": {
        "flag": "🇬🇧",
        "lang_name": "English",
        "title": "MINU Minutes AI",
        "subtitle": "Turn audio into professional transcripts.",
        "slogan": "In minutes instead of hours.",
        "slogan_gradient": "precise, fast, AI-powered",
        "audio_conversion_text": "Your audio file will be converted into a minutes file in minutes.",
        "upload_info_title": "💡 Tip for better results",
        "upload_info_text": "If you mention the meeting title and introduce all participants at the beginning of the audio recording, this information will be automatically included in the protocol. You can also add or change this later in the Word document.<br><br>🌍 <b>Language:</b> The protocol will be generated in your selected language – regardless of what language is spoken in the audio. Your recording can be in English, German, French or Italian.",
        "upload_label": "Upload audio file",
        "upload_help": "MP3, WAV, M4A, OGG, WEBM, MP4",
        "upload_hint": "MP3, WAV, M4A · Max. 200 MB",
        "processing": "Processing...",
        "transcribing": "Transcribing audio...",
        "generating": "Generating protocol...",
        "creating_pdf": "Creating PDF...",
        "creating_word": "Creating Word...",
        "done": "Done!",
        "file_uploaded_success": "Audio file successfully uploaded",
        "your_protocol_ready": "✓ Your protocol is ready!",
        "download_pdf": "Download PDF",
        "download_word": "Download Word",
        "send_email": "Send via email",
        "email_placeholder": "Enter email address",
        "email_format_label": "Format:",
        "send_button": "Send",
        "email_sent": "✓ Email sent!",
        "email_format_error": "Please select at least one format",
        "email_error": "Please enter email address",
        "show_protocol": "Show protocol",
        "new_protocol": "Create new protocol",
        "login_title": "🔐 MINU Minutes AI",
        "login_subtitle": "Please enter password",
        "password_placeholder": "Password",
        "login_button": "Login",
        "wrong_password": "Wrong password",
        "logout": "Logout",
        "my_account": "My Account",
        "subscription": "Subscription",
        "valid_until": "Valid until",
        "minutes_remaining": "Minutes remaining",
        "upgrade_plan": "Upgrade plan",
        "beta_free_access": "Free access for minutes of meetings - beta phase",
        "company_label": "Company / Organization",
        "company_placeholder": "Your company or organization",
        "email_label": "Email address",
        "email_placeholder_reg": "your.email@company.com",
        "register_button": "Start free",
        "consent_checkbox": "I confirm I am acting as a business/professional and accept the Terms and Privacy Policy",
        "privacy_link": "Privacy Policy",
        "error_company_required": "Please enter your company/organization",
        "error_email_required": "Please enter a valid email address",
        "error_email_exists": "This email is already registered.",
        "error_downgrade_not_allowed": "Downgrade not possible. Please cancel your current subscription first and re-register.",
        "error_plan_already_active": "You already have an active subscription with this email address.",
        "error_country_blocked": "MINU-AI is not available in your country. According to our Terms of Service (§5), only customers from Switzerland, EU, EEA and UK can purchase.",
        "error_consent_required": "Please accept the Terms & Conditions",
        "login_title": "Login",
        "login_subtitle": "Continue with your existing account",
        "login_button": "Login",
        "login_link": "Already registered? Login",
        "register_link": "New here? Register",
        "error_email_not_found": "Email not found. Please register first.",
        "privacy_title": "Privacy Policy",
        "sidebar_title": "MINU Minutes AI",
        "sidebar_subtitle": "Meeting Protocol Generator",
        "status": "STATUS",
        "transcript_ready": "✓ Transcript",
        "protocol_ready": "✓ Protocol",
        "documents_ready": "✓ Documents",
        "ready_to_start": "_Ready to start_",
        "technology": "TECHNOLOGY",
        "installation": "INSTALLATION",
        "install_as_app": "📱 Install as app",
        "start_recording": "🎙️ Start recording",
        "stop_recording": "⏹️ Stop recording",
        "recording_in_progress": "Recording...",
        "or_upload_file": "Or upload audio file",
        "upload_file_button": "Upload file",
        "mic_permission_needed": "Please allow microphone access",
        "recording_not_supported": "Recording not supported in this browser",
        "recording_warning_45min": "⚠️ 45 minutes reached. Recording will stop automatically at 60 minutes.",
        "recording_stopped_limit": "Recording stopped: Maximum 60 minutes reached.",
        "admin": "🔧 ADMIN",
        "activity_log": "📊 Activity Log",
        "no_activities": "No activities",
        "file_size": "File size",
        "small_file": "Small file - direct transcription...",
        "large_file_splitting": "Large file - splitting...",
        "parts_created": "audio parts created",
        "transcribing_part": "Transcribing part",
        "words_transcribed": "words transcribed",
        "total_words": "TOTAL",
        "words": "words",
        "characters": "characters",
        "sending_to_mistral": "Sending to Mistral AI...",
        "protocol_generated": "Protocol generated",
        "warning_too_short": "⚠️ WARNING: Protocol too short!",
        "processing_log": "📋 Processing Log",
        "error": "Error",
        "retry": "Retry",
        "file_too_large": "File is too large",
        "maximum": "Maximum",
        "ffmpeg_needed": "ffmpeg is needed. Please run in terminal: brew install ffmpeg",
        "email_subject": "Meeting Protocol from",
        "email_body_greeting": "Hello,",
        "email_body_text": "Please find attached the meeting protocol as PDF and Word document.",
        "email_body_closing": "Best regards",
        "email_rating_intro": "How satisfied are you with MINU Minutes AI?",
        "email_rating_thanks": "Thank you for your feedback!",
        "welcome_email_subject": "Welcome to MINU Minutes AI! 🎉",
        "welcome_email_greeting": "Welcome to MINU Minutes AI!",
        "welcome_email_intro": "Thank you for registering. You now have access to our AI-powered meeting protocol tool.",
        "welcome_email_benefits_title": "Your benefits:",
        "welcome_email_benefit1": "⏱️ Save time – Protocols in minutes, not hours",
        "welcome_email_benefit2": "📱 Flexible – Record directly from your phone or upload audio",
        "welcome_email_benefit3": "✨ Professional – Ready-to-use protocols as PDF and Word",
        "welcome_email_benefit4": "🔒 Secure – Your data stays in Switzerland (GDPR compliant)",
        "welcome_email_benefit5": "🗣️ Dialect-friendly – Regional dialects and accents are transcribed without any problems",
        "welcome_email_cta": "Start now and create your first protocol:",
        "welcome_email_help": "Questions? Simply reply to this email.",
        "whisper_language": "en",
        "gpt_language": "English",
        "limit_reached_total": "⚠️ Beta limit reached: Maximum 300 protocols have been created. The beta phase has ended.",
        "limit_reached_daily": "⚠️ Daily limit reached: You can create max. 4 protocols per day. Please try again tomorrow.",
        "limit_no_minutes": "⚠️ No transcription minutes remaining. Please upgrade your plan to continue.",
        "remaining_today": "Remaining protocols today",
        "beta_total_remaining": "Total beta protocols remaining",
        "verify_email_subject": "Confirm your email – MINU Minutes AI",
        "verify_email_heading": "Confirm your email address",
        "verify_email_text": "Thank you for registering! Please click the button below to confirm your email address and activate your account.",
        "verify_email_button": "Confirm email",
        "verify_email_expiry": "This link is valid for 48 hours.",
        "verify_email_ignore": "If you did not register, you can ignore this email.",
        "verify_email_sent_title": "Check your inbox",
        "verify_email_sent_text": "We have sent a confirmation email to <b>{email}</b>. Please click the link in the email to activate your account.",
        "verify_email_spam_hint": "Not received? Check your spam folder.",
        "verify_resend_button": "Resend email",
        "verify_resend_wait": "Please wait {seconds} seconds before resending.",
        "verify_resend_success": "✓ Confirmation email resent!",
        "verify_success_title": "Email confirmed!",
        "verify_success_text": "Your email has been successfully confirmed. You can now log in.",
        "verify_success_login": "Log in now",
        "verify_expired_title": "Link expired",
        "verify_expired_text": "This confirmation link has expired. Please register again.",
        "verify_invalid_title": "Invalid link",
        "verify_invalid_text": "This confirmation link is invalid.",
        "error_email_not_verified": "Please confirm your email address first. Check your inbox for the confirmation email."
    },
    "de": {
        "flag": "🇩🇪",
        "lang_name": "Deutsch",
        "title": "MINU Protokoll KI",
        "subtitle": "Verwandle Audio in professionelle Protokolle.",
        "slogan": "In Minuten statt Stunden.",
        "slogan_gradient": "präzise, schnell, KI-gestützt",
        "audio_conversion_text": "Ihre Audio-Datei wird in Minuten zum Protokoll.",
        "upload_info_title": "💡 Tipp für bessere Ergebnisse",
        "upload_info_text": "Wenn du den Sitzungstitel nennst und alle Teilnehmenden zu Beginn der Audioaufnahme vorstellst, werden diese Informationen automatisch ins Protokoll übernommen. Du kannst dies auch nachträglich im Word-Dokument ergänzen oder ändern.<br><br>🌍 <b>Sprache:</b> Das Protokoll wird in deiner gewählten Sprache erstellt – unabhängig davon, welche Sprache im Audio gesprochen wird. Deine Aufnahme kann auf Englisch, Deutsch, Französisch oder Italienisch sein.",
        "upload_label": "Audio-Datei hochladen",
        "upload_help": "MP3, WAV, M4A, OGG, WEBM, MP4",
        "upload_hint": "MP3, WAV, M4A · Max. 200 MB",
        "processing": "Verarbeitung...",
        "transcribing": "Transkribiere Audio...",
        "generating": "Erstelle Protokoll...",
        "creating_pdf": "Generiere PDF...",
        "creating_word": "Generiere Word...",
        "done": "Fertig!",
        "file_uploaded_success": "Audio-Datei erfolgreich übermittelt",
        "your_protocol_ready": "✓ Dein Protokoll ist fertig!",
        "download_pdf": "PDF laden",
        "download_word": "Word laden",
        "send_email": "Per E-Mail versenden",
        "email_placeholder": "E-Mail-Adresse eingeben",
        "email_format_label": "Format:",
        "send_button": "Senden",
        "email_sent": "✓ E-Mail gesendet!",
        "email_format_error": "Bitte mindestens ein Format auswählen",
        "email_error": "Bitte E-Mail-Adresse eingeben",
        "show_protocol": "Protokoll anzeigen",
        "new_protocol": "Neues Protokoll erstellen",
        "login_title": "🔐 MINU Protokoll KI",
        "login_subtitle": "Bitte Passwort eingeben",
        "password_placeholder": "Passwort",
        "login_button": "Anmelden",
        "wrong_password": "Falsches Passwort",
        "logout": "Abmelden",
        "my_account": "Mein Konto",
        "subscription": "Abonnement",
        "valid_until": "Gültig bis",
        "minutes_remaining": "Minuten verbleibend",
        "upgrade_plan": "Plan upgraden",
        "beta_free_access": "Kostenloser Zugang für Protokolle - Beta-Phase",
        "company_label": "Firma / Organisation",
        "company_placeholder": "Ihre Firma oder Organisation",
        "email_label": "E-Mail-Adresse",
        "email_placeholder_reg": "ihre.email@firma.ch",
        "register_button": "Kostenlos starten",
        "consent_checkbox": "Ich bestätige, als Unternehmer/Gewerbetreibender zu handeln und akzeptiere die AGB und Datenschutzerklärung",
        "privacy_link": "Datenschutzerklärung",
        "error_company_required": "Bitte geben Sie Ihre Firma/Organisation an",
        "error_email_required": "Bitte geben Sie eine gültige E-Mail-Adresse an",
        "error_email_exists": "Diese E-Mail ist bereits registriert.",
        "error_downgrade_not_allowed": "Downgrade nicht möglich. Bitte kündigen Sie zuerst Ihr aktuelles Abo und registrieren Sie sich neu.",
        "error_plan_already_active": "Sie haben bereits ein aktives Abonnement mit dieser E-Mail-Adresse.",
        "error_country_blocked": "MINU-AI ist in Ihrem Land nicht verfügbar. Gemäss unseren AGB (§5) können nur Kunden aus der Schweiz, EU, EWR und UK kaufen.",
        "error_consent_required": "Bitte akzeptieren Sie die AGB",
        "login_title": "Anmelden",
        "login_subtitle": "Mit bestehendem Konto fortfahren",
        "login_button": "Anmelden",
        "login_link": "Bereits registriert? Anmelden",
        "register_link": "Neu hier? Registrieren",
        "error_email_not_found": "E-Mail nicht gefunden. Bitte registrieren Sie sich zuerst.",
        "privacy_title": "Datenschutzerklärung",
        "sidebar_title": "MINU Protokoll KI",
        "sidebar_subtitle": "Meeting-Protokoll Generator",
        "status": "STATUS",
        "transcript_ready": "✓ Transkript",
        "protocol_ready": "✓ Protokoll",
        "documents_ready": "✓ Dokumente",
        "ready_to_start": "_Bereit zum Start_",
        "technology": "TECHNOLOGIE",
        "installation": "INSTALLATION",
        "install_as_app": "📱 Als App installieren",
        "start_recording": "🎙️ Aufnahme starten",
        "stop_recording": "⏹️ Aufnahme beenden",
        "recording_in_progress": "Aufnahme läuft...",
        "or_upload_file": "Oder Audio-Datei hochladen",
        "upload_file_button": "Datei hochladen",
        "mic_permission_needed": "Bitte Mikrofon-Zugriff erlauben",
        "recording_not_supported": "Aufnahme wird in diesem Browser nicht unterstützt",
        "recording_warning_45min": "⚠️ 45 Minuten erreicht. Aufnahme stoppt automatisch bei 60 Minuten.",
        "recording_stopped_limit": "Aufnahme gestoppt: Maximum 60 Minuten erreicht.",
        "admin": "🔧 ADMIN",
        "activity_log": "📊 Aktivitäts-Log",
        "no_activities": "Keine Aktivitäten",
        "file_size": "Dateigrösse",
        "small_file": "Kleine Datei - direkte Transkription...",
        "large_file_splitting": "Grosse Datei - wird gesplittet...",
        "parts_created": "Audio-Teile erstellt",
        "transcribing_part": "Transkribiere Teil",
        "words_transcribed": "Wörter transkribiert",
        "total_words": "TOTAL",
        "words": "Wörter",
        "characters": "Zeichen",
        "sending_to_mistral": "Sende an Mistral AI...",
        "protocol_generated": "Protokoll generiert",
        "warning_too_short": "⚠️ WARNUNG: Protokoll zu kurz!",
        "processing_log": "📋 Verarbeitungs-Log",
        "error": "Fehler",
        "retry": "Erneut versuchen",
        "file_too_large": "Datei ist zu gross",
        "maximum": "Maximum",
        "ffmpeg_needed": "ffmpeg wird benötigt. Bitte im Terminal ausführen: brew install ffmpeg",
        "email_subject": "Meeting-Protokoll vom",
        "email_body_greeting": "Guten Tag,",
        "email_body_text": "Im Anhang finden Sie das Meeting-Protokoll als PDF und Word-Dokument.",
        "email_body_closing": "Freundliche Grüsse",
        "email_rating_intro": "Wie zufrieden sind Sie mit MINU Protokoll KI?",
        "email_rating_thanks": "Vielen Dank für Ihr Feedback!",
        "welcome_email_subject": "Willkommen bei MINU Protokoll KI! 🎉",
        "welcome_email_greeting": "Willkommen bei MINU Protokoll KI!",
        "welcome_email_intro": "Vielen Dank für Ihre Registrierung. Sie haben jetzt Zugang zu unserem KI-gestützten Protokoll-Tool.",
        "welcome_email_benefits_title": "Ihre Vorteile:",
        "welcome_email_benefit1": "⏱️ Zeitersparnis – Protokolle in Minuten statt Stunden",
        "welcome_email_benefit2": "📱 Flexibel – Direkt vom Handy aufnehmen oder Audio hochladen",
        "welcome_email_benefit3": "✨ Professionell – Fertige Protokolle als PDF und Word",
        "welcome_email_benefit4": "🔒 Sicher – Ihre Daten bleiben in der Schweiz (DSGVO-konform)",
        "welcome_email_benefit5": "🗣️ Dialektfreundlich – Schweizerdeutsch, Mundart und Akzente werden problemlos erkannt",
        "welcome_email_cta": "Starten Sie jetzt und erstellen Sie Ihr erstes Protokoll:",
        "welcome_email_help": "Fragen? Antworten Sie einfach auf diese E-Mail.",
        "whisper_language": "de",
        "gpt_language": "German (Swiss style)",
        "limit_reached_total": "⚠️ Beta-Limit erreicht: Es wurden maximal 300 Protokolle erstellt. Die Beta-Phase ist beendet.",
        "limit_reached_daily": "⚠️ Tageslimit erreicht: Sie können max. 4 Protokolle pro Tag erstellen. Bitte versuchen Sie es morgen erneut.",
        "limit_no_minutes": "⚠️ Keine Transkriptionsminuten mehr übrig. Bitte upgraden Sie Ihren Plan, um fortzufahren.",
        "remaining_today": "Verbleibende Protokolle heute",
        "beta_total_remaining": "Verbleibende Beta-Protokolle total",
        "verify_email_subject": "E-Mail bestätigen – MINU Protokoll KI",
        "verify_email_heading": "E-Mail-Adresse bestätigen",
        "verify_email_text": "Vielen Dank für Ihre Registrierung! Bitte klicken Sie auf den Button, um Ihre E-Mail-Adresse zu bestätigen und Ihr Konto zu aktivieren.",
        "verify_email_button": "E-Mail bestätigen",
        "verify_email_expiry": "Dieser Link ist 48 Stunden gültig.",
        "verify_email_ignore": "Wenn Sie sich nicht registriert haben, können Sie diese E-Mail ignorieren.",
        "verify_email_sent_title": "Überprüfen Sie Ihr Postfach",
        "verify_email_sent_text": "Wir haben eine Bestätigungs-E-Mail an <b>{email}</b> gesendet. Bitte klicken Sie auf den Link in der E-Mail, um Ihr Konto zu aktivieren.",
        "verify_email_spam_hint": "Nicht erhalten? Prüfen Sie Ihren Spam-Ordner.",
        "verify_resend_button": "E-Mail erneut senden",
        "verify_resend_wait": "Bitte warten Sie {seconds} Sekunden, bevor Sie erneut senden.",
        "verify_resend_success": "✓ Bestätigungs-E-Mail erneut gesendet!",
        "verify_success_title": "E-Mail bestätigt!",
        "verify_success_text": "Ihre E-Mail-Adresse wurde erfolgreich bestätigt. Sie können sich jetzt anmelden.",
        "verify_success_login": "Jetzt anmelden",
        "verify_expired_title": "Link abgelaufen",
        "verify_expired_text": "Dieser Bestätigungslink ist abgelaufen. Bitte registrieren Sie sich erneut.",
        "verify_invalid_title": "Ungültiger Link",
        "verify_invalid_text": "Dieser Bestätigungslink ist ungültig.",
        "error_email_not_verified": "Bitte bestätigen Sie zuerst Ihre E-Mail-Adresse. Prüfen Sie Ihr Postfach für die Bestätigungs-E-Mail."
    },
    "fr": {
        "flag": "🇫🇷",
        "lang_name": "Français",
        "title": "MINU Minutes IA",
        "subtitle": "Transformez l'audio en procès-verbaux professionnels.",
        "slogan": "En minutes au lieu d'heures.",
        "slogan_gradient": "précis, rapide, propulsé par l'IA",
        "audio_conversion_text": "Votre fichier audio sera converti en procès-verbal en quelques minutes.",
        "upload_info_title": "💡 Conseil pour de meilleurs résultats",
        "upload_info_text": "Si vous mentionnez le titre de la réunion et présentez tous les participants au début de l'enregistrement audio, ces informations seront automatiquement incluses dans le procès-verbal. Vous pouvez également ajouter ou modifier ces informations ultérieurement dans le document Word.<br><br>🌍 <b>Langue:</b> Le procès-verbal sera généré dans la langue sélectionnée – quelle que soit la langue parlée dans l'audio. Votre enregistrement peut être en anglais, allemand, français ou italien.",
        "upload_label": "Télécharger le fichier audio",
        "upload_help": "MP3, WAV, M4A, OGG, WEBM, MP4",
        "upload_hint": "MP3, WAV, M4A · Max. 200 MB",
        "processing": "Traitement...",
        "transcribing": "Transcription audio...",
        "generating": "Génération du procès-verbal...",
        "creating_pdf": "Création du PDF...",
        "creating_word": "Création du Word...",
        "done": "Terminé!",
        "file_uploaded_success": "Fichier audio téléchargé avec succès",
        "your_protocol_ready": "✓ Votre procès-verbal est prêt!",
        "download_pdf": "Télécharger PDF",
        "download_word": "Télécharger Word",
        "send_email": "Envoyer par e-mail",
        "email_placeholder": "Entrer l'adresse e-mail",
        "email_format_label": "Format:",
        "send_button": "Envoyer",
        "email_sent": "✓ E-mail envoyé!",
        "email_format_error": "Veuillez sélectionner au moins un format",
        "email_error": "Veuillez entrer une adresse e-mail",
        "show_protocol": "Afficher le procès-verbal",
        "new_protocol": "Créer un nouveau procès-verbal",
        "login_title": "🔐 MINU Minutes IA",
        "login_subtitle": "Veuillez entrer le mot de passe",
        "password_placeholder": "Mot de passe",
        "login_button": "Connexion",
        "wrong_password": "Mot de passe incorrect",
        "logout": "Déconnexion",
        "my_account": "Mon compte",
        "subscription": "Abonnement",
        "valid_until": "Valide jusqu'au",
        "minutes_remaining": "Minutes restantes",
        "upgrade_plan": "Changer de forfait",
        "beta_free_access": "Accès gratuit pour procès-verbaux - phase bêta",
        "company_label": "Entreprise / Organisation",
        "company_placeholder": "Votre entreprise ou organisation",
        "email_label": "Adresse e-mail",
        "email_placeholder_reg": "votre.email@entreprise.ch",
        "register_button": "Commencer gratuitement",
        "consent_checkbox": "Je confirme agir en tant que professionnel/entreprise et j'accepte les CGV et la politique de confidentialité",
        "privacy_link": "Politique de confidentialité",
        "error_company_required": "Veuillez entrer votre entreprise/organisation",
        "error_email_required": "Veuillez entrer une adresse e-mail valide",
        "error_email_exists": "Cet e-mail est déjà enregistré.",
        "error_downgrade_not_allowed": "Downgrade impossible. Veuillez d'abord annuler votre abonnement actuel et vous réinscrire.",
        "error_plan_already_active": "Vous avez déjà un abonnement actif avec cette adresse e-mail.",
        "error_country_blocked": "MINU-AI n'est pas disponible dans votre pays. Selon nos CGV (§5), seuls les clients de Suisse, UE, EEE et Royaume-Uni peuvent acheter.",
        "error_consent_required": "Veuillez accepter les CGV",
        "login_title": "Connexion",
        "login_subtitle": "Continuer avec votre compte existant",
        "login_button": "Se connecter",
        "login_link": "Déjà inscrit? Se connecter",
        "register_link": "Nouveau? S'inscrire",
        "error_email_not_found": "E-mail non trouvé. Veuillez d'abord vous inscrire.",
        "privacy_title": "Politique de confidentialité",
        "sidebar_title": "MINU Minutes IA",
        "sidebar_subtitle": "Générateur de procès-verbaux",
        "status": "STATUT",
        "transcript_ready": "✓ Transcription",
        "protocol_ready": "✓ Procès-verbal",
        "documents_ready": "✓ Documents",
        "ready_to_start": "_Prêt à démarrer_",
        "technology": "TECHNOLOGIE",
        "installation": "INSTALLATION",
        "install_as_app": "📱 Installer comme app",
        "start_recording": "🎙️ Démarrer l'enregistrement",
        "stop_recording": "⏹️ Arrêter l'enregistrement",
        "recording_in_progress": "Enregistrement en cours...",
        "or_upload_file": "Ou télécharger un fichier audio",
        "upload_file_button": "Télécharger fichier",
        "mic_permission_needed": "Veuillez autoriser l'accès au microphone",
        "recording_not_supported": "L'enregistrement n'est pas pris en charge dans ce navigateur",
        "recording_warning_45min": "⚠️ 45 minutes atteintes. L'enregistrement s'arrêtera automatiquement à 60 minutes.",
        "recording_stopped_limit": "Enregistrement arrêté: Maximum 60 minutes atteint.",
        "admin": "🔧 ADMIN",
        "activity_log": "📊 Journal d'activité",
        "no_activities": "Aucune activité",
        "file_size": "Taille du fichier",
        "small_file": "Petit fichier - transcription directe...",
        "large_file_splitting": "Grand fichier - division en cours...",
        "parts_created": "parties audio créées",
        "transcribing_part": "Transcription partie",
        "words_transcribed": "mots transcrits",
        "total_words": "TOTAL",
        "words": "mots",
        "characters": "caractères",
        "sending_to_mistral": "Envoi à Mistral AI...",
        "protocol_generated": "Procès-verbal généré",
        "warning_too_short": "⚠️ ATTENTION: Procès-verbal trop court!",
        "processing_log": "📋 Journal de traitement",
        "error": "Erreur",
        "retry": "Réessayer",
        "file_too_large": "Le fichier est trop volumineux",
        "maximum": "Maximum",
        "ffmpeg_needed": "ffmpeg est nécessaire. Exécutez dans le terminal: brew install ffmpeg",
        "email_subject": "Procès-verbal de réunion du",
        "email_body_greeting": "Bonjour,",
        "email_body_text": "Veuillez trouver ci-joint le procès-verbal de la réunion en PDF et Word.",
        "email_body_closing": "Cordialement",
        "email_rating_intro": "Êtes-vous satisfait de MINU Minutes IA?",
        "email_rating_thanks": "Merci pour votre avis!",
        "welcome_email_subject": "Bienvenue chez MINU Minutes IA! 🎉",
        "welcome_email_greeting": "Bienvenue chez MINU Minutes IA!",
        "welcome_email_intro": "Merci pour votre inscription. Vous avez maintenant accès à notre outil de procès-verbal propulsé par l'IA.",
        "welcome_email_benefits_title": "Vos avantages:",
        "welcome_email_benefit1": "⏱️ Gain de temps – Des procès-verbaux en minutes au lieu d'heures",
        "welcome_email_benefit2": "📱 Flexible – Enregistrez directement depuis votre téléphone ou téléchargez l'audio",
        "welcome_email_benefit3": "✨ Professionnel – Procès-verbaux prêts à l'emploi en PDF et Word",
        "welcome_email_benefit4": "🔒 Sécurisé – Vos données restent en Suisse (conforme RGPD)",
        "welcome_email_benefit5": "🗣️ Dialectes reconnus – Les dialectes régionaux et accents sont transcrits sans problème",
        "welcome_email_cta": "Commencez maintenant et créez votre premier procès-verbal:",
        "welcome_email_help": "Des questions? Répondez simplement à cet e-mail.",
        "whisper_language": "fr",
        "gpt_language": "French",
        "limit_reached_total": "⚠️ Limite bêta atteinte: Maximum 300 procès-verbaux ont été créés. La phase bêta est terminée.",
        "limit_reached_daily": "⚠️ Limite quotidienne atteinte: Vous pouvez créer max. 4 procès-verbaux par jour. Réessayez demain.",
        "limit_no_minutes": "⚠️ Plus de minutes de transcription disponibles. Veuillez mettre à niveau votre plan pour continuer.",
        "remaining_today": "Procès-verbaux restants aujourd'hui",
        "beta_total_remaining": "Procès-verbaux bêta restants au total",
        "verify_email_subject": "Confirmez votre e-mail – MINU Minutes IA",
        "verify_email_heading": "Confirmez votre adresse e-mail",
        "verify_email_text": "Merci pour votre inscription! Veuillez cliquer sur le bouton ci-dessous pour confirmer votre adresse e-mail et activer votre compte.",
        "verify_email_button": "Confirmer l'e-mail",
        "verify_email_expiry": "Ce lien est valide pendant 48 heures.",
        "verify_email_ignore": "Si vous ne vous êtes pas inscrit, vous pouvez ignorer cet e-mail.",
        "verify_email_sent_title": "Vérifiez votre boîte de réception",
        "verify_email_sent_text": "Nous avons envoyé un e-mail de confirmation à <b>{email}</b>. Veuillez cliquer sur le lien dans l'e-mail pour activer votre compte.",
        "verify_email_spam_hint": "Pas reçu? Vérifiez votre dossier spam.",
        "verify_resend_button": "Renvoyer l'e-mail",
        "verify_resend_wait": "Veuillez patienter {seconds} secondes avant de renvoyer.",
        "verify_resend_success": "✓ E-mail de confirmation renvoyé!",
        "verify_success_title": "E-mail confirmé!",
        "verify_success_text": "Votre adresse e-mail a été confirmée avec succès. Vous pouvez maintenant vous connecter.",
        "verify_success_login": "Se connecter",
        "verify_expired_title": "Lien expiré",
        "verify_expired_text": "Ce lien de confirmation a expiré. Veuillez vous inscrire à nouveau.",
        "verify_invalid_title": "Lien invalide",
        "verify_invalid_text": "Ce lien de confirmation est invalide.",
        "error_email_not_verified": "Veuillez d'abord confirmer votre adresse e-mail. Vérifiez votre boîte de réception pour l'e-mail de confirmation."
    },
    "it": {
        "flag": "🇮🇹",
        "lang_name": "Italiano",
        "title": "MINU Verbali IA",
        "subtitle": "Trasforma l'audio in verbali professionali.",
        "slogan": "In minuti invece di ore.",
        "slogan_gradient": "preciso, veloce, basato su IA",
        "audio_conversion_text": "Il tuo file audio sarà convertito in verbale in pochi minuti.",
        "upload_info_title": "💡 Consiglio per risultati migliori",
        "upload_info_text": "Se menzioni il titolo della riunione e presenti tutti i partecipanti all'inizio della registrazione audio, queste informazioni saranno automaticamente incluse nel verbale. Puoi anche aggiungere o modificare queste informazioni successivamente nel documento Word.<br><br>🌍 <b>Lingua:</b> Il verbale sarà generato nella lingua selezionata – indipendentemente dalla lingua parlata nell'audio. La tua registrazione può essere in inglese, tedesco, francese o italiano.",
        "upload_label": "Carica file audio",
        "upload_help": "MP3, WAV, M4A, OGG, WEBM, MP4",
        "upload_hint": "MP3, WAV, M4A · Max. 200 MB",
        "processing": "Elaborazione...",
        "transcribing": "Trascrizione audio...",
        "generating": "Generazione verbale...",
        "creating_pdf": "Creazione PDF...",
        "creating_word": "Creazione Word...",
        "done": "Fatto!",
        "file_uploaded_success": "File audio caricato con successo",
        "your_protocol_ready": "✓ Il tuo verbale è pronto!",
        "download_pdf": "Scarica PDF",
        "download_word": "Scarica Word",
        "send_email": "Invia via e-mail",
        "email_placeholder": "Inserisci indirizzo e-mail",
        "email_format_label": "Formato:",
        "send_button": "Invia",
        "email_sent": "✓ E-mail inviata!",
        "email_format_error": "Seleziona almeno un formato",
        "email_error": "Inserisci un indirizzo e-mail",
        "show_protocol": "Mostra verbale",
        "new_protocol": "Crea nuovo verbale",
        "login_title": "🔐 MINU Verbali IA",
        "login_subtitle": "Inserisci la password",
        "password_placeholder": "Password",
        "login_button": "Accedi",
        "wrong_password": "Password errata",
        "logout": "Esci",
        "my_account": "Il mio account",
        "subscription": "Abbonamento",
        "valid_until": "Valido fino al",
        "minutes_remaining": "Minuti rimanenti",
        "upgrade_plan": "Cambia piano",
        "beta_free_access": "Accesso gratuito per verbali - fase beta",
        "company_label": "Azienda / Organizzazione",
        "company_placeholder": "La tua azienda o organizzazione",
        "email_label": "Indirizzo e-mail",
        "email_placeholder_reg": "tua.email@azienda.ch",
        "register_button": "Inizia gratis",
        "consent_checkbox": "Confermo di agire come professionista/azienda e accetto le CGC e l'informativa sulla privacy",
        "privacy_link": "Informativa sulla privacy",
        "error_company_required": "Inserisci la tua azienda/organizzazione",
        "error_email_required": "Inserisci un indirizzo e-mail valido",
        "error_email_exists": "Questa e-mail è già registrata.",
        "error_downgrade_not_allowed": "Downgrade non possibile. Annulla prima il tuo abbonamento attuale e registrati di nuovo.",
        "error_plan_already_active": "Hai già un abbonamento attivo con questo indirizzo e-mail.",
        "error_country_blocked": "MINU-AI non è disponibile nel tuo paese. Secondo le nostre CGC (§5), solo i clienti di Svizzera, UE, SEE e Regno Unito possono acquistare.",
        "error_consent_required": "Accetta le CGC",
        "login_title": "Accedi",
        "login_subtitle": "Continua con il tuo account esistente",
        "login_button": "Accedi",
        "login_link": "Già registrato? Accedi",
        "register_link": "Nuovo? Registrati",
        "error_email_not_found": "E-mail non trovata. Registrati prima.",
        "privacy_title": "Informativa sulla privacy",
        "sidebar_title": "MINU Verbali IA",
        "sidebar_subtitle": "Generatore di verbali",
        "status": "STATO",
        "transcript_ready": "✓ Trascrizione",
        "protocol_ready": "✓ Verbale",
        "documents_ready": "✓ Documenti",
        "ready_to_start": "_Pronto per iniziare_",
        "technology": "TECNOLOGIA",
        "installation": "INSTALLAZIONE",
        "install_as_app": "📱 Installa come app",
        "start_recording": "🎙️ Avvia registrazione",
        "stop_recording": "⏹️ Termina registrazione",
        "recording_in_progress": "Registrazione in corso...",
        "or_upload_file": "Oppure carica file audio",
        "upload_file_button": "Carica file",
        "mic_permission_needed": "Consenti l'accesso al microfono",
        "recording_not_supported": "Registrazione non supportata in questo browser",
        "recording_warning_45min": "⚠️ 45 minuti raggiunti. La registrazione si fermerà automaticamente a 60 minuti.",
        "recording_stopped_limit": "Registrazione fermata: Massimo 60 minuti raggiunto.",
        "admin": "🔧 ADMIN",
        "activity_log": "📊 Registro attività",
        "no_activities": "Nessuna attività",
        "file_size": "Dimensione file",
        "small_file": "File piccolo - trascrizione diretta...",
        "large_file_splitting": "File grande - divisione in corso...",
        "parts_created": "parti audio create",
        "transcribing_part": "Trascrizione parte",
        "words_transcribed": "parole trascritte",
        "total_words": "TOTALE",
        "words": "parole",
        "characters": "caratteri",
        "sending_to_mistral": "Invio a Mistral AI...",
        "protocol_generated": "Verbale generato",
        "warning_too_short": "⚠️ ATTENZIONE: Verbale troppo corto!",
        "processing_log": "📋 Registro elaborazione",
        "error": "Errore",
        "retry": "Riprova",
        "file_too_large": "Il file è troppo grande",
        "maximum": "Massimo",
        "ffmpeg_needed": "ffmpeg è necessario. Esegui nel terminale: brew install ffmpeg",
        "email_subject": "Verbale della riunione del",
        "email_body_greeting": "Buongiorno,",
        "email_body_text": "In allegato trova il verbale della riunione in formato PDF e Word.",
        "email_body_closing": "Cordiali saluti",
        "email_rating_intro": "Quanto sei soddisfatto di MINU Verbali IA?",
        "email_rating_thanks": "Grazie per il tuo feedback!",
        "welcome_email_subject": "Benvenuto in MINU Verbali IA! 🎉",
        "welcome_email_greeting": "Benvenuto in MINU Verbali IA!",
        "welcome_email_intro": "Grazie per la registrazione. Ora hai accesso al nostro strumento di verbali basato sull'IA.",
        "welcome_email_benefits_title": "I tuoi vantaggi:",
        "welcome_email_benefit1": "⏱️ Risparmio di tempo – Verbali in minuti invece di ore",
        "welcome_email_benefit2": "📱 Flessibile – Registra direttamente dal telefono o carica l'audio",
        "welcome_email_benefit3": "✨ Professionale – Verbali pronti all'uso in PDF e Word",
        "welcome_email_benefit4": "🔒 Sicuro – I tuoi dati rimangono in Svizzera (conforme GDPR)",
        "welcome_email_benefit5": "🗣️ Dialetti riconosciuti – I dialetti regionali e gli accenti vengono trascritti senza problemi",
        "welcome_email_cta": "Inizia ora e crea il tuo primo verbale:",
        "welcome_email_help": "Domande? Rispondi semplicemente a questa e-mail.",
        "whisper_language": "it",
        "gpt_language": "Italian",
        "limit_reached_total": "⚠️ Limite beta raggiunto: Sono stati creati il massimo di 300 verbali. La fase beta è terminata.",
        "limit_reached_daily": "⚠️ Limite giornaliero raggiunto: Puoi creare max. 4 verbali al giorno. Riprova domani.",
        "limit_no_minutes": "⚠️ Nessun minuto di trascrizione rimasto. Aggiorna il tuo piano per continuare.",
        "remaining_today": "Verbali rimanenti oggi",
        "beta_total_remaining": "Verbali beta rimanenti in totale",
        "verify_email_subject": "Conferma la tua e-mail – MINU Verbali IA",
        "verify_email_heading": "Conferma il tuo indirizzo e-mail",
        "verify_email_text": "Grazie per la registrazione! Clicca sul pulsante qui sotto per confermare il tuo indirizzo e-mail e attivare il tuo account.",
        "verify_email_button": "Conferma e-mail",
        "verify_email_expiry": "Questo link è valido per 48 ore.",
        "verify_email_ignore": "Se non ti sei registrato, puoi ignorare questa e-mail.",
        "verify_email_sent_title": "Controlla la tua casella di posta",
        "verify_email_sent_text": "Abbiamo inviato un'e-mail di conferma a <b>{email}</b>. Clicca sul link nell'e-mail per attivare il tuo account.",
        "verify_email_spam_hint": "Non ricevuto? Controlla la cartella spam.",
        "verify_resend_button": "Reinvia e-mail",
        "verify_resend_wait": "Attendi {seconds} secondi prima di reinviare.",
        "verify_resend_success": "✓ E-mail di conferma reinviata!",
        "verify_success_title": "E-mail confermata!",
        "verify_success_text": "Il tuo indirizzo e-mail è stato confermato con successo. Ora puoi accedere.",
        "verify_success_login": "Accedi ora",
        "verify_expired_title": "Link scaduto",
        "verify_expired_text": "Questo link di conferma è scaduto. Registrati di nuovo.",
        "verify_invalid_title": "Link non valido",
        "verify_invalid_text": "Questo link di conferma non è valido.",
        "error_email_not_verified": "Conferma prima il tuo indirizzo e-mail. Controlla la tua casella di posta per l'e-mail di conferma."
    }
}

def get_text(key: str) -> str:
    """Holt den übersetzten Text für den aktuellen Sprachcode."""
    lang = st.session_state.get("language", "en")
    return TRANSLATIONS.get(lang, TRANSLATIONS["en"]).get(key, key)

def get_lang() -> str:
    """Gibt den aktuellen Sprachcode zurück."""
    return st.session_state.get("language", "en")

# ============================================================================
# Protokoll-Limits für Beta-Phase
# ============================================================================

BETA_TOTAL_LIMIT = 200  # Maximale Anzahl Protokolle in der Beta-Phase
DAILY_USER_LIMIT = 4    # Maximale Protokolle pro Benutzer pro Tag

# ============================================================================
# Subscription-System - Pläne und Preise
# ============================================================================

SUBSCRIPTION_PLANS = {
    "free_trial": {
        "name": "Free Trial",
        "name_de": "Kostenlose Testphase",
        "price_monthly": 0,
        "price_yearly": 0,
        "max_users": 1,
        "max_transcription_minutes": 60,  # 60 Minuten Transkription
        "trial_days": 21,     # 3 Wochen
        "features": [
            "60 Min. Transkription",
            "100% DSGVO-konform",
            "Swiss Developer 🇨🇭",
            "Unbegrenzte Protokolle",
        ],
    },
    "basic_solo": {
        "name": "Starter",
        "name_de": "Starter",
        "price_monthly": 19,
        "price_yearly": 190,  # 2 Monate gratis
        "max_users": 1,
        "max_transcription_minutes": 180,  # 180 Min/Monat
        "features": [
            "180 Min. Transkription/Monat",
            "100% DSGVO-konform",
            "Swiss Developer 🇨🇭",
            "Unbegrenzte Protokolle",
            "Priority Support",
        ],
    },
    "team": {
        "name": "Pro",
        "name_de": "Pro",
        "price_monthly": 79,
        "price_yearly": 790,  # 2 Monate gratis
        "max_users": 5,
        "max_transcription_minutes": 600,  # 600 Min/Monat
        "features": [
            "600 Min. Transkription/Monat",
            "100% DSGVO-konform",
            "Swiss Developer 🇨🇭",
            "Unbegrenzte Protokolle",
            "Bis zu 5 Benutzer",
            "Priority Support",
            "Pro-Verwaltung",
        ],
    }
}

# ============================================================================
# Stripe Payment Integration
# ============================================================================

def get_stripe_price_id(plan: str, billing_cycle: str) -> str:
    """Holt die Stripe Preis-ID für einen Plan und Abrechnungszyklus."""
    price_map = {
        ("basic_solo", "monthly"): "STRIPE_PRICE_STARTER_MONTHLY",
        ("basic_solo", "yearly"): "STRIPE_PRICE_STARTER_YEARLY",
        ("team", "monthly"): "STRIPE_PRICE_PRO_MONTHLY",
        ("team", "yearly"): "STRIPE_PRICE_PRO_YEARLY",
    }
    env_key = price_map.get((plan, billing_cycle))
    if env_key:
        return os.getenv(env_key, "")
    return ""

def create_stripe_checkout_session(user_email: str, plan: str, billing_cycle: str = "monthly") -> str:
    """
    Erstellt eine Stripe Checkout Session für ein Abonnement.

    Args:
        user_email: E-Mail des Benutzers
        plan: 'basic_solo' (Starter) oder 'team' (PRO)
        billing_cycle: 'monthly' oder 'yearly'

    Returns:
        URL zur Stripe Checkout-Seite
    """
    # Stripe API Key setzen
    stripe.api_key = os.getenv("STRIPE_SECRET_KEY", "")
    if not stripe.api_key:
        print("STRIPE_SECRET_KEY nicht gefunden in .env")
        return None

    # Preis-ID basierend auf Plan und Abrechnungszyklus
    price_id = get_stripe_price_id(plan, billing_cycle)

    if not price_id:
        print(f"Keine Preis-ID gefunden für Plan: {plan}, Zyklus: {billing_cycle}")
        return None

    try:
        # Basis-URL für Erfolg/Abbruch (automatisch localhost oder production)
        base_url = os.getenv("APP_BASE_URL", "http://localhost:8502")

        checkout_session = stripe.checkout.Session.create(
            payment_method_types=['card'],
            line_items=[{
                'price': price_id,
                'quantity': 1,
            }],
            mode='subscription',
            success_url=f"{base_url}/?payment=success&session_id={{CHECKOUT_SESSION_ID}}",
            cancel_url=f"{base_url}/?payment=cancelled",
            customer_email=user_email,
            # Rechnungsadresse erforderlich - Länderblockierung via Stripe Radar Rules
            billing_address_collection='required',
            metadata={
                'user_email': user_email,
                'plan': plan,
                'billing_cycle': billing_cycle,
            },
            subscription_data={
                'metadata': {
                    'user_email': user_email,
                    'plan': plan,
                }
            }
        )
        return checkout_session.url
    except stripe.error.StripeError as e:
        print(f"Stripe Fehler: {e}")
        return None

def verify_stripe_payment(session_id: str) -> dict:
    """
    Verifiziert eine Stripe-Zahlung anhand der Session-ID.

    Returns:
        dict mit Zahlungsinformationen oder None bei Fehler
    """
    # Stripe API Key setzen
    stripe.api_key = os.getenv("STRIPE_SECRET_KEY", "")
    if not stripe.api_key:
        return {"success": False}

    try:
        session = stripe.checkout.Session.retrieve(session_id)

        if session.payment_status == "paid":
            return {
                "success": True,
                "user_email": session.metadata.get("user_email"),
                "plan": session.metadata.get("plan"),
                "billing_cycle": session.metadata.get("billing_cycle"),
                "subscription_id": session.subscription,
                "customer_id": session.customer,
            }
    except stripe.error.StripeError as e:
        print(f"Stripe Verifizierungsfehler: {e}")

    return {"success": False}


def create_stripe_portal_session(customer_id: str, return_url: str) -> str:
    """
    Erstellt eine Stripe Customer Portal Session für Abo-Verwaltung.

    Args:
        customer_id: Stripe Customer ID
        return_url: URL wohin der Kunde nach dem Portal zurückgeleitet wird

    Returns:
        URL zum Stripe Customer Portal oder None bei Fehler
    """
    stripe.api_key = os.getenv("STRIPE_SECRET_KEY", "")
    if not stripe.api_key or not customer_id:
        return None

    try:
        session = stripe.billing_portal.Session.create(
            customer=customer_id,
            return_url=return_url,
        )
        return session.url
    except stripe.error.StripeError as e:
        print(f"Stripe Portal Fehler: {e}")
        return None


def upgrade_user_subscription(user_email: str, plan: str, billing_cycle: str, subscription_id: str = None, company: str = None):
    """
    Aktualisiert das Abonnement eines Benutzers nach erfolgreicher Zahlung.
    Bestehende Subscription wird aktualisiert, keine Duplikate.
    """
    subs = load_subscriptions()
    email_key = user_email.lower().strip()  # Normalisierte Email als Key

    # Bestehende Daten beibehalten oder neu initialisieren
    if email_key not in subs:
        subs[email_key] = {
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "protocols_created": 0,
        }

    # Plan-Konfiguration holen
    plan_config = SUBSCRIPTION_PLANS.get(plan, SUBSCRIPTION_PLANS["free_trial"])

    # Subscription aktualisieren (bestehende Werte beibehalten)
    subs[email_key].update({
        "plan": plan,
        "status": "active",
        "billing_cycle": billing_cycle,
        "stripe_subscription_id": subscription_id,
        "upgraded_at": datetime.now().isoformat(),
        "transcription_minutes_used": 0.0,  # Reset bei Upgrade
    })

    # Company aktualisieren falls übergeben
    if company:
        subs[email_key]["company"] = company

    save_subscriptions(subs)
    print(f"[STRIPE] Upgrade erfolgreich: {email_key} -> {plan} ({billing_cycle})")

def get_subscriptions_file_path():
    """Gibt den Pfad zur Subscriptions-Datei zurück."""
    return PROJECT_ROOT / "subscriptions.json"

def load_subscriptions():
    """Lädt alle Subscription-Daten."""
    subs_file = get_subscriptions_file_path()
    try:
        if subs_file.exists():
            with open(subs_file, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception as e:
        print(f"Fehler beim Laden der Subscriptions: {e}")
    return {}

def save_subscriptions(data):
    """Speichert die Subscription-Daten."""
    subs_file = get_subscriptions_file_path()
    try:
        with open(subs_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Fehler beim Speichern der Subscriptions: {e}")

def get_user_subscription(user_email: str) -> dict:
    """Holt die Subscription-Daten eines Benutzers (case-insensitive)."""
    if not user_email:
        print(f"[DEBUG get_user_subscription] No user_email provided")
        return None
    subs = load_subscriptions()
    email_lower = user_email.lower().strip()
    print(f"[DEBUG get_user_subscription] Looking for: '{email_lower}' in keys: {list(subs.keys())}")
    # Case-insensitive lookup
    if email_lower in subs:
        print(f"[DEBUG get_user_subscription] FOUND: {subs[email_lower]}")
        return subs[email_lower]
    # Fallback: Suche case-insensitive durch alle Keys
    for key in subs:
        if key.lower().strip() == email_lower:
            return subs[key]
    return None

def create_trial_subscription(user_email: str, company: str = ""):
    """Erstellt eine neue Trial-Subscription für einen Benutzer."""
    subs = load_subscriptions()
    email_key = user_email.lower().strip()  # Normalisierte Email als Key

    now = datetime.now()
    trial_end = now + timedelta(days=SUBSCRIPTION_PLANS["free_trial"]["trial_days"])

    subs[email_key] = {
        "plan": "free_trial",
        "status": "active",  # active, expired, cancelled
        "company": company,
        "created_at": now.strftime("%Y-%m-%d %H:%M:%S"),
        "trial_start": now.strftime("%Y-%m-%d"),
        "trial_end": trial_end.strftime("%Y-%m-%d"),
        "transcription_minutes_used": 0.0,  # Verbrauchte Transkriptionsminuten
        "protocols_created": 0,  # Anzahl erstellter Protokolle (für Statistik)
        "stripe_customer_id": None,
        "stripe_subscription_id": None,
        "billing_cycle": None,  # monthly, yearly
    }

    save_subscriptions(subs)
    return subs[email_key]

def add_transcription_minutes(user_email: str, minutes: float):
    """Fügt verbrauchte Transkriptionsminuten für einen Benutzer hinzu."""
    subs = load_subscriptions()
    email_key = user_email.lower().strip()  # Normalisierte Email als Key
    if email_key in subs:
        subs[email_key]["transcription_minutes_used"] = subs[email_key].get("transcription_minutes_used", 0.0) + minutes
        subs[email_key]["protocols_created"] = subs[email_key].get("protocols_created", 0) + 1
        save_subscriptions(subs)

def get_remaining_minutes(user_email: str) -> float:
    """Gibt die verbleibenden Transkriptionsminuten für einen Benutzer zurück."""
    sub = get_user_subscription(user_email)
    if not sub:
        return 0.0

    plan = sub.get("plan", "free_trial")
    plan_config = SUBSCRIPTION_PLANS.get(plan, SUBSCRIPTION_PLANS["free_trial"])
    max_minutes = plan_config.get("max_transcription_minutes", 60)
    used_minutes = sub.get("transcription_minutes_used", 0.0)

    return max(0.0, max_minutes - used_minutes)

def check_subscription_status(user_email: str) -> dict:
    """
    Prüft den Subscription-Status eines Benutzers.

    Returns:
        dict: {
            "can_create": bool,
            "reason": str or None,
            "plan": str,
            "status": str,
            "trial_days_left": int or None,
            "minutes_left": float or None,
            "minutes_used": float,
            "max_minutes": float,
            "needs_upgrade": bool
        }
    """
    sub = get_user_subscription(user_email)

    # Kein Abo gefunden - neuer Benutzer
    if not sub:
        return {
            "can_create": False,
            "reason": "no_subscription",
            "plan": None,
            "status": None,
            "trial_days_left": None,
            "minutes_left": None,
            "minutes_used": 0,
            "max_minutes": 0,
            "needs_upgrade": False
        }

    plan = sub.get("plan", "free_trial")
    plan_config = SUBSCRIPTION_PLANS.get(plan, SUBSCRIPTION_PLANS["free_trial"])

    # Transkriptionsminuten prüfen
    minutes_used = sub.get("transcription_minutes_used", 0.0)
    max_minutes = plan_config.get("max_transcription_minutes", 60)
    minutes_left = max(0.0, max_minutes - minutes_used)

    # Free Trial prüfen
    if plan == "free_trial":
        # Trial-Zeit prüfen
        trial_end = datetime.strptime(sub["trial_end"], "%Y-%m-%d")
        now = datetime.now()
        days_left = (trial_end - now).days

        # Trial abgelaufen (Zeit)
        if days_left < 0:
            return {
                "can_create": False,
                "reason": "trial_expired_time",
                "plan": plan,
                "status": "expired",
                "trial_days_left": 0,
                "minutes_left": minutes_left,
                "minutes_used": minutes_used,
                "max_minutes": max_minutes,
                "needs_upgrade": True
            }

        # Transkriptionsminuten aufgebraucht
        if minutes_left <= 0:
            return {
                "can_create": False,
                "reason": "trial_expired_minutes",
                "plan": plan,
                "status": "expired",
                "trial_days_left": days_left,
                "minutes_left": 0,
                "minutes_used": minutes_used,
                "max_minutes": max_minutes,
                "needs_upgrade": True
            }

        # Trial aktiv
        return {
            "can_create": True,
            "reason": None,
            "plan": plan,
            "status": "active",
            "trial_days_left": days_left,
            "minutes_left": minutes_left,
            "minutes_used": minutes_used,
            "max_minutes": max_minutes,
            "needs_upgrade": False
        }

    # Bezahltes Abo prüfen
    if sub.get("status") == "active":
        # Minuten aufgebraucht bei bezahltem Abo
        if minutes_left <= 0:
            return {
                "can_create": False,
                "reason": "minutes_exhausted",
                "plan": plan,
                "status": "active",
                "trial_days_left": None,
                "minutes_left": 0,
                "minutes_used": minutes_used,
                "max_minutes": max_minutes,
                "needs_upgrade": True
            }

        return {
            "can_create": True,
            "reason": None,
            "plan": plan,
            "status": "active",
            "trial_days_left": None,
            "minutes_left": minutes_left,
            "minutes_used": minutes_used,
            "max_minutes": max_minutes,
            "needs_upgrade": False
        }

    # Abo gekündigt oder abgelaufen
    return {
        "can_create": False,
        "reason": "subscription_inactive",
        "plan": plan,
        "status": sub.get("status", "inactive"),
        "trial_days_left": None,
        "minutes_left": None,
        "minutes_used": minutes_used,
        "max_minutes": max_minutes,
        "needs_upgrade": True
    }

def upgrade_subscription(user_email: str, new_plan: str, billing_cycle: str,
                         stripe_customer_id: str = None, stripe_subscription_id: str = None):
    """Upgraded eine Subscription zu einem bezahlten Plan."""
    subs = load_subscriptions()

    if user_email not in subs:
        return False

    subs[user_email]["plan"] = new_plan
    subs[user_email]["status"] = "active"
    subs[user_email]["billing_cycle"] = billing_cycle
    subs[user_email]["upgraded_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    if stripe_customer_id:
        subs[user_email]["stripe_customer_id"] = stripe_customer_id
    if stripe_subscription_id:
        subs[user_email]["stripe_subscription_id"] = stripe_subscription_id

    save_subscriptions(subs)
    return True


@st.dialog("Free Plan Info", width="small")
def show_free_plan_info_dialog():
    """Zeigt den Free Plan Info Dialog."""
    lang = st.session_state.get("language", "en")

    # Übersetzungen
    if lang == "de":
        info_title = "Free Plan Inklusive"
        features = SUBSCRIPTION_PLANS["free_trial"]["features"]
        cta_text = "OK"
        info_text = "Kostenlos registrieren und direkt mit dem Free-Plan loslegen."
    elif lang == "fr":
        info_title = "Plan Gratuit Inclus"
        features = [
            "60 min. de transcription",
            "100% conforme RGPD",
            "Swiss Developer 🇨🇭",
            "Protocoles illimités",
        ]
        cta_text = "OK"
        info_text = "Inscrivez-vous gratuitement et commencez immédiatement."
    elif lang == "it":
        info_title = "Piano Gratuito Incluso"
        features = [
            "60 min. di trascrizione",
            "100% conforme GDPR",
            "Swiss Developer 🇨🇭",
            "Protocolli illimitati",
        ]
        cta_text = "OK"
        info_text = "Registrati gratuitamente e inizia subito."
    else:  # en
        info_title = "Free Plan Included"
        features = [
            "60 min. transcription",
            "100% GDPR compliant",
            "Swiss Developer 🇨🇭",
            "Unlimited protocols",
        ]
        cta_text = "OK"
        info_text = "Register for free and start immediately."

    # Feature list HTML
    features_html = "".join([f'<div style="display: flex; align-items: center; gap: 10px; margin-bottom: 12px;"><span style="color: #34c759; font-size: 16px;">✓</span><span style="color: #333; font-size: 15px;">{f}</span></div>' for f in features])

    st.markdown(f"""
    <div style="text-align: center; margin-bottom: 16px;">
        <div style="font-size: 24px; font-weight: 700; color: #1d1d1f;">{info_title}</div>
        <div style="font-size: 14px; color: #6b7280; margin-top: 8px;">{info_text}</div>
    </div>
    <div style="margin: 24px 0;">
        {features_html}
    </div>
    """, unsafe_allow_html=True)

    if st.button(cta_text, use_container_width=True, type="primary"):
        st.query_params.clear()
        st.rerun()


def render_free_plan_info_button():
    """Rendert den Info-Button für den Free Plan als st.expander."""
    lang = st.session_state.get("language", "en")

    # Übersetzungen
    if lang == "de":
        info_title = "ℹ️ Free Plan Inklusive"
        features = ["60 Min. Transkription", "100% DSGVO-konform", "Swiss Developer 🇨🇭", "Unbegrenzte Protokolle"]
    elif lang == "fr":
        info_title = "ℹ️ Plan Gratuit Inclus"
        features = ["60 min. de transcription", "100% conforme RGPD", "Swiss Developer 🇨🇭", "Protocoles illimités"]
    elif lang == "it":
        info_title = "ℹ️ Piano Gratuito Incluso"
        features = ["60 min. di trascrizione", "100% conforme GDPR", "Swiss Developer 🇨🇭", "Protocolli illimitati"]
    else:
        info_title = "ℹ️ Free Plan Included"
        features = ["60 min. transcription", "100% GDPR compliant", "Swiss Developer 🇨🇭", "Unlimited protocols"]

    # Expander für die Info
    with st.expander(info_title, expanded=False):
        for f in features:
            st.markdown(f"✓ {f}")

def get_usage_file_path():
    """Gibt den Pfad zur Usage-Datei zurück."""
    return PROJECT_ROOT / "protocol_usage.json"

def load_usage_data():
    """Lädt die Nutzungsdaten aus der JSON-Datei."""
    usage_file = get_usage_file_path()
    try:
        if usage_file.exists():
            with open(usage_file, 'r') as f:
                return json.load(f)
    except Exception as e:
        print(f"Fehler beim Laden der Usage-Daten: {e}")

    # Standard-Struktur zurückgeben
    return {
        "total_count": 0,
        "daily_usage": {}
    }

def save_usage_data(data):
    """Speichert die Nutzungsdaten in die JSON-Datei."""
    usage_file = get_usage_file_path()
    try:
        with open(usage_file, 'w') as f:
            json.dump(data, f, indent=2)
    except Exception as e:
        print(f"Fehler beim Speichern der Usage-Daten: {e}")

def get_today_string():
    """Gibt das heutige Datum als String zurück."""
    return datetime.now().strftime("%Y-%m-%d")

def check_protocol_limits(user_email: str) -> tuple:
    """
    Prüft, ob die Protokoll-Limits erreicht sind.

    Returns:
        tuple: (can_create: bool, error_message: str or None, remaining_today: int, total_remaining: int)
    """
    # Zuerst: Transkriptionsminuten prüfen
    remaining_minutes = get_remaining_minutes(user_email)
    if remaining_minutes <= 0:
        return (False, get_text("limit_no_minutes"), 0, 0)

    data = load_usage_data()
    today = get_today_string()

    # Gesamtlimit prüfen
    total_count = data.get("total_count", 0)
    total_remaining = BETA_TOTAL_LIMIT - total_count

    if total_count >= BETA_TOTAL_LIMIT:
        return (False, get_text("limit_reached_total"), 0, 0)

    # Tageslimit für Benutzer prüfen
    daily_usage = data.get("daily_usage", {})
    user_usage = daily_usage.get(user_email, {})
    today_count = user_usage.get(today, 0)
    remaining_today = DAILY_USER_LIMIT - today_count

    if today_count >= DAILY_USER_LIMIT:
        return (False, get_text("limit_reached_daily"), 0, total_remaining)

    return (True, None, remaining_today, total_remaining)

def increment_protocol_count(user_email: str):
    """Erhöht den Protokoll-Zähler für einen Benutzer."""
    data = load_usage_data()
    today = get_today_string()

    # Gesamtzähler erhöhen
    data["total_count"] = data.get("total_count", 0) + 1

    # Täglichen Zähler für Benutzer erhöhen
    if "daily_usage" not in data:
        data["daily_usage"] = {}
    if user_email not in data["daily_usage"]:
        data["daily_usage"][user_email] = {}

    data["daily_usage"][user_email][today] = data["daily_usage"][user_email].get(today, 0) + 1

    save_usage_data(data)

    # Log-Eintrag
    log_activity("Protokoll erstellt", f"User: {user_email}, Total: {data['total_count']}/{BETA_TOTAL_LIMIT}")

# Typische ffmpeg-Pfade auf macOS
FFMPEG_PATHS = [
    "/opt/homebrew/bin/ffmpeg",  # Apple Silicon Homebrew
    "/usr/local/bin/ffmpeg",      # Intel Homebrew
    "/usr/bin/ffmpeg",            # System
]

def find_ffmpeg():
    """Findet ffmpeg auf dem System oder im Projektordner."""
    # Erst im Projektordner suchen (falls PROJECT_ROOT existiert)
    try:
        local_ffmpeg = Path(__file__).resolve().parent / "ffmpeg"
        if local_ffmpeg.exists() and os.access(str(local_ffmpeg), os.X_OK):
            return str(local_ffmpeg)
    except:
        pass
    # Im PATH suchen
    path = shutil.which("ffmpeg")
    if path:
        return path
    # Bekannte Pfade prüfen
    for p in FFMPEG_PATHS:
        if os.path.isfile(p) and os.access(p, os.X_OK):
            return p
    return None

def install_ffmpeg_brew():
    """Installiert ffmpeg über Homebrew."""
    brew_paths = ["/opt/homebrew/bin/brew", "/usr/local/bin/brew"]
    brew = None
    for bp in brew_paths:
        if os.path.isfile(bp):
            brew = bp
            break
    if brew:
        try:
            result = subprocess.run([brew, "install", "ffmpeg"],
                                   capture_output=True, timeout=600)
            return result.returncode == 0
        except:
            pass
    return False

def get_ffmpeg_path():
    """Gibt den ffmpeg-Pfad zurück."""
    return find_ffmpeg()

# ffmpeg beim Start suchen
FFMPEG_PATH = find_ffmpeg()
FFMPEG_AVAILABLE = FFMPEG_PATH is not None

# .env laden (für lokale Entwicklung)
PROJECT_ROOT = Path(__file__).resolve().parent
load_dotenv(PROJECT_ROOT / ".env")

# Logo-Pfad (für App Logo)
LOGO_PATH = PROJECT_ROOT / "ICON.png"
LOGO_AVAILABLE = LOGO_PATH.exists()


def get_secret(key: str, default: str = "") -> str:
    """Holt Secret aus Streamlit Cloud oder .env."""
    try:
        return st.secrets.get(key, os.getenv(key, default))
    except Exception:
        return os.getenv(key, default)


# Konfiguration
AUDIO_EXTENSIONS = (".mp3", ".wav", ".m4a", ".ogg", ".webm", ".mp4", ".mpeg", ".mpga")
WHISPER_CHUNK_SIZE = 24 * 1024 * 1024  # 24 MB (Whisper Limit ist 25 MB)
CHUNK_DURATION_MS = 10 * 60 * 1000  # 10 Minuten pro Chunk
MAX_FILE_SIZE = 200 * 1024 * 1024  # Immer 200 MB erlauben
MAX_FILE_SIZE_MB = 200
MAX_AUDIO_DURATION_MINUTES = 120  # Maximale Audiodauer in Minuten

# ============================================================================
# PWA (Progressive Web App) Konfiguration
# ============================================================================

PWA_META_TAGS = """
<!-- PWA Meta Tags -->
<link rel="manifest" href="/static/manifest.json">
<meta name="theme-color" content="#4F46E5">
<meta name="apple-mobile-web-app-capable" content="yes">
<meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
<meta name="apple-mobile-web-app-title" content="MINU Minutes AI">
<link rel="apple-touch-icon" href="/static/icon-192.png">
<link rel="apple-touch-icon" sizes="192x192" href="/static/icon-192.png">
<link rel="apple-touch-icon" sizes="512x512" href="/static/icon-512.png">
<meta name="mobile-web-app-capable" content="yes">
<meta name="application-name" content="MINU Minutes AI">
<meta name="msapplication-TileColor" content="#4F46E5">
<meta name="msapplication-TileImage" content="/static/icon-192.png">

<!-- SEO Meta Tags -->
<meta name="description" content="Protokoll AI – Verwandle Audio-Aufnahmen in professionelle Meeting-Protokolle mit KI. Transkription mit Whisper, Protokollerstellung mit Mistral AI. DSGVO-konform, EU-basiert.">
<meta name="keywords" content="Protokoll AI, Meeting Protokoll, MINU-AI, KI Transkription, Whisper, Sitzungsprotokoll, Audio zu Text, Meeting Notes, Automatisches Protokoll">
<meta name="author" content="MINU-AI">
<meta name="robots" content="index, follow">
<meta name="language" content="de-CH">

<!-- Open Graph (Facebook, LinkedIn, WhatsApp) -->
<meta property="og:type" content="website">
<meta property="og:title" content="Protokoll AI – Meeting-Protokolle per KI erstellen">
<meta property="og:description" content="Audio hochladen, KI transkribiert und erstellt ein professionelles Protokoll. Export als PDF und Word. DSGVO-konform.">
<meta property="og:site_name" content="MINU-AI">
<meta property="og:locale" content="de_CH">
<meta property="og:image" content="./static/icon-512.png">

<!-- Twitter Card -->
<meta name="twitter:card" content="summary">
<meta name="twitter:title" content="Protokoll AI – Meeting-Protokolle per KI">
<meta name="twitter:description" content="Audio hochladen, KI erstellt professionelle Protokolle. PDF & Word Export. Swiss Developer.">
<meta name="twitter:image" content="./static/icon-512.png">

<!-- Structured Data: SoftwareApplication (JSON-LD) -->
<script type="application/ld+json">
{
  "@context": "https://schema.org",
  "@type": "SoftwareApplication",
  "name": "Protokoll AI",
  "alternateName": "MINU Minutes AI",
  "description": "KI-gestützte Erstellung von Meeting-Protokollen aus Audio-Aufnahmen. Transkription mit OpenAI Whisper, Protokollerstellung mit Mistral AI (EU). Export als PDF und DOCX.",
  "applicationCategory": "BusinessApplication",
  "operatingSystem": "Web",
  "inLanguage": ["de", "en", "fr", "it"],
  "offers": {
    "@type": "Offer",
    "price": "0",
    "priceCurrency": "CHF",
    "availability": "https://schema.org/InStock"
  },
  "creator": {
    "@type": "Organization",
    "name": "MINU-AI",
    "url": "https://minu-ai.ch"
  },
  "featureList": [
    "Audio-Transkription mit OpenAI Whisper",
    "Protokollerstellung mit Mistral AI (EU-basiert)",
    "Export als PDF und DOCX",
    "E-Mail-Versand",
    "4 Sprachen (DE, EN, FR, IT)",
    "DSGVO-konform",
    "Progressive Web App (PWA)"
  ]
}
</script>
"""

PWA_SERVICE_WORKER = """
<script>
    // PWA Service Worker Registration
    if ('serviceWorker' in navigator) {
        window.addEventListener('load', function() {
            navigator.serviceWorker.register('./static/service-worker.js')
                .then(function(registration) {
                    console.log('MINU Minutes AI ServiceWorker registered:', registration.scope);
                })
                .catch(function(error) {
                    console.log('MINU Minutes AI ServiceWorker registration failed:', error);
                });
        });
    }

    // PWA Install Prompt
    let deferredPrompt;
    window.addEventListener('beforeinstallprompt', (e) => {
        e.preventDefault();
        deferredPrompt = e;
        // Show install button in UI if needed
        console.log('MINU Minutes AI can be installed as PWA');
    });
</script>
"""

# ============================================================================
# Custom CSS - Apple-Style minimalistisches Design
# ============================================================================

CUSTOM_CSS = """
<style>
    /* Apple-Style: Clean, minimal, optimiert für Mobile */

    /* ============================================
       Gradient-Style für ALLE Buttons (wie Monthly Toggle)
       ============================================ */
    .stButton button,
    .stButton > button,
    [data-testid="stButton"] button,
    [data-testid="baseButton-primary"],
    [data-testid="baseButton-secondary"],
    button[kind="primary"],
    button[kind="secondary"],
    .stButton button[kind="primary"],
    .stButton button[kind="secondary"] {
        background: linear-gradient(135deg, #7c3aed 0%, #2dd4bf 100%) !important;
        background-color: transparent !important;
        border: none !important;
        color: white !important;
        font-weight: 600 !important;
        border-radius: 25px !important;
        padding: 0.75rem 1.5rem !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 4px 15px rgba(124, 58, 237, 0.3) !important;
    }

    .stButton button:hover,
    .stButton > button:hover,
    [data-testid="stButton"] button:hover,
    [data-testid="baseButton-primary"]:hover,
    [data-testid="baseButton-secondary"]:hover {
        background: linear-gradient(135deg, #6d28d9 0%, #14b8a6 100%) !important;
        box-shadow: 0 6px 20px rgba(124, 58, 237, 0.4) !important;
        transform: translateY(-1px) !important;
        border: none !important;
    }

    .stButton button:active,
    .stButton > button:active {
        transform: translateY(0) !important;
    }

    /* Override Streamlit's default button colors */
    .stButton button::before,
    .stButton button::after {
        display: none !important;
    }

    /* Hauptcontainer - weniger Padding für Mobile */
    .main .block-container {
        padding-top: 1rem;
        padding-bottom: 2rem;
        max-width: 680px;
    }

    /* Mobile-optimiert */
    @media (max-width: 768px) {
        .main .block-container {
            padding-top: 0.5rem;
            padding-bottom: 1rem;
        }
        h1 {
            font-size: 1.8rem !important;
            line-height: 1.2 !important;
        }
    }

    /* Streamlit Header und unnötigen Platz verstecken */
    header[data-testid="stHeader"] {
        display: none !important;
    }

    #MainMenu {
        display: none !important;
    }

    .stApp > header {
        display: none !important;
    }

    /* Entferne Top-Margin auf Mobile */
    @media (max-width: 768px) {
        .stApp {
            margin-top: -3rem !important;
        }
        .main > div:first-child {
            padding-top: 0 !important;
        }
        [data-testid="stAppViewContainer"] {
            padding-top: 0 !important;
        }
        [data-testid="stVerticalBlock"] {
            gap: 0.5rem !important;
        }
    }

    /* Grosse, klare Überschriften */
    h1 {
        font-size: 3rem !important;
        font-weight: 700 !important;
        letter-spacing: -0.02em;
        color: #1d1d1f !important;
        text-align: center;
        margin-bottom: 0.5rem !important;
    }

    h2, h3 {
        font-weight: 600 !important;
        color: #1d1d1f !important;
        letter-spacing: -0.01em;
    }

    /* Subheadline */
    .main p {
        color: #86868b;
        font-size: 1.1rem;
        line-height: 1.5;
        text-align: center;
    }

    /* Zentrierte Checkbox */
    .stCheckbox {
        display: flex;
        justify-content: center;
    }
    .stCheckbox label {
        text-align: center;
    }
    .stCheckbox label p {
        font-size: 14px !important;
        text-align: center !important;
    }

    /* Apple-Style Button - Blau, abgerundet */
    .stButton > button {
        background: #0071e3 !important;
        color: white !important;
        border: none !important;
        border-radius: 980px !important;
        padding: 12px 24px !important;
        font-size: 17px !important;
        font-weight: 400 !important;
        transition: all 0.3s ease !important;
        min-height: 50px;
    }

    .stButton > button:hover {
        background: #0077ED !important;
        transform: scale(1.02);
    }

    .stButton > button:disabled {
        background: #d2d2d7 !important;
        color: #86868b !important;
    }

    /* Download Buttons - Grün */
    .stDownloadButton > button {
        background: #34c759 !important;
        border: none !important;
        border-radius: 980px !important;
        font-size: 17px !important;
        min-height: 50px;
    }

    .stDownloadButton > button:hover {
        background: #30d158 !important;
    }

    /* File Uploader - Clean & Zentriert */
    .stFileUploader {
        margin: 2rem auto;
        text-align: center;
    }

    .stFileUploader label {
        font-size: 1rem;
        color: #1d1d1f;
        text-align: center;
    }

    /* Text Inputs - Zentriert mit dunklerem Hintergrund */
    .stTextInput {
        max-width: 400px;
        margin-left: auto;
        margin-right: auto;
    }

    .stTextInput input {
        background-color: #f5f5f7 !important;
        border: 1px solid #d2d2d7 !important;
        border-radius: 12px !important;
        padding: 12px !important;
        font-size: 16px !important;
        text-align: center !important;
    }

    .stTextInput input::placeholder {
        text-align: center !important;
    }

    .stTextInput input:focus {
        border-color: #0071e3 !important;
        box-shadow: 0 0 0 4px rgba(0, 113, 227, 0.1) !important;
    }

    /* Text Areas */
    .stTextArea textarea {
        border-radius: 12px;
        border: 1px solid #d2d2d7;
        font-size: 15px;
        padding: 12px;
    }

    .stTextArea textarea:focus {
        border-color: #0071e3;
        box-shadow: 0 0 0 4px rgba(0, 113, 227, 0.1);
    }

    /* Expander - Clean */
    .streamlit-expanderHeader {
        font-weight: 600;
        color: #1d1d1f;
        border-radius: 12px;
    }

    /* Sidebar - Minimal */
    [data-testid="stSidebar"] {
        background: #fbfbfd;
        border-right: 1px solid #d2d2d7;
    }

    [data-testid="stSidebar"] h1,
    [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] h3 {
        font-size: 1.2rem !important;
        text-align: left !important;
    }

    /* Alerts - Soft */
    .stAlert {
        border-radius: 12px;
        border: none;
    }

    /* Progress Bar */
    .stProgress > div > div {
        background: #0071e3 !important;
        border-radius: 10px;
    }

    /* Metric Cards */
    [data-testid="metric-container"] {
        background: #f5f5f7;
        border-radius: 18px;
        padding: 1.25rem;
        border: none;
    }

    /* Divider - Subtle */
    hr {
        border-color: #d2d2d7;
        margin: 2.5rem 0;
        opacity: 0.5;
    }

    /* Links */
    a {
        color: #0066cc;
        text-decoration: none;
    }

    a:hover {
        text-decoration: underline;
    }

    /* Success/Info Messages */
    .stSuccess, .stInfo {
        border-radius: 12px;
    }

    /* Hide Streamlit Branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}

    /* Hide GitHub icon and header buttons */
    header {visibility: hidden;}
    .stApp [data-testid="stToolbar"] {display: none !important;}
    .stApp [data-testid="stDecoration"] {display: none !important;}
    .stDeployButton {display: none !important;}

    /* Custom Footer */
    .custom-footer {
        position: fixed;
        bottom: 0;
        left: 0;
        width: 100%;
        background: #fbfbfd;
        border-top: 1px solid #d2d2d7;
        padding: 10px 0;
        text-align: center;
        font-size: 12px;
        color: #86868b;
        z-index: 999;
    }

    /* ============================================
       Moderne Kreisförmige Fortschrittsanzeige
       ============================================ */

    /* Container für Animation */
    .processing-animation {
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        padding: 2rem 1.5rem;
        background: #ffffff;
        border-radius: 16px;
        margin: 1rem 0;
        border: 1px solid #e5e5e5;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
    }

    /* Upload Erfolg Meldung */
    .upload-success {
        background: linear-gradient(90deg, #7c3aed, #2dd4bf);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        font-size: 15px;
        font-weight: 600;
        margin-bottom: 1rem;
        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    }

    /* Kreisförmiger Fortschrittsring */
    .circular-progress {
        position: relative;
        width: 140px;
        height: 140px;
        margin-bottom: 1rem;
    }

    .circular-progress svg {
        transform: rotate(-90deg);
        width: 140px;
        height: 140px;
    }

    .circular-progress .bg-circle {
        fill: none;
        stroke: #f0f0f0;
        stroke-width: 8;
    }

    .circular-progress .progress-circle {
        fill: none;
        stroke: url(#progressGradient);
        stroke-width: 8;
        stroke-linecap: round;
        stroke-dasharray: 390;
        stroke-dashoffset: 390;
        transition: stroke-dashoffset 0.5s ease;
    }

    .circular-progress .percentage {
        position: absolute;
        top: 50%;
        left: 50%;
        transform: translate(-50%, -50%);
        font-size: 32px;
        font-weight: 700;
        color: #1d1d1f;
        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    }

    /* Datei-Info */
    .file-info {
        color: #86868b;
        font-size: 14px;
        margin-bottom: 0.5rem;
        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    }

    /* Status Text */
    .status-text {
        color: #1d1d1f;
        font-size: 16px;
        text-align: center;
        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
        margin-top: 0.5rem;
    }

    /* Kleiner rotierender Spinner - zeigt Aktivität */
    .working-spinner {
        width: 24px;
        height: 24px;
        border: 3px solid transparent;
        border-top: 3px solid #7c3aed;
        border-right: 3px solid #2dd4bf;
        border-radius: 50%;
        animation: spinGradient 1s linear infinite;
        margin: 12px auto 0;
    }

    @keyframes spinGradient {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }

    /* Erfolgs-Animation */
    @keyframes successPulse {
        0%, 100% { transform: scale(1); }
        50% { transform: scale(1.05); }
    }

    .processing-animation.success {
        animation: successPulse 0.5s ease;
    }

    /* Logo oben links */
    .top-logo {
        position: fixed;
        top: 1rem;
        left: 1rem;
        z-index: 1000;
        height: 40px;
    }

    .top-logo img {
        height: 40px;
        width: auto;
    }
</style>
"""

# ============================================================================
# PDF-Klasse (aus create_pdf.py)
# ============================================================================

class ProtocolPDF(FPDF):
    """Professionelle PDF-Klasse für Meeting-Protokolle."""

    BLACK = (0, 0, 0)
    GRAY = (100, 100, 100)

    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=25)
        self.set_margins(left=20, top=20, right=20)
        self.is_first_page = True
        self.doc_title = ""

    def header(self):
        if not self.is_first_page and self.page_no() > 1:
            self.set_font("Helvetica", "B", 10)
            self.set_text_color(*self.BLACK)
            self.cell(0, 10, self.doc_title[:60], align="L")
            self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "", 9)
        self.set_text_color(*self.GRAY)
        page_info = f"{self.page_no()}/{{nb}}"
        self.cell(15, 10, page_info, align="L")
        self.cell(0, 10, self.doc_title[:50], align="L")

    def add_main_title(self, text: str):
        self.doc_title = text
        self.set_x(self.l_margin)
        self.set_font("Helvetica", "", 11)
        self.set_text_color(*self.GRAY)
        self.ln(10)
        self.cell(0, 6, text, align="L")
        self.ln(8)

    def add_protocol_title(self, text: str):
        self.set_x(self.l_margin)
        self.set_font("Helvetica", "B", 14)
        self.set_text_color(*self.BLACK)
        self.cell(0, 8, text, align="L")
        self.ln(12)

    def add_meta_label(self, label: str, value: str):
        self.set_x(self.l_margin)
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(*self.BLACK)
        self.cell(25, 6, label, align="L")
        self.set_font("Helvetica", "", 10)
        self.cell(0, 6, value, align="L")
        self.ln(6)

    def add_section_header(self, text: str):
        self.ln(6)
        self.set_x(self.l_margin)
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(*self.BLACK)
        text_width = self.get_string_width(text)
        self.cell(text_width, 6, text, align="L")
        self.ln(1)
        self.set_draw_color(*self.BLACK)
        self.set_line_width(0.4)
        self.line(self.l_margin, self.get_y(), self.l_margin + text_width, self.get_y())
        self.ln(5)

    def add_participant_row(self, name: str, role: str = ""):
        self.set_x(self.l_margin)
        self.set_font("Helvetica", "", 10)
        self.set_text_color(*self.BLACK)
        if role:
            self.cell(40, 5, name, align="L")
            self.set_text_color(*self.GRAY)
            self.cell(0, 5, role, align="L")
            self.ln(5)
        else:
            self.cell(0, 5, name, align="L")
            self.ln(5)

    def add_traktandum(self, number: str, text: str):
        self.set_x(self.l_margin)
        self.set_font("Helvetica", "", 10)
        self.set_text_color(*self.BLACK)
        self.cell(8, 5, number, align="L")
        self.cell(0, 5, text, align="L")
        self.ln(5)

    def add_content_title(self, number: str, text: str):
        self.ln(6)
        self.set_x(self.l_margin)
        self.set_font("Helvetica", "B", 11)
        self.set_text_color(*self.BLACK)
        if number:
            self.cell(10, 7, number, align="L")
        self.cell(0, 7, text, align="L")
        self.ln(9)

    def add_body_text(self, text: str):
        """Fügt Fliesstext mit Markdown-Bold-Unterstützung hinzu."""
        self.set_x(self.l_margin)
        self.set_text_color(*self.BLACK)
        text = text.replace("\u2022", "-").replace(chr(149), "-")

        # Prüfen ob der gesamte Text fett sein soll (z.B. **Titel**)
        if text.startswith("**") and text.endswith("**") and text.count("**") == 2:
            # Ganzer Text ist fett
            clean_text = text[2:-2]
            self.set_font("Helvetica", "B", 10)
            self.multi_cell(0, 5.5, clean_text)
            self.ln(2)
            return

        # Inline-Bold-Formatierung: **text** wird fett
        # Regex um **text** zu finden und in Teile aufzuteilen
        parts = re.split(r'(\*\*[^*]+\*\*)', text)

        if len(parts) == 1:
            # Kein Bold-Text gefunden, normal ausgeben
            self.set_font("Helvetica", "", 10)
            clean_text = text.replace("**", "")
            self.multi_cell(0, 5.5, clean_text)
            self.ln(2)
            return

        # Mixed content: Bold und Normal
        line_height = 5.5
        x_start = self.l_margin
        max_width = self.w - self.l_margin - self.r_margin

        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                # Bold text
                self.set_font("Helvetica", "B", 10)
                clean_part = part[2:-2]
            else:
                # Normal text
                self.set_font("Helvetica", "", 10)
                clean_part = part

            # Text ausgeben (FPDF write für inline-Text)
            self.write(line_height, clean_part)

        self.ln(line_height + 2)

    def add_task_row(self, task: str, responsible: str):
        self.set_x(self.l_margin)
        self.set_font("Helvetica", "", 10)
        self.set_text_color(*self.BLACK)
        page_width = self.w - self.l_margin - self.r_margin
        self.cell(5, 6, "-", align="L")
        task_width = page_width - 55
        self.cell(task_width, 6, task[:80], align="L")
        if responsible:
            self.set_font("Helvetica", "", 9)
            self.set_text_color(*self.GRAY)
            self.cell(50, 6, responsible, align="R")
        self.ln(6)

    def add_signature(self, name: str, date: str):
        self.ln(10)
        self.set_font("Helvetica", "", 10)
        self.set_text_color(*self.BLACK)
        self.cell(0, 5, f"{name}, {date}", align="L")


# ============================================================================
# Kernfunktionen
# ============================================================================

def get_ffprobe_path():
    """Findet ffprobe (liegt im gleichen Ordner wie ffmpeg)."""
    if FFMPEG_PATH:
        ffprobe = FFMPEG_PATH.replace("ffmpeg", "ffprobe")
        if os.path.isfile(ffprobe):
            return ffprobe
    return shutil.which("ffprobe")


def get_audio_duration(file_path: str) -> float:
    """Ermittelt die Dauer einer Audio-Datei in Sekunden mit ffprobe."""
    ffprobe = get_ffprobe_path()
    print(f"[DURATION] ffprobe path: {ffprobe}")
    if not ffprobe:
        print("[DURATION] FEHLER: ffprobe nicht gefunden!")
        return 0
    try:
        result = subprocess.run(
            [ffprobe, "-v", "error", "-show_entries", "format=duration",
             "-of", "default=noprint_wrappers=1:nokey=1", file_path],
            capture_output=True, text=True, timeout=30
        )
        print(f"[DURATION] ffprobe output: '{result.stdout.strip()}', stderr: '{result.stderr.strip()}'")
        duration = float(result.stdout.strip())
        print(f"[DURATION] Dauer: {duration} Sekunden")
        return duration
    except Exception as e:
        print(f"[DURATION] FEHLER: {e}")
        return 0


def split_audio_file(file_path: str, chunk_duration_ms: int = CHUNK_DURATION_MS) -> list:
    """Teilt eine Audio-Datei in kleinere Chunks auf mit ffmpeg."""
    print(f"[SPLIT] Start - ffmpeg available: {FFMPEG_AVAILABLE}, path: {FFMPEG_PATH}")

    if not FFMPEG_AVAILABLE or not FFMPEG_PATH:
        print("[SPLIT] FEHLER: ffmpeg nicht verfügbar!")
        return [file_path]

    try:
        # Audio-Dauer ermitteln
        duration_sec = get_audio_duration(file_path)
        chunk_duration_sec = chunk_duration_ms / 1000
        print(f"[SPLIT] Audio-Dauer: {duration_sec} Sekunden ({duration_sec/60:.1f} Minuten)")
        print(f"[SPLIT] Chunk-Dauer: {chunk_duration_sec} Sekunden")

        # Wenn Audio kurz genug ist oder Dauer unbekannt, nicht splitten
        if duration_sec <= 0:
            print("[SPLIT] FEHLER: Konnte Audio-Dauer nicht ermitteln!")
            return [file_path]

        if duration_sec <= chunk_duration_sec:
            print("[SPLIT] Audio kurz genug, kein Splitting nötig")
            return [file_path]

        # In Chunks aufteilen mit ffmpeg
        chunks = []
        num_chunks = math.ceil(duration_sec / chunk_duration_sec)
        base_path = os.path.splitext(file_path)[0]

        for i in range(num_chunks):
            start_sec = i * chunk_duration_sec
            chunk_path = f"{base_path}_chunk{i}.mp3"

            # ffmpeg Befehl: Segment extrahieren und als MP3 speichern
            cmd = [
                FFMPEG_PATH, "-y", "-i", file_path,
                "-ss", str(start_sec),
                "-t", str(chunk_duration_sec),
                "-acodec", "libmp3lame", "-b:a", "128k",
                "-loglevel", "error",
                chunk_path
            ]

            result = subprocess.run(cmd, capture_output=True, timeout=120)
            if result.returncode == 0 and os.path.exists(chunk_path):
                chunks.append(chunk_path)
            else:
                # Bei Fehler: Aufräumen und Original zurückgeben
                for c in chunks:
                    if os.path.exists(c):
                        os.remove(c)
                return [file_path]

        return chunks
    except Exception as e:
        # Bei Fehler: Original-Datei zurückgeben
        return [file_path]


def get_audio_duration_minutes(audio_file) -> float:
    """
    Ermittelt die Dauer einer Audiodatei in Minuten.
    Verwendet pydub falls verfügbar, sonst Schätzung basierend auf Dateigrösse.
    """
    try:
        from pydub import AudioSegment
        import tempfile
        import os

        file_ext = os.path.splitext(audio_file.name)[1].lower() or ".mp3"

        # Temporäre Datei erstellen
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
            tmp.write(audio_file.getvalue())
            tmp_path = tmp.name

        try:
            # Audio laden und Dauer ermitteln
            audio = AudioSegment.from_file(tmp_path)
            duration_minutes = len(audio) / 1000 / 60  # ms -> Minuten
            return duration_minutes
        finally:
            # Temporäre Datei löschen
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    except Exception as e:
        # Fallback: Schätzung basierend auf Dateigrösse
        # Durchschnittlich ~1 MB pro Minute bei MP3 128kbps
        file_size_mb = len(audio_file.getvalue()) / (1024 * 1024)
        estimated_minutes = file_size_mb * 1.0  # Konservative Schätzung
        return estimated_minutes


def transcribe_audio(audio_file, client: OpenAI, progress_callback=None, status_callback=None, language: str = None) -> str:
    """Transkribiert eine Audio-Datei mit OpenAI Whisper. Unterstützt große Dateien durch automatisches Splitting."""
    file_ext = os.path.splitext(audio_file.name)[1].lower() or ".mp3"

    # Sprache aus Session State holen falls nicht übergeben
    if language is None:
        lang_code = st.session_state.get("language", "en")
        language = TRANSLATIONS.get(lang_code, TRANSLATIONS["en"]).get("whisper_language", "en")

    # Temporäre Datei erstellen
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
        tmp.write(audio_file.read())
        tmp_path = tmp.name

    chunk_paths = []

    try:
        # Dateigröße prüfen
        file_size = os.path.getsize(tmp_path)

        if status_callback:
            status_callback(f"📁 {get_text('file_size')}: {file_size // (1024*1024)} MB")

        if file_size <= WHISPER_CHUNK_SIZE:
            # Kleine Datei - direkt transkribieren
            if status_callback:
                status_callback(f"📝 {get_text('small_file')}")
            with open(tmp_path, "rb") as f:
                transcript = client.audio.transcriptions.create(
                    model="whisper-1",
                    file=f,
                    language=language,
                    response_format="text",
                    temperature=0,  # Reduziert Halluzinationen
                    prompt="This is a meeting recording. Transcribe only what is actually spoken. Do not add or invent any content."
                )
            return transcript
        else:
            # Große Datei - in Chunks aufteilen
            if status_callback:
                status_callback(f"✂️ {get_text('large_file_splitting')} (ffmpeg: {FFMPEG_PATH})...")

            chunk_paths = split_audio_file(tmp_path)

            if status_callback:
                status_callback(f"📦 {len(chunk_paths)} {get_text('parts_created')}")

            # Prüfen ob wirklich gesplittet wurde
            if len(chunk_paths) == 1 and chunk_paths[0] == tmp_path:
                if status_callback:
                    status_callback("⚠️ WARNING: File was NOT split!")

            transcripts = []
            for i, chunk_path in enumerate(chunk_paths):
                if progress_callback:
                    progress_callback(i + 1, len(chunk_paths))
                if status_callback:
                    chunk_size = os.path.getsize(chunk_path) // (1024*1024)
                    status_callback(f"🎙️ {get_text('transcribing_part')} {i+1}/{len(chunk_paths)} ({chunk_size} MB)...")

                with open(chunk_path, "rb") as f:
                    chunk_transcript = client.audio.transcriptions.create(
                        model="whisper-1",
                        file=f,
                        language=language,
                        response_format="text",
                        temperature=0,  # Reduziert Halluzinationen
                        prompt="This is a meeting recording. Transcribe only what is actually spoken. Do not add or invent any content."
                    )
                transcripts.append(chunk_transcript)

                if status_callback:
                    words_in_chunk = len(chunk_transcript.split())
                    status_callback(f"✓ Part {i+1}: {words_in_chunk} {get_text('words_transcribed')}")

            # Alle Transkripte zusammenführen
            full_transcript = " ".join(transcripts)
            if status_callback:
                total_words = len(full_transcript.split())
                status_callback(f"✅ {get_text('total_words')}: {total_words} {get_text('words')} from {len(chunk_paths)} parts")

            return full_transcript

    except Exception as e:
        error_msg = str(e)
        if "400" in error_msg:
            raise Exception(f"Dateiformat-Fehler: Die Datei '{audio_file.name}' konnte nicht verarbeitet werden. Versuche MP3 oder WAV.")
        raise
    finally:
        # Aufräumen
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        for chunk_path in chunk_paths:
            if chunk_path != tmp_path and os.path.exists(chunk_path):
                os.unlink(chunk_path)


def generate_protocol_text(transcript: str, mistral_api_key: str) -> str:
    """Generiert ein strukturiertes Protokoll aus dem Transkript mit Mistral AI (EU-basiert)."""

    # Debug: Transkript-Länge
    transcript_words = len(transcript.split())
    transcript_chars = len(transcript)
    print(f"[PROTOKOLL] Transkript-Eingabe: {transcript_words} Wörter, {transcript_chars} Zeichen")

    # Sprache aus Session State holen
    lang_code = st.session_state.get("language", "en")
    gpt_language = TRANSLATIONS.get(lang_code, TRANSLATIONS["en"]).get("gpt_language", "English")

    # Sprachspezifische Prompts (OR-konform für CH)
    language_instructions = {
        "en": {
            "style": "professional English",
            "date_format": "Date:",
            "location": "Location:",
            "participants": "Participants",
            "agenda": "Agenda",
            "tasks": "Tasks",
            "task_overview": "Tasks (Overview)",
            "responsible": "Responsible",
            "deadline": "Deadline",
            "protocol": "Meeting Minutes",
            "decisions": "Decisions",
            "decisions_overview": "Decisions (Overview)",
            "speaker_contributions": "Speaker Contributions",
            "signature_author": "Minutes taken by",
            "signature_chair": "Chairperson"
        },
        "de": {
            "style": "Schweizer Hochdeutsch",
            "date_format": "Datum:",
            "location": "Ort:",
            "participants": "Teilnehmende",
            "agenda": "Traktanden",
            "tasks": "Pendenzen",
            "task_overview": "Pendenzen (Gesamtübersicht)",
            "responsible": "Zuständig",
            "deadline": "Termin",
            "protocol": "Protokoll der Sitzung",
            "decisions": "Beschlüsse",
            "decisions_overview": "Beschlüsse (Gesamtübersicht)",
            "speaker_contributions": "Beiträge der Teilnehmenden",
            "signature_author": "Protokollführung",
            "signature_chair": "Vorsitz"
        },
        "fr": {
            "style": "français professionnel",
            "date_format": "Date:",
            "location": "Lieu:",
            "participants": "Participants",
            "agenda": "Ordre du jour",
            "tasks": "Actions",
            "task_overview": "Actions (Résumé)",
            "responsible": "Responsable",
            "deadline": "Échéance",
            "protocol": "Procès-verbal de la réunion",
            "decisions": "Décisions",
            "decisions_overview": "Décisions (Résumé)",
            "speaker_contributions": "Contributions des participants",
            "signature_author": "Procès-verbal rédigé par",
            "signature_chair": "Présidence"
        },
        "it": {
            "style": "italiano professionale",
            "date_format": "Data:",
            "location": "Luogo:",
            "participants": "Partecipanti",
            "agenda": "Ordine del giorno",
            "tasks": "Azioni",
            "task_overview": "Azioni (Riepilogo)",
            "responsible": "Responsabile",
            "deadline": "Scadenza",
            "protocol": "Verbale della riunione",
            "decisions": "Delibere",
            "decisions_overview": "Delibere (Riepilogo)",
            "speaker_contributions": "Contributi dei partecipanti",
            "signature_author": "Verbale redatto da",
            "signature_chair": "Presidenza"
        }
    }

    lang_config = language_instructions.get(lang_code, language_instructions["en"])

    system_prompt = f"""You are a professional meeting minute taker. Create a protocol in {gpt_language}.

⚠️ CRITICAL ACCURACY REQUIREMENT ⚠️
- ONLY include information that is EXPLICITLY stated in the transcript
- DO NOT invent, assume, fabricate, or add ANY content not in the transcript
- If the transcript is short, the protocol MUST be short
- If information is missing (date, location, names), write "Not specified" - DO NOT guess
- NEVER add plausible-sounding content that wasn't said
- If unsure about something, omit it rather than guess

===FORMAT===

# [Project/Topic]
## {lang_config['protocol']}

**{lang_config['date_format']}** [from transcript or "Not specified"]
**{lang_config['location']}** [from transcript or omit]

**{lang_config['participants']}**
| Name | Function/Organization |

**{lang_config['agenda']}**
1. [Topic 1]
2. [Topic 2]
...

---

## 1 [First Agenda Item]

[COMPREHENSIVE flowing text - at least 2-4 paragraphs:
- Background and context
- What was discussed (all points!)
- What opinions/positions were expressed
- What was decided]

**{lang_config['tasks']}:**
| Task | {lang_config['responsible']} | {lang_config['deadline']} |

## 2 [Second Agenda Item]

[Again comprehensive - 2-4 paragraphs]

[... more agenda items ...]

---

## {lang_config['speaker_contributions']}
[For each identified speaker, write 2-3 sentences summarizing their main contributions and positions. Only include speakers who are clearly identifiable in the transcript.]

**[Speaker Name 1]:** [What they contributed, proposed, or emphasized]

**[Speaker Name 2]:** [What they contributed, proposed, or emphasized]

[Continue for all identified speakers...]

---

## {lang_config['decisions_overview']}
| Nr. | {lang_config['decisions']} |
[List all decisions made during the meeting with voting results if mentioned]

## {lang_config['task_overview']}
| Nr. | Task | {lang_config['responsible']} | {lang_config['deadline']} |

---

**{lang_config['signature_author']}:** _________________________ [Name], [Date]

**{lang_config['signature_chair']}:** _________________________ [Name], [Date]

===RULES===
- ONLY include what is EXPLICITLY in the transcript
- DO NOT add, invent, or assume any information
- If transcript is short, protocol is short - that's OK!
- Write in {lang_config['style']}
- Flowing text, no bullet points in main text
- If information is missing, write "Not specified" or omit
- ACCURACY over length - never fabricate content"""

    user_prompt = f"""Here is the meeting transcript ({transcript_words} words).

⚠️ CRITICAL: Create an ACCURATE protocol in {gpt_language}.
- ONLY include information from this transcript
- DO NOT add or invent ANY content
- If the transcript is short, the protocol should be proportionally short
- ACCURACY is more important than length

TRANSCRIPT:
{transcript}"""

    # Mistral API via HTTP (ohne SDK)
    headers = {
        "Authorization": f"Bearer {mistral_api_key}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "mistral-large-latest",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "temperature": 0.1,  # Niedrig für maximale Genauigkeit, minimale Halluzinationen
        "max_tokens": 12000
    }

    response = requests.post(
        "https://api.mistral.ai/v1/chat/completions",
        headers=headers,
        json=payload,
        timeout=120
    )

    if response.status_code != 200:
        raise Exception(f"Mistral API Error: {response.status_code} - {response.text}")

    result = response.json()["choices"][0]["message"]["content"]
    result_words = len(result.split())
    print(f"[PROTOKOLL] Generiertes Protokoll: {result_words} Wörter")

    return result


def sanitize_text_for_pdf(text: str) -> str:
    """Entfernt oder ersetzt Sonderzeichen, die von Helvetica nicht unterstützt werden."""
    # Häufige problematische Unicode-Zeichen ersetzen
    replacements = {
        '"': '"',  # Typografische Anführungszeichen
        '"': '"',
        ''': "'",
        ''': "'",
        '–': '-',  # En-dash
        '—': '-',  # Em-dash
        '…': '...',
        '•': '-',
        '→': '->',
        '←': '<-',
        '≤': '<=',
        '≥': '>=',
        '≠': '!=',
        '×': 'x',
        '÷': '/',
        '€': 'EUR',
        '£': 'GBP',
        '¥': 'JPY',
        '©': '(c)',
        '®': '(R)',
        '™': '(TM)',
        '°': ' Grad',
        '±': '+/-',
        '½': '1/2',
        '¼': '1/4',
        '¾': '3/4',
        '\u200b': '',  # Zero-width space
        '\u00a0': ' ',  # Non-breaking space
        '\ufeff': '',  # BOM
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    # Alle verbleibenden nicht-ASCII Zeichen durch ASCII-Äquivalente ersetzen
    # Behalte deutsche Umlaute, da FPDF diese mit latin-1 unterstützt
    result = []
    for char in text:
        if ord(char) < 128:  # ASCII
            result.append(char)
        elif char in 'äöüÄÖÜß':  # Deutsche Umlaute
            result.append(char)
        elif char in 'àâáãåçèéêëìíîïñòóôõùúûýÿ':  # Andere latin-1 Zeichen
            result.append(char)
        elif char in 'ÀÂÁÃÅÇÈÉÊËÌÍÎÏÑÒÓÔÕÙÚÛÝ':
            result.append(char)
        else:
            result.append('')  # Unbekannte Zeichen entfernen

    return ''.join(result)

def parse_markdown_to_pdf(markdown_text: str) -> bytes:
    """Konvertiert Markdown-Protokoll zu PDF und gibt Bytes zurück."""
    # Text für PDF bereinigen
    markdown_text = sanitize_text_for_pdf(markdown_text)
    pdf = ProtocolPDF()
    pdf.alias_nb_pages()
    pdf.add_page()

    lines = markdown_text.split("\n")
    i = 0
    in_participants = False
    in_traktanden = False
    in_tasks = False

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            in_participants = False
            in_traktanden = False
            in_tasks = False
            i += 1
            continue

        if line in ["---", "===DECKBLATT===", "===INHALT===", "===ABSCHLUSS==="]:
            if line == "===INHALT===":
                pdf.is_first_page = False
                pdf.ln(8)
                pdf.set_draw_color(*pdf.GRAY)
                pdf.set_line_width(0.5)
                pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
                pdf.ln(4)
            i += 1
            continue

        if line.startswith("|"):
            if "---" in line:
                i += 1
                continue

            if "Aufgabe" in line or "Zuständig" in line or "Termin" in line:
                in_tasks = True
                i += 1
                continue

            if "Name" in line or "Funktion" in line:
                i += 1
                continue

            parts = [p.strip() for p in line.split("|") if p.strip()]
            if len(parts) >= 1:
                if in_tasks:
                    responsible = parts[1] if len(parts) > 1 else ""
                    pdf.add_task_row(parts[0], responsible)
                else:
                    role = parts[1] if len(parts) > 1 else ""
                    pdf.add_participant_row(parts[0], role)
            i += 1
            continue

        if line.startswith("# "):
            title = line[2:].strip()
            pdf.add_main_title(title)
            i += 1
            continue

        if line.startswith("## "):
            subtitle = line[3:].strip()
            match = re.match(r"^(\d+)\s+(.+)$", subtitle)
            if match:
                pdf.add_content_title(match.group(1), match.group(2))
            elif "Protokoll" in subtitle:
                pdf.add_protocol_title(subtitle)
            else:
                pdf.add_content_title("", subtitle)
            i += 1
            continue

        if line.startswith("**") and ":**" in line:
            match = re.match(r"\*\*(.+?):\*\*\s*(.*)", line)
            if match:
                label = match.group(1) + ":"
                value = match.group(2)

                if label == "Teilnehmende:":
                    pdf.add_section_header("Teilnehmende")
                    in_participants = True
                elif label == "Entschuldigte:":
                    pdf.add_section_header("Entschuldigte")
                    in_participants = True
                elif label == "Traktanden:":
                    pdf.add_section_header("Traktanden")
                    in_traktanden = True
                else:
                    pdf.add_meta_label(label, value)
            i += 1
            continue

        if in_traktanden and re.match(r"^\d+\.", line):
            match = re.match(r"^(\d+)\.\s*(.+)", line)
            if match:
                pdf.add_traktandum(match.group(1), match.group(2))
            i += 1
            continue

        if re.match(r"^\d+\.\s", line) and not in_traktanden:
            match = re.match(r"^(\d+)\.\s*(.+)", line)
            if match:
                pdf.add_traktandum(match.group(1), match.group(2))
            i += 1
            continue

        if re.match(r"^[A-Z][a-z]+\s+[A-Z][a-z]+,\s+\d", line):
            parts = line.split(",", 1)
            if len(parts) == 2:
                pdf.add_signature(parts[0].strip(), parts[1].strip())
            i += 1
            continue

        if "[Protokollführer" in line or "[Datum" in line:
            i += 1
            continue

        # Bold-Formatierung (**text**) wird in add_body_text verarbeitet
        clean_line = line.replace("\u2022", "-").strip()
        if len(clean_line) > 0:
            pdf.add_body_text(clean_line)
        i += 1

    return bytes(pdf.output())


def parse_markdown_to_docx(markdown_text: str) -> bytes:
    """Konvertiert Markdown-Protokoll zu Word-Dokument und gibt Bytes zurück."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = markdown_text.split("\n")
    i = 0
    in_table = False
    table_data = []

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            if in_table and table_data:
                if len(table_data) > 0:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_text in enumerate(row_data):
                            table.rows[row_idx].cells[col_idx].text = cell_text
                    doc.add_paragraph()
                table_data = []
                in_table = False
            i += 1
            continue

        if line in ["---", "===DECKBLATT===", "===INHALT===", "===ABSCHLUSS==="]:
            if line == "===INHALT===":
                doc.add_paragraph("_" * 60)
            i += 1
            continue

        if line.startswith("|"):
            if "---" in line:
                i += 1
                continue

            in_table = True
            parts = [p.strip() for p in line.split("|") if p.strip()]
            if parts:
                table_data.append(parts)
            i += 1
            continue

        if in_table and table_data:
            if len(table_data) > 0:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < len(table.rows[row_idx].cells):
                            table.rows[row_idx].cells[col_idx].text = cell_text
                doc.add_paragraph()
            table_data = []
            in_table = False

        if line.startswith("# "):
            title = line[2:].strip()
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(100, 100, 100)
            i += 1
            continue

        if line.startswith("## "):
            subtitle = line[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(subtitle)
            run.bold = True
            run.font.size = Pt(13)
            i += 1
            continue

        if line.startswith("**") and ":**" in line:
            match = re.match(r"\*\*(.+?):\*\*\s*(.*)", line)
            if match:
                label = match.group(1) + ":"
                value = match.group(2)
                p = doc.add_paragraph()
                run_label = p.add_run(label + " ")
                run_label.bold = True
                p.add_run(value)
            i += 1
            continue

        if re.match(r"^\d+\.\s", line):
            match = re.match(r"^(\d+)\.\s*(.+)", line)
            if match:
                p = doc.add_paragraph(f"{match.group(1)}. {match.group(2)}")
            i += 1
            continue

        if "[Protokollführer" in line or "[Datum" in line:
            i += 1
            continue

        if re.match(r"^[A-Z][a-z]+\s+[A-Z][a-z]+,\s+\d", line):
            doc.add_paragraph()
            p = doc.add_paragraph(line)
            p.paragraph_format.space_before = Pt(24)
            i += 1
            continue

        clean_line = line.replace("**", "").replace("\u2022", "-").strip()
        if clean_line:
            doc.add_paragraph(clean_line)
        i += 1

    if in_table and table_data and len(table_data) > 0:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(table.rows[row_idx].cells):
                    table.rows[row_idx].cells[col_idx].text = cell_text

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def send_email_with_protocol(pdf_bytes: bytes, docx_bytes: bytes, recipient: str, filename_base: str, send_pdf: bool = True, send_word: bool = True) -> tuple[bool, str]:
    """Versendet PDF und/oder Word-Dokument per E-Mail in der gewählten Sprache."""
    smtp_email = get_secret("SMTP_EMAIL")
    smtp_password = get_secret("SMTP_PASSWORD")
    smtp_server = get_secret("SMTP_SERVER", "smtp.gmail.com")
    smtp_port = int(get_secret("SMTP_PORT", "587"))

    if not smtp_email or not smtp_password:
        return False, "SMTP configuration missing in .env"

    msg = MIMEMultipart()
    msg["From"] = smtp_email
    msg["To"] = recipient
    msg["Subject"] = f"{get_text('email_subject')} {datetime.now().strftime('%d.%m.%Y')}"

    # HTML E-Mail im einheitlichen Design
    body_html = f"""
    <html>
    <body style="font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; color: #1d1d1f; line-height: 1.6; max-width: 600px; margin: 0 auto; padding: 20px;">
        <div style="text-align: center; margin-bottom: 30px;">
            <h1 style="font-size: 28px; font-weight: 600; margin: 0; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text;">
                MINU Minutes AI
            </h1>
        </div>

        <p style="font-size: 16px; color: #1d1d1f; text-align: center;">{get_text('email_body_greeting')}</p>
        <p style="font-size: 16px; color: #1d1d1f; text-align: center;">{get_text('email_body_text')}</p>
        <p style="font-size: 16px; color: #1d1d1f; text-align: center;">{get_text('email_body_closing')}<br>
        <strong>MINU Minutes AI</strong></p>

        <hr style="border: none; border-top: 1px solid #d2d2d7; margin: 30px 0;">

        <p style="color: #86868b; font-size: 14px; text-align: center;">{get_text('email_rating_intro')}</p>
        <table cellpadding="0" cellspacing="0" style="margin: 10px auto;">
            <tr>
                <td><a href="{APP_URL}/?feedback=1" style="text-decoration: none; font-size: 28px; padding: 5px 8px;">⭐</a></td>
                <td><a href="{APP_URL}/?feedback=2" style="text-decoration: none; font-size: 28px; padding: 5px 8px;">⭐</a></td>
                <td><a href="{APP_URL}/?feedback=3" style="text-decoration: none; font-size: 28px; padding: 5px 8px;">⭐</a></td>
                <td><a href="{APP_URL}/?feedback=4" style="text-decoration: none; font-size: 28px; padding: 5px 8px;">⭐</a></td>
                <td><a href="{APP_URL}/?feedback=5" style="text-decoration: none; font-size: 28px; padding: 5px 8px;">⭐</a></td>
            </tr>
            <tr style="font-size: 11px; color: #86868b; text-align: center;">
                <td>1</td>
                <td>2</td>
                <td>3</td>
                <td>4</td>
                <td>5</td>
            </tr>
        </table>
        <p style="color: #86868b; font-size: 12px; text-align: center;">{get_text('email_rating_thanks')}</p>

        <hr style="border: none; border-top: 1px solid #d2d2d7; margin: 30px 0;">

        <p style="font-size: 12px; color: #86868b; text-align: center;">
            <a href="{APP_URL}" style="color: #667eea;">{APP_DOMAIN}</a>
        </p>
        <p style="font-size: 11px; color: #86868b; text-align: center; margin-top: 15px;">
            🇨🇭 Swiss Developer · 🇪🇺 EU-DSGVO konform
        </p>
    </body>
    </html>
    """
    msg.attach(MIMEText(body_html, "html", "utf-8"))

    # PDF anhängen (wenn ausgewählt)
    if send_pdf and pdf_bytes:
        pdf_attachment = MIMEApplication(pdf_bytes, _subtype="pdf")
        pdf_attachment.add_header("Content-Disposition", "attachment", filename=f"{filename_base}.pdf")
        msg.attach(pdf_attachment)

    # Word anhängen (wenn ausgewählt)
    if send_word and docx_bytes:
        docx_attachment = MIMEApplication(docx_bytes, _subtype="vnd.openxmlformats-officedocument.wordprocessingml.document")
        docx_attachment.add_header("Content-Disposition", "attachment", filename=f"{filename_base}.docx")
        msg.attach(docx_attachment)

    try:
        # Port 465 = SSL, Port 587 = STARTTLS
        if smtp_port == 465:
            with smtplib.SMTP_SSL(smtp_server, smtp_port) as server:
                server.login(smtp_email, smtp_password)
                server.send_message(msg)
        else:
            with smtplib.SMTP(smtp_server, smtp_port) as server:
                server.starttls()
                server.login(smtp_email, smtp_password)
                server.send_message(msg)
        return True, f"E-Mail mit PDF und Word erfolgreich an {recipient} gesendet!"
    except smtplib.SMTPAuthenticationError:
        return False, "SMTP-Authentifizierung fehlgeschlagen. Prüfe .env-Datei."
    except Exception as e:
        return False, f"Fehler beim Versand: {str(e)}"


def send_verification_email(recipient: str, company: str, token: str, language: str = "de") -> tuple[bool, str]:
    """Sendet eine Verifizierungs-E-Mail via Resend API (oder SMTP Fallback)."""
    verify_url = f"{APP_URL}?verify={token}"

    # Temporär Sprache setzen für get_text
    original_lang = st.session_state.get("language", "de")
    st.session_state.language = language

    subject = get_text("verify_email_subject")
    heading = get_text("verify_email_heading")
    text = get_text("verify_email_text")
    button_text = get_text("verify_email_button")
    expiry_text = get_text("verify_email_expiry")
    ignore_text = get_text("verify_email_ignore")

    # Sprache zurücksetzen
    st.session_state.language = original_lang

    body_html = f"""
    <html>
    <body style="font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; color: #1d1d1f; line-height: 1.6; max-width: 600px; margin: 0 auto; padding: 20px; background-color: #f5f5f7;">
        <div style="background: white; border-radius: 16px; padding: 40px; box-shadow: 0 2px 8px rgba(0,0,0,0.06);">
            <div style="text-align: center; margin-bottom: 30px;">
                <h1 style="font-size: 24px; font-weight: 700; margin: 0; background: linear-gradient(135deg, #7c3aed 0%, #2dd4bf 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text;">
                    MINU-AI
                </h1>
            </div>

            <h2 style="font-size: 22px; font-weight: 600; color: #1d1d1f; margin-bottom: 16px; text-align: center;">
                {heading}
            </h2>

            <p style="font-size: 16px; color: #333; margin-bottom: 30px; text-align: center;">
                {text}
            </p>

            <div style="text-align: center; margin: 32px 0;">
                <a href="{verify_url}" style="display: inline-block; background: linear-gradient(135deg, #7c3aed 0%, #2dd4bf 100%); color: white; text-decoration: none; padding: 14px 40px; border-radius: 25px; font-weight: 600; font-size: 16px;">
                    {button_text}
                </a>
            </div>

            <p style="font-size: 13px; color: #86868b; text-align: center; margin-top: 24px;">
                {expiry_text}
            </p>

            <p style="font-size: 12px; color: #aaa; text-align: center; margin-top: 16px;">
                {ignore_text}
            </p>

            <hr style="border: none; border-top: 1px solid #e4e3df; margin: 24px 0;">

            <p style="font-size: 11px; color: #86868b; text-align: center;">
                <a href="{APP_URL}" style="color: #7c3aed;">{APP_DOMAIN}</a> · 🇨🇭 Swiss Developer · 🇪🇺 EU-DSGVO konform
            </p>
        </div>
    </body>
    </html>
    """

    # Versuch 1: Resend API
    resend_api_key = get_secret("RESEND_API_KEY", "")
    if RESEND_AVAILABLE and resend_api_key and not resend_api_key.startswith("re_DEIN"):
        try:
            resend_lib.api_key = resend_api_key
            # Absender: verifizierte Domain oder Resend-Testdomain
            smtp_email = get_secret("SMTP_EMAIL", "onboarding@resend.dev")
            params = {
                "from": f"MINU-AI <{smtp_email}>",
                "to": [recipient],
                "subject": subject,
                "html": body_html,
            }
            email_response = resend_lib.Emails.send(params)
            print(f"[VERIFY] Resend E-Mail an {recipient} gesendet: {email_response}")
            return True, f"Verification email sent via Resend to {recipient}"
        except Exception as e:
            print(f"[VERIFY] Resend Fehler: {e}, versuche SMTP Fallback...")

    # Versuch 2: SMTP Fallback
    smtp_email = get_secret("SMTP_EMAIL")
    smtp_password = get_secret("SMTP_PASSWORD")
    smtp_server = get_secret("SMTP_SERVER", "smtp.gmail.com")
    smtp_port = int(get_secret("SMTP_PORT", "587"))

    if not smtp_email or not smtp_password:
        return False, "Weder Resend noch SMTP konfiguriert"

    try:
        msg = MIMEMultipart()
        msg["From"] = smtp_email
        msg["To"] = recipient
        msg["Subject"] = subject
        msg.attach(MIMEText(body_html, "html", "utf-8"))

        if smtp_port == 465:
            with smtplib.SMTP_SSL(smtp_server, smtp_port) as server:
                server.login(smtp_email, smtp_password)
                server.send_message(msg)
        else:
            with smtplib.SMTP(smtp_server, smtp_port) as server:
                server.starttls()
                server.login(smtp_email, smtp_password)
                server.send_message(msg)
        print(f"[VERIFY] SMTP E-Mail erfolgreich an {recipient} gesendet")
        return True, f"Verification email sent via SMTP to {recipient}"
    except Exception as e:
        print(f"[VERIFY] SMTP Fehler: {e}")
        return False, f"Error sending verification email: {str(e)}"


def send_welcome_email(recipient: str, company: str, language: str = "de") -> tuple[bool, str]:
    """Sendet eine Willkommens-E-Mail an neue Benutzer."""
    smtp_email = get_secret("SMTP_EMAIL")
    smtp_password = get_secret("SMTP_PASSWORD")
    smtp_server = get_secret("SMTP_SERVER", "smtp.gmail.com")
    smtp_port = int(get_secret("SMTP_PORT", "587"))

    if not smtp_email or not smtp_password:
        return False, "SMTP configuration missing"

    # Temporär Sprache setzen für get_text
    original_lang = st.session_state.get("language", "de")
    st.session_state.language = language

    msg = MIMEMultipart()
    msg["From"] = smtp_email
    msg["To"] = recipient
    msg["Subject"] = get_text("welcome_email_subject")

    body_html = f"""
    <html>
    <body style="font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; color: #1d1d1f; line-height: 1.6; max-width: 600px; margin: 0 auto; padding: 20px;">
        <div style="text-align: center; margin-bottom: 30px;">
            <h1 style="font-size: 28px; font-weight: 600; margin: 0; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text;">
                MINU Minutes AI
            </h1>
        </div>

        <h2 style="font-size: 22px; font-weight: 600; color: #1d1d1f; margin-bottom: 20px; text-align: center;">
            {get_text('welcome_email_greeting')}
        </h2>

        <p style="font-size: 16px; color: #1d1d1f; margin-bottom: 25px; text-align: center;">
            {get_text('welcome_email_intro')}
        </p>

        <div style="background: linear-gradient(135deg, #f5f5f7 0%, #e8e8ed 100%); border-radius: 12px; padding: 25px; margin: 25px 0;">
            <h3 style="font-size: 16px; font-weight: 600; color: #1d1d1f; margin: 0 0 15px 0;">
                {get_text('welcome_email_benefits_title')}
            </h3>
            <ul style="margin: 0; padding-left: 0; list-style: none; color: #1d1d1f;">
                <li style="margin-bottom: 10px;">{get_text('welcome_email_benefit1')}</li>
                <li style="margin-bottom: 10px;">{get_text('welcome_email_benefit2')}</li>
                <li style="margin-bottom: 10px;">{get_text('welcome_email_benefit3')}</li>
                <li style="margin-bottom: 10px;">{get_text('welcome_email_benefit4')}</li>
                <li style="margin-bottom: 10px;">{get_text('welcome_email_benefit5')}</li>
            </ul>
        </div>

        <p style="font-size: 16px; color: #1d1d1f; margin: 25px 0; text-align: center;">
            {get_text('welcome_email_cta')}
        </p>

        <div style="text-align: center; margin: 30px 0;">
            <a href="{APP_URL}" style="display: inline-block; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; text-decoration: none; padding: 14px 32px; border-radius: 25px; font-weight: 600; font-size: 16px;">
                minu-ai.ch
            </a>
        </div>

        <p style="font-size: 14px; color: #86868b; margin-top: 30px; text-align: center;">
            {get_text('welcome_email_help')}
        </p>

        <hr style="border: none; border-top: 1px solid #d2d2d7; margin: 30px 0;">

        <p style="font-size: 12px; color: #86868b; text-align: center;">
            <a href="{APP_URL}" style="color: #667eea;">{APP_DOMAIN}</a>
        </p>
        <p style="font-size: 11px; color: #86868b; text-align: center; margin-top: 15px;">
            🇨🇭 Swiss Developer · 🇪🇺 EU-DSGVO konform
        </p>
    </body>
    </html>
    """

    msg.attach(MIMEText(body_html, "html", "utf-8"))

    # Sprache zurücksetzen
    st.session_state.language = original_lang

    try:
        if smtp_port == 465:
            with smtplib.SMTP_SSL(smtp_server, smtp_port) as server:
                server.login(smtp_email, smtp_password)
                server.send_message(msg)
        else:
            with smtplib.SMTP(smtp_server, smtp_port) as server:
                server.starttls()
                server.login(smtp_email, smtp_password)
                server.send_message(msg)
        print(f"[WELCOME] E-Mail erfolgreich an {recipient} gesendet")
        return True, f"Welcome email sent to {recipient}"
    except Exception as e:
        print(f"[WELCOME] Fehler beim Versand: {e}")
        return False, f"Error sending welcome email: {str(e)}"


# ============================================================================
# UI Komponenten
# ============================================================================

def render_progress_tracker(current_step: int):
    """Rendert einen minimalistischen Fortschritts-Tracker im Apple-Stil."""
    steps = ["Upload", "Transkription", "Protokoll", "Dokumente", "Fertig"]

    # Einfache Progress Bar
    progress_value = (current_step - 1) / (len(steps) - 1) if current_step > 1 else 0
    st.progress(progress_value)

    # Aktueller Schritt als Text
    st.markdown(
        f"<p style='text-align:center; color:#86868b; font-size:14px; margin-top:8px;'>"
        f"Schritt {current_step} von {len(steps)}: <strong style='color:#1d1d1f;'>{steps[current_step-1]}</strong></p>",
        unsafe_allow_html=True
    )


def render_sidebar():
    """Rendert eine minimalistische Sidebar im Apple-Stil."""
    with st.sidebar:
        st.markdown("")  # Spacing

        # Logo - Clean
        if LOGO_AVAILABLE:
            st.image(str(LOGO_PATH), width=120)
            st.markdown("")
        st.markdown(f"### {get_text('sidebar_title')}")
        st.caption(get_text("sidebar_subtitle"))

        st.markdown("---")

        # Status - Minimal
        st.caption(get_text("status"))

        if st.session_state.get("transcript"):
            st.markdown(get_text("transcript_ready"))
        if st.session_state.get("protocol"):
            st.markdown(get_text("protocol_ready"))
        if st.session_state.get("pdf_bytes"):
            st.markdown(get_text("documents_ready"))

        if not any([st.session_state.get("transcript"), st.session_state.get("protocol"), st.session_state.get("pdf_bytes")]):
            st.markdown(get_text("ready_to_start"))

        # Beta-Limits anzeigen
        if st.session_state.get("authenticated"):
            user_email = st.session_state.get("user_email", "unknown")
            can_create, _, remaining_today, total_remaining = check_protocol_limits(user_email)
            st.markdown("---")
            st.caption("BETA LIMITS")
            st.markdown(f"📅 {get_text('remaining_today')}: **{remaining_today}**/{DAILY_USER_LIMIT}")
            st.markdown(f"📊 {get_text('beta_total_remaining')}: **{total_remaining}**/{BETA_TOTAL_LIMIT}")

        st.markdown("---")

        # Info - Minimal
        st.caption(get_text("technology"))
        st.markdown("Whisper · Mistral AI")

        st.markdown("")
        st.caption(get_text("installation"))
        st.markdown(f"{get_text('install_as_app')}", help="iOS: Share → Add to Home Screen\nChrome: Menu → Install App")

        # Admin-Bereich: Aktivitäts-Log (nur für Admins)
        if st.session_state.get("is_admin"):
            st.markdown("---")
            st.caption(get_text("admin"))

            with st.expander(get_text("activity_log")):
                logs = get_activity_logs()
                if logs:
                    # Neueste zuerst
                    for log in reversed(logs[-20:]):
                        st.text(f"{log['timestamp']}")
                        st.caption(f"{log['action']}: {log['details']}")
                        st.markdown("")
                else:
                    st.caption(get_text("no_activities"))

        # Abmelden für alle Benutzer
        if st.session_state.get("authenticated"):
            st.markdown("---")
            if st.button(get_text("logout"), use_container_width=True):
                log_activity("Logout", "Admin" if st.session_state.get("is_admin") else "User")
                st.session_state.authenticated = False
                st.session_state.is_admin = False
                st.rerun()


def get_current_step() -> int:
    """Ermittelt den aktuellen Schritt basierend auf dem Session State."""
    if st.session_state.get("pdf_bytes"):
        return 5
    elif st.session_state.get("protocol"):
        return 4
    elif st.session_state.get("transcript"):
        return 3
    elif st.session_state.get("uploaded_file_name"):
        return 2
    return 1


# ============================================================================
# Aktivitäts-Logging
# ============================================================================

ACTIVITY_LOG_FILE = PROJECT_ROOT / "activity_log.json"

def log_activity(action: str, details: str = ""):
    """Speichert eine Aktivität im Log."""
    try:
        # Bestehende Logs laden
        if ACTIVITY_LOG_FILE.exists():
            with open(ACTIVITY_LOG_FILE, "r", encoding="utf-8") as f:
                logs = json.load(f)
        else:
            logs = []

        # Neue Aktivität hinzufügen
        logs.append({
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "action": action,
            "details": details
        })

        # Nur letzte 100 Einträge behalten
        logs = logs[-100:]

        # Speichern
        with open(ACTIVITY_LOG_FILE, "w", encoding="utf-8") as f:
            json.dump(logs, f, ensure_ascii=False, indent=2)
    except:
        pass  # Logging-Fehler ignorieren


def get_activity_logs() -> list:
    """Lädt alle Aktivitäts-Logs."""
    try:
        if ACTIVITY_LOG_FILE.exists():
            with open(ACTIVITY_LOG_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except:
        pass
    return []


# ============================================================================
# Streamlit App
# ============================================================================

def check_password():
    """Prüft Zugang - Landing Page mit Pricing und Registrierung."""

    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if "is_admin" not in st.session_state:
        st.session_state.is_admin = False
    if "language" not in st.session_state:
        # GeoIP-basierte Spracherkennung beim ersten Besuch
        if "language_auto_detected" not in st.session_state:
            st.session_state.language = detect_language_from_ip()
            st.session_state.language_auto_detected = True
        else:
            st.session_state.language = "en"
    if "show_privacy" not in st.session_state:
        st.session_state.show_privacy = False
    if "show_register_form" not in st.session_state:
        st.session_state.show_register_form = False
    if "current_page" not in st.session_state:
        st.session_state.current_page = "landing"

    # Sprachwechsel für ALLE Benutzer (auch authentifizierte) HIER behandeln
    query_params = st.query_params
    if "lang" in query_params:
        new_lang = query_params["lang"]
        if new_lang in ["en", "de", "fr", "it"] and new_lang != st.session_state.get("language", "en"):
            st.session_state.language = new_lang
            st.query_params.clear()
            st.rerun()

    # ============ E-MAIL-VERIFIZIERUNG ============
    # Query-Parameter ?verify=<token> abfangen
    verify_token = st.query_params.get("verify")
    if verify_token:
        st.query_params.clear()
        result = verify_email_token(verify_token)

        # Einfacher Header für Verifizierungsseiten
        st.markdown('''
        <div style="position:fixed;top:0;left:0;width:100%;background:#fbfbfd;border-bottom:1px solid #d2d2d7;padding:8px 20px;z-index:1000;box-sizing:border-box;">
            <a href="/" style="display:flex;align-items:center;gap:9px;text-decoration:none;">
                <span style="font-size:18px;font-weight:700;background:linear-gradient(90deg,#7c3aed,#2dd4bf);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;">MINU-AI</span>
            </a>
        </div>
        <div style="height: 50px;"></div>
        ''', unsafe_allow_html=True)

        if result is None:
            # Token nicht gefunden
            st.markdown(f"""
            <div style="max-width: 500px; margin: 60px auto; text-align: center; padding: 40px; background: white; border-radius: 16px; box-shadow: 0 4px 20px rgba(0,0,0,0.08);">
                <div style="font-size: 48px; margin-bottom: 16px;">❌</div>
                <h2 style="color: #1d1d1f; margin-bottom: 12px;">{get_text('verify_invalid_title')}</h2>
                <p style="color: #86868b; font-size: 16px;">{get_text('verify_invalid_text')}</p>
                <a href="/" style="display: inline-block; margin-top: 24px; background: linear-gradient(135deg, #7c3aed 0%, #2dd4bf 100%); color: white; text-decoration: none; padding: 12px 32px; border-radius: 25px; font-weight: 600;">← {get_text('register_link')}</a>
            </div>
            """, unsafe_allow_html=True)
            st.stop()

        elif result.get("status") == "expired":
            # Token abgelaufen
            st.markdown(f"""
            <div style="max-width: 500px; margin: 60px auto; text-align: center; padding: 40px; background: white; border-radius: 16px; box-shadow: 0 4px 20px rgba(0,0,0,0.08);">
                <div style="font-size: 48px; margin-bottom: 16px;">⏰</div>
                <h2 style="color: #1d1d1f; margin-bottom: 12px;">{get_text('verify_expired_title')}</h2>
                <p style="color: #86868b; font-size: 16px;">{get_text('verify_expired_text')}</p>
                <a href="/" style="display: inline-block; margin-top: 24px; background: linear-gradient(135deg, #7c3aed 0%, #2dd4bf 100%); color: white; text-decoration: none; padding: 12px 32px; border-radius: 25px; font-weight: 600;">← {get_text('register_link')}</a>
            </div>
            """, unsafe_allow_html=True)
            st.stop()

        elif result.get("status") == "success":
            # Verifizierung erfolgreich!
            verified_email = result.get("email", "")
            verified_company = result.get("company", "")

            # Trial-Subscription erstellen (jetzt erst nach Verifizierung)
            create_trial_subscription(verified_email, verified_company)

            # Willkommens-E-Mail senden
            try:
                send_welcome_email(verified_email, verified_company, st.session_state.get("language", "de"))
            except Exception:
                pass  # E-Mail-Fehler ignorieren

            log_activity("Email Verified", f"{verified_company} - {verified_email}")

            st.markdown(f"""
            <div style="max-width: 500px; margin: 60px auto; text-align: center; padding: 40px; background: white; border-radius: 16px; box-shadow: 0 4px 20px rgba(0,0,0,0.08);">
                <div style="font-size: 48px; margin-bottom: 16px;">✅</div>
                <h2 style="color: #1d1d1f; margin-bottom: 12px;">{get_text('verify_success_title')}</h2>
                <p style="color: #86868b; font-size: 16px; margin-bottom: 24px;">{get_text('verify_success_text')}</p>
                <a href="/" style="display: inline-block; background: linear-gradient(135deg, #7c3aed 0%, #2dd4bf 100%); color: white; text-decoration: none; padding: 14px 40px; border-radius: 25px; font-weight: 600; font-size: 16px;">→ {get_text('verify_success_login')}</a>
            </div>
            """, unsafe_allow_html=True)
            st.stop()

    if st.session_state.authenticated:
        return True

    # ============ VERIFY PENDING SEITE ============
    if st.session_state.get("current_page") == "verify_pending":
        pending_email = st.session_state.get("verify_pending_email", "")
        current_lang = st.session_state.get("language", "de")

        # Header
        back_text = {"de": "← Übersicht", "en": "← Overview", "fr": "← Aperçu", "it": "← Panoramica"}
        st.markdown(f'''
        <div style="
            position: fixed; top: 0; left: 0; width: 100%;
            background: #fbfbfd; border-bottom: 1px solid #d2d2d7;
            padding: 8px 20px; display: flex; justify-content: space-between;
            align-items: center; z-index: 1000; box-sizing: border-box;
        ">
            <a href="/" style="display:flex;align-items:center;gap:9px;text-decoration:none;color:#1a1a1a;">
                <span style="font-size:18px;font-weight:700;background:linear-gradient(90deg,#7c3aed,#2dd4bf);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;">MINU-AI</span>
            </a>
            <a href="/" style="color: #86868b; text-decoration: none; font-size: 12px;">{back_text.get(current_lang, "← Overview")}</a>
        </div>
        <div style="height: 50px;"></div>
        ''', unsafe_allow_html=True)

        # E-Mail bestätigen Seite
        sent_text = get_text("verify_email_sent_text").replace("{email}", pending_email)
        st.markdown(f"""
        <div style="max-width: 500px; margin: 40px auto; text-align: center; padding: 40px; background: white; border-radius: 16px; box-shadow: 0 4px 20px rgba(0,0,0,0.08);">
            <div style="font-size: 48px; margin-bottom: 16px;">📧</div>
            <h2 style="color: #1d1d1f; margin-bottom: 12px;">{get_text('verify_email_sent_title')}</h2>
            <p style="color: #333; font-size: 16px; line-height: 1.6; margin-bottom: 16px;">{sent_text}</p>
            <p style="color: #86868b; font-size: 14px;">{get_text('verify_email_spam_hint')}</p>
        </div>
        """, unsafe_allow_html=True)

        # E-Mail erneut senden Button (mit Rate-Limiting)
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            last_resend = st.session_state.get("last_verify_resend", 0)
            seconds_since = int(datetime.now().timestamp() - last_resend)
            cooldown = 60  # 60 Sekunden Cooldown

            if seconds_since < cooldown:
                remaining = cooldown - seconds_since
                wait_text = get_text("verify_resend_wait").replace("{seconds}", str(remaining))
                st.info(wait_text)
            else:
                if st.button(f"📨 {get_text('verify_resend_button')}", use_container_width=True, key="resend_verify"):
                    # Token für diese E-Mail neu generieren und senden
                    reg = get_registration_by_email(pending_email)
                    if reg and not reg.get("verified", True):
                        token = reg.get("verification_token", "")
                        company = reg.get("company", "")
                        if token:
                            send_verification_email(pending_email, company, token, current_lang)
                            st.session_state.last_verify_resend = datetime.now().timestamp()
                            st.success(get_text("verify_resend_success"))
                            st.rerun()

            st.markdown("<br>", unsafe_allow_html=True)

            if st.button("← Zurück", use_container_width=True, key="back_from_verify"):
                st.session_state.current_page = "landing"
                st.rerun()

        return False

    # ============ REGISTRIERUNGSSEITE ============
    if st.session_state.current_page == "register":
        selected_plan = st.session_state.get("selected_plan", "free_trial")

        # Query-Parameter für Sprachauswahl verarbeiten
        query_params = st.query_params
        if "lang" in query_params:
            new_lang = query_params["lang"]
            if new_lang in ["en", "de", "fr", "it"] and new_lang != st.session_state.get("language", "en"):
                st.session_state.language = new_lang
                st.query_params.clear()
                st.rerun()

        current_lang = st.session_state.get("language", "de")

        # Produktname-Übersetzung
        product_subtitle = {"en": "Minutes AI", "de": "Protokoll KI", "fr": "Minutes IA", "it": "Verbali IA"}.get(current_lang, "Protokoll AI")

        # Plan-Namen für verschiedene Sprachen
        plan_names = {
            "de": {"free_trial": "Free Trial", "basic_solo": "Starter", "team": "PRO"},
            "en": {"free_trial": "Free Trial", "basic_solo": "Starter", "team": "PRO"},
            "fr": {"free_trial": "Essai Gratuit", "basic_solo": "Starter", "team": "PRO"},
            "it": {"free_trial": "Prova Gratuita", "basic_solo": "Starter", "team": "PRO"}
        }

        # Billing Cycle aus Session holen
        billing_cycle = st.session_state.get("billing_cycle", "monthly")
        is_yearly = billing_cycle == "yearly"

        # Plan-Details mit korrekten Preisen je nach Abrechnungszyklus
        if is_yearly:
            plan_details = {
                "free_trial": {"minutes": "60", "price": "CHF 0"},
                "basic_solo": {"minutes": "180", "price": "CHF 182/Jahr"},
                "team": {"minutes": "600", "price": "CHF 758/Jahr"}
            }
        else:
            plan_details = {
                "free_trial": {"minutes": "60", "price": "CHF 0"},
                "basic_solo": {"minutes": "180", "price": "CHF 19/Monat"},
                "team": {"minutes": "600", "price": "CHF 79/Monat"}
            }

        # Button-Text basierend auf Plan
        button_texts = {
            "de": {"free_trial": "Kostenlos starten", "basic_solo": "Starter kaufen", "team": "PRO kaufen"},
            "en": {"free_trial": "Start free", "basic_solo": "Buy Starter", "team": "Buy PRO"},
            "fr": {"free_trial": "Commencer gratuitement", "basic_solo": "Acheter Starter", "team": "Acheter PRO"},
            "it": {"free_trial": "Inizia gratis", "basic_solo": "Acquista Starter", "team": "Acquista PRO"}
        }
        register_btn_text = button_texts.get(current_lang, button_texts["de"]).get(selected_plan, "Kostenlos starten")

        # CSS für Registrierungsseite
        st.markdown("""
        <style>
            .register-container {
                max-width: 500px;
                margin: 0 auto;
                padding: 40px;
                background: white;
                border-radius: 16px;
                box-shadow: 0 4px 20px rgba(0,0,0,0.08);
            }
            .register-header {
                text-align: center;
                margin-bottom: 32px;
            }
            .register-title {
                font-size: 28px;
                font-weight: 700;
                color: #1d1d1f;
                margin-bottom: 8px;
            }
            .register-subtitle {
                font-size: 16px;
                color: #6b7280;
            }
            .plan-badge {
                display: inline-block;
                background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 100%);
                color: white;
                padding: 8px 20px;
                border-radius: 20px;
                font-weight: 600;
                font-size: 14px;
                margin-bottom: 24px;
            }
            .plan-info {
                background: #f9fafb;
                border-radius: 12px;
                padding: 16px;
                margin-bottom: 24px;
                text-align: center;
                max-width: 400px;
                margin-left: auto;
                margin-right: auto;
            }
            .plan-info-minutes {
                font-size: 32px;
                font-weight: 700;
                color: #6366f1;
            }
            .plan-info-label {
                font-size: 13px;
                color: #6b7280;
            }
            .back-link {
                text-align: center;
                margin-top: 20px;
            }
            .back-link a {
                color: #6366f1;
                text-decoration: none;
            }
        </style>
        """, unsafe_allow_html=True)

        # Header im Fusszeilen-Stil (wie Landing Page)
        back_text = {"de": "← Übersicht", "en": "← Overview", "fr": "← Aperçu", "it": "← Panoramica"}
        st.markdown(f'''
        <div style="
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            background: #fbfbfd;
            border-bottom: 1px solid #d2d2d7;
            padding: 8px 20px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            z-index: 1000;
            box-sizing: border-box;
        ">
            <div style="display: flex; align-items: center; gap: 14px;">
                <a href="/" style="display:flex;align-items:center;gap:9px;text-decoration:none;color:#1a1a1a;">
                    <span style="font-size:18px;font-weight:700;background:linear-gradient(90deg,#7c3aed,#2dd4bf);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;">MINU-AI</span>
                </a>
                <div style="display: flex; gap: 8px; margin-left: 8px;">
                    <a href="?lang=en" style="text-decoration: none; font-size: 16px; opacity: {'1' if current_lang == 'en' else '0.5'};" title="English">🇬🇧</a>
                    <a href="?lang=de" style="text-decoration: none; font-size: 16px; opacity: {'1' if current_lang == 'de' else '0.5'};" title="Deutsch">🇩🇪</a>
                    <a href="?lang=fr" style="text-decoration: none; font-size: 16px; opacity: {'1' if current_lang == 'fr' else '0.5'};" title="Français">🇫🇷</a>
                    <a href="?lang=it" style="text-decoration: none; font-size: 16px; opacity: {'1' if current_lang == 'it' else '0.5'};" title="Italiano">🇮🇹</a>
                </div>
            </div>
            <a href="?home=1" style="color: #86868b; text-decoration: none; font-size: 12px;">{back_text.get(current_lang, '← Overview')}</a>
        </div>
        <div style="height: 20px;"></div>
        ''', unsafe_allow_html=True)

        # Header
        plan_name = plan_names.get(current_lang, plan_names["de"]).get(selected_plan, "Free Trial")
        # Grossbuchstaben für STARTER und PRO
        plan_name_display = plan_name.upper() if selected_plan in ["basic_solo", "team"] else plan_name

        # Untertitel basierend auf Plan
        if selected_plan == "free_trial":
            subtitle = get_text('register_button')
        else:
            subtitle = {"de": "Abo abschliessen", "en": "Complete subscription", "fr": "Finaliser l'abonnement", "it": "Completa l'abbonamento"}.get(current_lang, "Abo abschliessen")

        # Badge oben und grösser für bezahlte Pläne
        if selected_plan in ["basic_solo", "team"]:
            st.markdown(f"""
            <div class="register-header">
                <div class="plan-badge" style="font-size: 20px; padding: 12px 32px; margin-bottom: 20px;">{plan_name_display}</div>
                <div class="register-title"><span style="font-size: 2.5em; display: block; background: linear-gradient(90deg, #7c3aed, #2dd4bf); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; font-weight: 800; font-style: italic;">MINU-AI</span>{product_subtitle}</div>
                <div class="register-subtitle">{subtitle}</div>
            </div>
            """, unsafe_allow_html=True)

            # Billing Toggle für bezahlte Pläne
            toggle_labels = {
                "de": ("Monat", "Jahr"),
                "en": ("Monthly", "Yearly"),
                "fr": ("Mois", "Année"),
                "it": ("Mese", "Anno")
            }
            toggle_monthly, toggle_yearly = toggle_labels.get(current_lang, ("Monthly", "Yearly"))

            # Initialisiere billing_cycle falls nicht vorhanden
            if "billing_cycle" not in st.session_state:
                st.session_state.billing_cycle = "monthly"

            # Toggle CSS (gleicher Style wie Startseite)
            st.markdown("""
            <style>
                /* Pill-Style Toggle - zentriert */
                div[data-testid="stHorizontalBlock"]:has(div[data-testid="stRadio"]) {
                    justify-content: center !important;
                }
                div[data-testid="stRadio"] {
                    display: flex !important;
                    justify-content: center !important;
                    width: 100% !important;
                }
                div[data-testid="stRadio"] > div {
                    background: #f3f4f6 !important;
                    border-radius: 30px !important;
                    padding: 4px !important;
                    display: inline-flex !important;
                    gap: 0 !important;
                }
                div[data-testid="stRadio"] > div > div {
                    display: flex !important;
                    gap: 0 !important;
                }
                [data-testid="column"]:has([data-testid="stRadio"]) {
                    display: flex !important;
                    justify-content: center !important;
                }
                [data-testid="column"]:has([data-testid="stRadio"]) > div {
                    display: flex !important;
                    justify-content: center !important;
                    width: 100% !important;
                }
                div[data-testid="stRadio"] label {
                    background: transparent !important;
                    border-radius: 26px !important;
                    padding: 10px 24px !important;
                    margin: 0 !important;
                    cursor: pointer !important;
                    transition: all 0.3s ease !important;
                    font-weight: 500 !important;
                    color: #6b7280 !important;
                    border: none !important;
                }
                div[data-testid="stRadio"] label:has(input:checked) {
                    background: linear-gradient(135deg, #7c3aed, #2dd4bf) !important;
                    color: white !important;
                    box-shadow: 0 2px 8px rgba(124, 58, 237, 0.3) !important;
                }
                div[data-testid="stRadio"] label:has(input:checked) span {
                    color: white !important;
                }
                div[data-testid="stRadio"] label span {
                    color: inherit !important;
                }
                div[data-testid="stRadio"] input {
                    display: none !important;
                }
                /* Hide the radio circle */
                div[data-testid="stRadio"] label > div:first-child {
                    display: none !important;
                }
            </style>
            """, unsafe_allow_html=True)

            st.markdown("<div style='height: 10px;'></div>", unsafe_allow_html=True)
            col_l, col_toggle, col_r = st.columns([1, 1, 1])
            with col_toggle:
                billing = st.radio(
                    "Abrechnungszeitraum",
                    options=["monthly", "yearly"],
                    format_func=lambda x: toggle_monthly if x == "monthly" else toggle_yearly,
                    horizontal=True,
                    label_visibility="collapsed",
                    key="billing_toggle_reg",
                    index=0 if st.session_state.billing_cycle == "monthly" else 1
                )
                st.session_state.billing_cycle = billing
        else:
            st.markdown(f"""
            <div class="register-header">
                <div class="register-title"><span style="font-size: 2.5em; display: block; background: linear-gradient(90deg, #7c3aed, #2dd4bf); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; font-weight: 800; font-style: italic;">MINU-AI</span>{product_subtitle}</div>
                <div class="register-subtitle">{subtitle}</div>
                <br>
                <div class="plan-badge">{plan_name_display}</div>
            </div>
            """, unsafe_allow_html=True)

        # Preise basierend auf Billing Cycle aktualisieren
        is_yearly = st.session_state.get("billing_cycle", "monthly") == "yearly"
        if is_yearly:
            plan_details = {
                "free_trial": {"minutes": "60", "price": "CHF 0"},
                "basic_solo": {"minutes": "180", "price": "CHF 182/Jahr"},
                "team": {"minutes": "600", "price": "CHF 758/Jahr"}
            }
        else:
            plan_details = {
                "free_trial": {"minutes": "60", "price": "CHF 0"},
                "basic_solo": {"minutes": "180", "price": "CHF 19/Monat"},
                "team": {"minutes": "600", "price": "CHF 79/Monat"}
            }

        # Plan-Info Box
        plan_info = plan_details.get(selected_plan, plan_details["free_trial"])
        per_month_label = {"de": "/Monat", "en": "/Month", "fr": "/Mois", "it": "/Mese"}.get(current_lang, "/Monat")
        st.markdown(f"""
        <div class="plan-info">
            <div class="plan-info-minutes">{plan_info['minutes']} Min.{per_month_label}</div>
            <div class="plan-info-label">{"Transkriptionsminuten" if current_lang == "de" else "Transcription minutes"}</div>
            <div style="font-size: 18px; font-weight: 600; color: #1d1d1f; margin-top: 8px;">{plan_info['price']}</div>
        </div>
        """, unsafe_allow_html=True)

        # Formular
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            company = st.text_input(
                get_text("company_label"),
                placeholder=get_text("company_placeholder"),
                label_visibility="collapsed"
            )
            email = st.text_input(
                get_text("email_label"),
                placeholder=get_text("email_placeholder_reg"),
                label_visibility="collapsed"
            )
            consent = st.checkbox(get_text("consent_checkbox"))

            st.markdown("<br>", unsafe_allow_html=True)

            if st.button(f"✓ {register_btn_text}", use_container_width=True, type="primary", key="register_submit"):
                # Prüfen ob es ein Upgrade/Downgrade ist
                is_upgrade = False
                is_downgrade = False
                is_same_paid_plan = False
                current_subscription = get_user_subscription(email)  # Direkt Subscription prüfen

                if current_subscription:
                    current_plan = current_subscription.get("plan", "free_trial")

                    # Plan-Hierarchie: free_trial < basic_solo < team
                    plan_rank = {"free_trial": 0, "basic_solo": 1, "team": 2}
                    current_rank = plan_rank.get(current_plan, 0)
                    selected_rank = plan_rank.get(selected_plan, 0)

                    if selected_rank > current_rank:
                        is_upgrade = True  # Upgrade erlaubt
                    elif selected_rank < current_rank:
                        is_downgrade = True  # Downgrade blockieren
                    elif current_plan == selected_plan and selected_plan in ["basic_solo", "team"]:
                        # Gleicher bezahlter Plan bereits aktiv
                        is_same_paid_plan = True

                if not company:
                    st.error(get_text("error_company_required"))
                elif not email or "@" not in email:
                    st.error(get_text("error_email_required"))
                elif is_downgrade:
                    # Downgrade nicht erlaubt
                    st.error(get_text("error_downgrade_not_allowed"))
                elif is_same_paid_plan:
                    # Gleicher bezahlter Plan bereits vorhanden
                    st.error(get_text("error_plan_already_active"))
                elif selected_plan == "free_trial" and current_subscription:
                    # Free Trial nur einmal möglich - E-Mail hat bereits eine Subscription
                    st.error(get_text("error_email_exists"))
                elif selected_plan == "free_trial" and is_email_already_registered(email):
                    # E-Mail bereits verifiziert registriert
                    st.error(get_text("error_email_exists"))
                elif not consent:
                    st.error(get_text("error_consent_required"))
                else:
                    try:
                        if is_upgrade:
                            # Upgrade-Flow: User ist bereits verifiziert, direkt weiter
                            st.session_state.authenticated = True
                            st.session_state.is_admin = False
                            st.session_state.user_email = email.lower().strip()
                            st.session_state.user_company = company
                            log_activity("Upgrade", f"{company} - {email.lower().strip()} - Plan: {selected_plan}")

                            # Länderprüfung und Stripe-Checkout
                            country_allowed, detected_country = is_country_allowed()
                            if not country_allowed:
                                country_name = get_country_name(detected_country, current_lang) if detected_country else "Unknown"
                                st.error(f"{get_text('error_country_blocked')} (Detected: {country_name})")
                                st.session_state.current_page = "landing"
                                st.stop()

                            billing_cycle = st.session_state.get("billing_cycle", "monthly")
                            checkout_url = create_stripe_checkout_session(email, selected_plan, billing_cycle)
                            if checkout_url:
                                st.session_state.current_page = "landing"
                                st.markdown(f'<meta http-equiv="refresh" content="0;url={checkout_url}">', unsafe_allow_html=True)
                                st.info("Weiterleitung zur Zahlung...")
                                st.stop()
                            else:
                                st.error("Fehler beim Erstellen der Checkout-Session. Bitte versuchen Sie es erneut.")
                                st.session_state.current_page = "landing"

                        elif selected_plan != "free_trial":
                            # Bezahlter Plan (Neuregistrierung): E-Mail-Verifizierung + Stripe
                            token = save_registration(company, email)
                            if token:
                                send_verification_email(email, company, token, current_lang)
                                log_activity("Registration (pending)", f"{company} - {email} - Plan: {selected_plan}")
                                st.session_state.verify_pending_email = email.lower().strip()
                                st.session_state.verify_pending_plan = selected_plan
                                st.session_state.current_page = "verify_pending"
                                st.rerun()
                            else:
                                st.error(get_text("error_email_exists"))

                        else:
                            # Free Trial: Registrierung mit E-Mail-Verifizierung
                            token = save_registration(company, email)
                            if token:
                                # Verifizierungs-E-Mail senden
                                send_verification_email(email, company, token, current_lang)
                                log_activity("Registration (pending)", f"{company} - {email} - Plan: free_trial")
                                # Zur Verifizierungs-Warteseite weiterleiten
                                st.session_state.verify_pending_email = email.lower().strip()
                                st.session_state.verify_pending_plan = selected_plan
                                st.session_state.current_page = "verify_pending"
                                st.rerun()
                            else:
                                st.error(get_text("error_email_exists"))

                    except Exception as e:
                        st.error(f"Registrierung fehlgeschlagen: {str(e)}")

            st.markdown("<br>", unsafe_allow_html=True)

        # Textlinks zentriert
        st.markdown("""
            <style>
            div[data-testid="stButton"]:has(button[kind="secondary"]) {
                display: flex !important;
                justify-content: center !important;
                width: 100% !important;
            }
            div[data-testid="stButton"]:has(button[kind="secondary"]) button {
                background: none !important;
                border: none !important;
                color: #6B7280 !important;
                font-size: 14px !important;
                box-shadow: none !important;
                padding: 3px 10px !important;
                min-height: auto !important;
                width: auto !important;
            }
            div[data-testid="stButton"]:has(button[kind="secondary"]) button:hover {
                color: #4F46E5 !important;
                text-decoration: underline !important;
                background: none !important;
            }
            </style>
        """, unsafe_allow_html=True)

        if st.button(f"🔑 {get_text('login_link')}", key="goto_login", type="secondary", use_container_width=True):
            st.session_state.current_page = "login"
            st.rerun()
        if st.button("← Zurück zur Übersicht", key="back_to_landing", type="secondary", use_container_width=True):
            st.session_state.current_page = "landing"
            st.rerun()

        return False

    # ============ LOGIN-SEITE ============
    if st.session_state.current_page == "login":
        # Query-Parameter für Sprachauswahl verarbeiten
        query_params = st.query_params
        if "lang" in query_params:
            new_lang = query_params["lang"]
            if new_lang in ["en", "de", "fr", "it"] and new_lang != st.session_state.get("language", "en"):
                st.session_state.language = new_lang
                st.query_params.clear()
                st.rerun()

        current_lang = st.session_state.get("language", "de")

        # CSS für Login-Seite (ähnlich wie Registrierung)
        st.markdown("""
        <style>
            .login-container {
                max-width: 500px;
                margin: 0 auto;
                padding: 40px;
                background: white;
                border-radius: 16px;
                box-shadow: 0 4px 20px rgba(0,0,0,0.08);
            }
            .login-header {
                text-align: center;
                margin-bottom: 32px;
            }
            .login-title {
                font-size: 28px;
                font-weight: 700;
                color: #1d1d1f;
                margin-bottom: 8px;
            }
            .login-subtitle {
                font-size: 16px;
                color: #6b7280;
            }
            .back-link {
                text-align: center;
                margin-top: 20px;
            }
            .back-link a {
                color: #6366f1;
                text-decoration: none;
            }
        </style>
        """, unsafe_allow_html=True)

        # Header im Fusszeilen-Stil (wie Landing Page)
        back_text = {"de": "← Übersicht", "en": "← Overview", "fr": "← Aperçu", "it": "← Panoramica"}
        st.markdown(f'''
        <div style="
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            background: #fbfbfd;
            border-bottom: 1px solid #d2d2d7;
            padding: 8px 20px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            z-index: 1000;
            box-sizing: border-box;
        ">
            <div style="display: flex; align-items: center; gap: 14px;">
                <a href="/" style="display:flex;align-items:center;gap:9px;text-decoration:none;color:#1a1a1a;">
                    <span style="font-size:18px;font-weight:700;background:linear-gradient(90deg,#7c3aed,#2dd4bf);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;">MINU-AI</span>
                </a>
                <div style="display: flex; gap: 8px; margin-left: 8px;">
                    <a href="?lang=en" style="text-decoration: none; font-size: 16px; opacity: {'1' if current_lang == 'en' else '0.5'};" title="English">🇬🇧</a>
                    <a href="?lang=de" style="text-decoration: none; font-size: 16px; opacity: {'1' if current_lang == 'de' else '0.5'};" title="Deutsch">🇩🇪</a>
                    <a href="?lang=fr" style="text-decoration: none; font-size: 16px; opacity: {'1' if current_lang == 'fr' else '0.5'};" title="Français">🇫🇷</a>
                    <a href="?lang=it" style="text-decoration: none; font-size: 16px; opacity: {'1' if current_lang == 'it' else '0.5'};" title="Italiano">🇮🇹</a>
                </div>
            </div>
            <a href="?home=1" style="color: #86868b; text-decoration: none; font-size: 12px;">{back_text.get(current_lang, '← Overview')}</a>
        </div>
        <div style="height: 20px;"></div>
        ''', unsafe_allow_html=True)

        # Zentrierter Login-Container
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.markdown(f"""
            <div class="login-header">
                <div class="login-title">{get_text("login_title")}</div>
                <div class="login-subtitle">{get_text("login_subtitle")}</div>
            </div>
            """, unsafe_allow_html=True)

            email = st.text_input(
                get_text("email_label"),
                placeholder=get_text("email_placeholder_reg"),
                label_visibility="collapsed",
                key="login_email"
            )

            st.markdown("<br>", unsafe_allow_html=True)

            if st.button(f"→ {get_text('login_button')}", use_container_width=True, type="primary", key="login_submit"):
                if not email or "@" not in email:
                    st.error(get_text("error_email_required"))
                elif is_email_pending_verification(email):
                    # E-Mail registriert aber noch nicht verifiziert
                    st.warning(get_text("error_email_not_verified"))
                elif not is_email_already_registered(email):
                    st.error(get_text("error_email_not_found"))
                else:
                    # Benutzer-Daten aus registrations.json laden
                    try:
                        registrations_file = PROJECT_ROOT / "registrations.json"
                        with open(registrations_file, "r", encoding="utf-8") as f:
                            registrations = json.load(f)

                        email_lower = email.lower().strip()
                        company = ""
                        for reg in registrations:
                            if reg.get("email", "").lower().strip() == email_lower:
                                company = reg.get("company", "")
                                break

                        # Session setzen (Email normalisieren für konsistenten Lookup)
                        st.session_state.authenticated = True
                        st.session_state.is_admin = False
                        st.session_state.user_email = email_lower  # Normalisierte Email verwenden
                        st.session_state.user_company = company
                        st.session_state.current_page = "landing"

                        log_activity("Login", f"{company} - {email_lower}")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Login fehlgeschlagen: {str(e)}")

            st.markdown("<br>", unsafe_allow_html=True)

            # Link zur Registrierung
            if st.button(f"📝 {get_text('register_link')}", key="goto_register", use_container_width=True):
                st.session_state.current_page = "register"
                st.session_state.selected_plan = "free_trial"
                st.rerun()

            st.markdown("<br>", unsafe_allow_html=True)

            # Zurück-Button
            if st.button("← Zurück zur Übersicht", key="back_to_landing_login", use_container_width=True):
                st.session_state.current_page = "landing"
                st.rerun()

        return False

    # ============ LANDING PAGE ============
    # Query-Parameter für Sprachauswahl verarbeiten
    query_params = st.query_params
    if "lang" in query_params:
        new_lang = query_params["lang"]
        if new_lang in ["en", "de", "fr", "it"] and new_lang != st.session_state.get("language", "en"):
            st.session_state.language = new_lang
            st.query_params.clear()
            st.rerun()

    # Login-Link aus Header behandeln
    if "login" in query_params:
        st.session_state.current_page = "login"
        st.query_params.clear()
        st.rerun()

    # Home/Übersicht-Link aus Header behandeln
    if "home" in query_params:
        st.session_state.current_page = "landing"
        st.query_params.clear()
        st.rerun()

    current_lang = st.session_state.get("language", "en")

    # Übersetzungen für Pricing
    if current_lang == "de":
        pricing_texts = {
            "free_title": "Free Trial",
            "free_subtitle": "Kostenlos testen",
            "basic_title": "Starter",
            "basic_subtitle": "Für Selbstständige",
            "team_title": "PRO",
            "team_subtitle": "Für Teams & Profis",
            "per_month": "/Monat",
            "per_year": "/Jahr",
            "minutes": "Min. Transkription",
            "minutes_month": "Min./Monat",
            "audio_files": "Audiodateien bis 200MB",
            "gdpr": "100% DSGVO-konform",
            "swiss": "Swiss Developer",
            "dialect": "Mundart & Dialekt kompatibel",
            "support": "Priority Support",
            "team_mgmt": "Team-Verwaltung",
            "start_free": "Kostenlos starten",
            "buy_plan": "Plan wählen",
            "choose_starter": "Starter wählen",
            "choose_pro": "PRO wählen",
            "popular": "Beliebt",
            "save": "spare",
            "trial_days": "3 Wochen testen",
        }
    elif current_lang == "fr":
        pricing_texts = {
            "free_title": "Essai Gratuit",
            "free_subtitle": "Testez gratuitement",
            "basic_title": "Starter",
            "basic_subtitle": "Pour indépendants",
            "team_title": "PRO",
            "team_subtitle": "Pour équipes & pros",
            "per_month": "/mois",
            "per_year": "/an",
            "minutes": "min. transcription",
            "minutes_month": "min./mois",
            "audio_files": "Fichiers audio jusqu'à 200MB",
            "gdpr": "100% conforme RGPD",
            "swiss": "Swiss Developer",
            "dialect": "Compatible dialectes",
            "support": "Support prioritaire",
            "team_mgmt": "Gestion d'équipe",
            "start_free": "Commencer gratuit",
            "buy_plan": "Choisir ce plan",
            "choose_starter": "Choisir Starter",
            "choose_pro": "Choisir PRO",
            "popular": "Populaire",
            "save": "économisez",
            "trial_days": "3 semaines d'essai",
        }
    elif current_lang == "it":
        pricing_texts = {
            "free_title": "Prova Gratuita",
            "free_subtitle": "Prova gratis",
            "basic_title": "Starter",
            "basic_subtitle": "Per autonomi",
            "team_title": "PRO",
            "team_subtitle": "Per team & professionisti",
            "per_month": "/mese",
            "per_year": "/anno",
            "minutes": "min. trascrizione",
            "minutes_month": "min./mese",
            "audio_files": "File audio fino a 200MB",
            "gdpr": "100% conforme GDPR",
            "swiss": "Swiss Developer",
            "dialect": "Compatibile con dialetti",
            "support": "Supporto prioritario",
            "team_mgmt": "Gestione team",
            "start_free": "Inizia gratis",
            "buy_plan": "Scegli piano",
            "choose_starter": "Scegli Starter",
            "choose_pro": "Scegli PRO",
            "popular": "Popolare",
            "save": "risparmia",
            "trial_days": "3 settimane di prova",
        }
    else:  # en
        pricing_texts = {
            "free_title": "Free Trial",
            "free_subtitle": "Try for free",
            "basic_title": "Starter",
            "basic_subtitle": "For self-employed",
            "team_title": "PRO",
            "team_subtitle": "For teams & pros",
            "per_month": "/month",
            "per_year": "/year",
            "minutes": "min. transcription",
            "minutes_month": "min./month",
            "audio_files": "Audio files up to 200MB",
            "gdpr": "100% GDPR compliant",
            "swiss": "Swiss Developer",
            "dialect": "Dialect compatible",
            "support": "Priority Support",
            "team_mgmt": "Team management",
            "start_free": "Start free",
            "buy_plan": "Choose plan",
            "choose_starter": "Choose Starter",
            "choose_pro": "Choose PRO",
            "popular": "Popular",
            "save": "save",
            "trial_days": "3 weeks trial",
        }

    # Header im Fusszeilen-Stil (schmal, gleicher Style)
    login_text = {"de": "Anmelden", "en": "Login", "fr": "Connexion", "it": "Accedi"}

    st.markdown(f'''
    <div class="custom-header" style="
        position: fixed;
        top: 0;
        left: 0;
        width: 100%;
        background: #fbfbfd;
        border-bottom: 1px solid #d2d2d7;
        padding: 8px 20px;
        display: flex;
        justify-content: space-between;
        align-items: center;
        z-index: 1000;
        box-sizing: border-box;
    ">
        <div style="display: flex; align-items: center; gap: 14px;">
            <a href="/" style="display:flex;align-items:center;text-decoration:none;">
                <span style="font-size:18px;font-weight:700;background:linear-gradient(90deg,#7c3aed,#2dd4bf);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;">MINU-AI</span>
            </a>
            <div style="display: flex; gap: 8px; margin-left: 8px;">
                <a href="?lang=en" style="text-decoration: none; font-size: 16px; opacity: {'1' if current_lang == 'en' else '0.5'};" title="English">🇬🇧</a>
                <a href="?lang=de" style="text-decoration: none; font-size: 16px; opacity: {'1' if current_lang == 'de' else '0.5'};" title="Deutsch">🇩🇪</a>
                <a href="?lang=fr" style="text-decoration: none; font-size: 16px; opacity: {'1' if current_lang == 'fr' else '0.5'};" title="Français">🇫🇷</a>
                <a href="?lang=it" style="text-decoration: none; font-size: 16px; opacity: {'1' if current_lang == 'it' else '0.5'};" title="Italiano">🇮🇹</a>
            </div>
        </div>
        <a href="?login=1" style="color: #86868b; text-decoration: none; font-size: 12px;">🔑 {login_text.get(current_lang, 'Login')}</a>
    </div>
    <div style="height: 5px;"></div>
    ''', unsafe_allow_html=True)

    # Titel im neuen Layout-Stil
    subtitle_text = {"en": "Minutes AI", "de": "Protokoll KI", "fr": "Minutes IA", "it": "Verbali IA"}.get(current_lang, "Minutes AI")
    st.markdown(f"""
        <h1 style="text-align: center; font-size: 2.5rem; font-weight: 700; letter-spacing: -0.02em; color: #1d1d1f; margin-bottom: 0; margin-top: 0; line-height: 1.2;">
            <span style="font-size: 5rem; display: block; margin-bottom: -10px; background: linear-gradient(90deg, #7c3aed, #2dd4bf); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; font-weight: 800; font-style: italic;">MINU-AI</span>
            <span style="background: linear-gradient(90deg, #7c3aed, #2dd4bf); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text;">
                {get_text('slogan_gradient')}
            </span><br>
            {subtitle_text}
        </h1>
        <p style="text-align: center; font-size: 1rem; color: #6b7280; margin-top: 0.5rem; margin-bottom: 1.5rem;">
            {get_text('audio_conversion_text')}
        </p>
    """, unsafe_allow_html=True)

    # Übersetzungen für Toggle (kurz, damit auf einer Zeile)
    if current_lang == "de":
        toggle_monthly = "Monat"
        toggle_yearly = "Jahr"
    elif current_lang == "fr":
        toggle_monthly = "Mois"
        toggle_yearly = "Année"
    elif current_lang == "it":
        toggle_monthly = "Mese"
        toggle_yearly = "Anno"
    else:
        toggle_monthly = "Monthly"
        toggle_yearly = "Yearly"

    # Billing Toggle mit Streamlit
    if "billing_cycle" not in st.session_state:
        st.session_state.billing_cycle = "monthly"

    # Styled Toggle CSS mit Gradient wie Popular Badge
    st.markdown("""
    <style>
        /* Pill-Style Toggle - zentriert */
        div[data-testid="stHorizontalBlock"]:has(div[data-testid="stRadio"]) {
            justify-content: center;
        }
        div[data-testid="stRadio"] {
            display: flex;
            justify-content: center;
            width: 100%;
        }
        div[data-testid="stRadio"] > div {
            background: #f3f4f6;
            border-radius: 30px;
            padding: 4px;
            display: inline-flex;
            gap: 0;
        }
        /* Zentriere den Container */
        [data-testid="column"]:has([data-testid="stRadio"]) {
            display: flex;
            justify-content: center;
        }
        [data-testid="column"]:has([data-testid="stRadio"]) > div {
            display: flex;
            justify-content: center;
            width: 100%;
        }
        div[data-testid="stRadio"] label {
            background: transparent;
            border-radius: 26px;
            padding: 10px 24px;
            margin: 0;
            cursor: pointer;
            transition: all 0.3s ease;
            font-weight: 500;
            color: #6b7280;
            border: none;
        }
        div[data-testid="stRadio"] label:has(input:checked) {
            background: linear-gradient(135deg, #7c3aed, #2dd4bf);
            color: white;
            box-shadow: 0 2px 8px rgba(124, 58, 237, 0.3);
        }
        div[data-testid="stRadio"] label:has(input:checked) span {
            color: white !important;
        }
        div[data-testid="stRadio"] label span {
            color: inherit !important;
        }
        div[data-testid="stRadio"] input {
            display: none;
        }
        /* Hide the radio circle */
        div[data-testid="stRadio"] label > div:first-child {
            display: none !important;
        }
    </style>
    """, unsafe_allow_html=True)

    # Zentrierter Toggle
    col_l, col_toggle, col_r = st.columns([1, 1, 1])
    with col_toggle:
        billing = st.radio(
            "Abrechnungszeitraum",
            options=["monthly", "yearly"],
            format_func=lambda x: toggle_monthly if x == "monthly" else toggle_yearly,
            horizontal=True,
            label_visibility="collapsed",
            key="billing_toggle"
        )
        st.session_state.billing_cycle = billing

    is_yearly = st.session_state.billing_cycle == "yearly"

    # Preise basierend auf Toggle (20% Rabatt für Jahresabo)
    if is_yearly:
        basic_price = "CHF 182"
        basic_period = pricing_texts['per_year']
        basic_note = f"= CHF 15.17{pricing_texts['per_month']}"
        team_price = "CHF 758"
        team_period = pricing_texts['per_year']
        team_note = f"= CHF 63.17{pricing_texts['per_month']}"
    else:
        basic_price = "CHF 19"
        basic_period = pricing_texts['per_month']
        basic_note = f"CHF 182{pricing_texts['per_year']} ({pricing_texts['save']} 20%)"
        team_price = "CHF 79"
        team_period = pricing_texts['per_month']
        team_note = f"CHF 758{pricing_texts['per_year']} ({pricing_texts['save']} 20%)"

    # "IDEAL FÜR" Übersetzungen (B2B / KMU fokussiert)
    if current_lang == "de":
        ideal_for = "IDEAL FÜR"
        free_ideals = ["KMU-Evaluation", "Pilotprojekte", "Testphase vor Entscheid"]
        basic_ideals = ["Selbstständige", "Kleinunternehmen", "Berater & Coaches"]
        team_ideals = ["KMU-Teams", "Abteilungen", "Mittelständische Betriebe"]
    elif current_lang == "fr":
        ideal_for = "IDÉAL POUR"
        free_ideals = ["Évaluation PME", "Projets pilotes", "Phase de test"]
        basic_ideals = ["Indépendants", "Petites entreprises", "Consultants"]
        team_ideals = ["Équipes PME", "Départements", "Entreprises moyennes"]
    elif current_lang == "it":
        ideal_for = "IDEALE PER"
        free_ideals = ["Valutazione PMI", "Progetti pilota", "Fase di test"]
        basic_ideals = ["Lavoratori autonomi", "Piccole imprese", "Consulenti"]
        team_ideals = ["Team PMI", "Reparti", "Medie imprese"]
    else:
        ideal_for = "IDEAL FOR"
        free_ideals = ["SME evaluation", "Pilot projects", "Trial before decision"]
        basic_ideals = ["Self-employed", "Small businesses", "Consultants & Coaches"]
        team_ideals = ["SME teams", "Departments", "Mid-sized companies"]

    # Pricing Cards - CSS und HTML in einem Block
    pricing_html = f"""
    <style>
        /* Pricing Buttons mit Gradient - auch auf Mobile */
        div[data-testid="stButton"] button[kind="primary"] {{
            font-weight: 700 !important;
            background: linear-gradient(135deg, #7c3aed, #2dd4bf) !important;
            color: white !important;
            border: none !important;
        }}
        div[data-testid="stButton"] button[kind="primary"]:hover {{
            background: linear-gradient(135deg, #6d28d9, #14b8a6) !important;
            box-shadow: 0 4px 15px rgba(124, 58, 237, 0.4) !important;
        }}
        div[data-testid="stButton"] button[kind="primary"] p {{
            color: white !important;
        }}

        /* Mobile: Toggle zentrieren */
        @media (max-width: 768px) {{
            div[data-testid="stRadio"] {{
                display: flex !important;
                justify-content: center !important;
                width: 100% !important;
            }}
            div[data-testid="stRadio"] > div {{
                margin: 0 auto !important;
            }}
        }}
        .pricing-container {{
            display: flex;
            justify-content: center;
            gap: 12px;
            flex-wrap: nowrap;
            margin: 20px auto;
            max-width: 100%;
        }}
        .pricing-card {{
            background: white;
            border-radius: 16px;
            padding: 24px;
            padding-bottom: 20px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.08);
            border: 1px solid #e5e7eb;
            position: relative;
            transition: transform 0.2s, box-shadow 0.2s;
            min-height: 580px;
            display: flex;
            flex-direction: column;
        }}
        .pricing-card:hover {{
            transform: translateY(-4px);
            box-shadow: 0 8px 30px rgba(0,0,0,0.12);
        }}
        .pricing-card.popular {{
            border: 2px solid #7c3aed;
        }}
        @media (max-width: 700px) {{
            .pricing-container {{
                flex-wrap: wrap;
            }}
            .pricing-card {{
                flex: 1 1 100%;
                max-width: 100%;
            }}
        }}
        .popular-badge {{
            position: absolute;
            top: -12px;
            left: 50%;
            transform: translateX(-50%);
            background: linear-gradient(135deg, #7c3aed, #2dd4bf);
            color: white;
            padding: 4px 16px;
            border-radius: 20px;
            font-size: 12px;
            font-weight: 600;
        }}
        .popular-badge-inline {{
            display: inline-block;
            background: linear-gradient(135deg, #7c3aed, #2dd4bf);
            color: white;
            padding: 4px 16px;
            border-radius: 20px;
            font-size: 12px;
            font-weight: 600;
            margin-bottom: 12px;
        }}
        .pricing-title {{
            font-size: 20px;
            font-weight: 700;
            color: #1d1d1f;
            margin-bottom: 4px;
        }}
        .pricing-subtitle {{
            font-size: 13px;
            color: #6b7280;
            margin-bottom: 16px;
        }}
        .pricing-price {{
            font-size: 36px;
            font-weight: 700;
            color: #1d1d1f;
        }}
        .pricing-price-small {{
            font-size: 16px;
            font-weight: 400;
            color: #6b7280;
        }}
        .pricing-yearly {{
            font-size: 13px;
            color: #22c55e;
            margin-top: 4px;
            margin-bottom: 16px;
        }}
        .pricing-features {{
            margin: 20px 0;
        }}
        .pricing-feature {{
            display: flex;
            align-items: center;
            gap: 8px;
            margin-bottom: 10px;
            font-size: 14px;
            color: #374151;
        }}
        .pricing-feature-icon {{
            color: #22c55e;
            font-weight: bold;
        }}
        .pricing-minutes {{
            background: linear-gradient(135deg, rgba(124, 58, 237, 0.1), rgba(45, 212, 191, 0.1));
            border-radius: 8px;
            padding: 12px;
            text-align: center;
            margin-bottom: 16px;
        }}
        .pricing-minutes-value {{
            font-size: 56px;
            font-weight: 700;
            background: linear-gradient(90deg, #7c3aed, #2dd4bf);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
        }}
        .pricing-minutes-label {{
            font-size: 12px;
            color: #6b7280;
        }}
        .stButton > button {{
            width: 100% !important;
            border-radius: 8px !important;
            padding: 14px 20px !important;
            font-weight: 600 !important;
        }}
        .pricing-features {{
            min-height: 150px;
            flex-grow: 1;
        }}
        .pricing-ideal {{
            background: #f9fafb;
            border-radius: 8px;
            padding: 12px;
            margin-top: auto;
        }}
        .pricing-ideal-title {{
            font-size: 11px;
            font-weight: 700;
            color: #7c3aed;
            letter-spacing: 0.5px;
            margin-bottom: 8px;
        }}
        .pricing-ideal-item {{
            font-size: 13px;
            color: #4b5563;
            margin-bottom: 4px;
        }}
        .pricing-ideal-item::before {{
            content: "• ";
            color: #7c3aed;
        }}
    </style>
    """
    st.markdown(pricing_html, unsafe_allow_html=True)

    # Pricing Cards
    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown(f"""
        <div class="pricing-card">
            <div class="pricing-subtitle">{pricing_texts['free_subtitle']}</div>
            <div class="pricing-title">{pricing_texts['free_title']}</div>
            <div class="pricing-price">CHF 0</div>
            <div style="color: #22c55e; font-size: 13px; margin-bottom: 16px;">{pricing_texts['trial_days']}</div>
            <div class="pricing-minutes">
                <div class="pricing-minutes-value">60</div>
                <div class="pricing-minutes-label">{pricing_texts['minutes']}</div>
            </div>
            <div class="pricing-features">
                <div class="pricing-feature"><span class="pricing-feature-icon">✓</span> {pricing_texts['audio_files']}</div>
                <div class="pricing-feature"><span class="pricing-feature-icon">✓</span> {pricing_texts['gdpr']}</div>
                <div class="pricing-feature"><span class="pricing-feature-icon">✓</span> {pricing_texts['dialect']}</div>
                <div class="pricing-feature"><span class="pricing-feature-icon">✓</span> 🇨🇭 {pricing_texts['swiss']}</div>
            </div>
            <div class="pricing-ideal">
                <div class="pricing-ideal-title">{ideal_for}</div>
                <div class="pricing-ideal-item">{free_ideals[0]}</div>
                <div class="pricing-ideal-item">{free_ideals[1]}</div>
                <div class="pricing-ideal-item">{free_ideals[2]}</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button(f"▶ {pricing_texts['start_free']}", key="btn_free", use_container_width=True, type="primary"):
            st.session_state.current_page = "register"
            st.session_state.selected_plan = "free_trial"
            st.rerun()

    with col2:
        st.markdown(f"""
        <div class="pricing-card popular">
            <div class="popular-badge">{pricing_texts['popular']}</div>
            <div class="pricing-subtitle">{pricing_texts['basic_subtitle']}</div>
            <div class="pricing-title">{pricing_texts['basic_title']}</div>
            <div class="pricing-price">{basic_price}<span class="pricing-price-small">{basic_period}</span></div>
            <div style="color: #22c55e; font-size: 13px; margin-bottom: 16px;">{basic_note}</div>
            <div class="pricing-minutes">
                <div class="pricing-minutes-value">180</div>
                <div class="pricing-minutes-label">{pricing_texts['minutes_month']}</div>
            </div>
            <div class="pricing-features">
                <div class="pricing-feature"><span class="pricing-feature-icon">✓</span> {pricing_texts['audio_files']}</div>
                <div class="pricing-feature"><span class="pricing-feature-icon">✓</span> {pricing_texts['gdpr']}</div>
                <div class="pricing-feature"><span class="pricing-feature-icon">✓</span> {pricing_texts['dialect']}</div>
                <div class="pricing-feature"><span class="pricing-feature-icon">✓</span> 🇨🇭 {pricing_texts['swiss']}</div>
            </div>
            <div class="pricing-ideal">
                <div class="pricing-ideal-title">{ideal_for}</div>
                <div class="pricing-ideal-item">{basic_ideals[0]}</div>
                <div class="pricing-ideal-item">{basic_ideals[1]}</div>
                <div class="pricing-ideal-item">{basic_ideals[2]}</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button(pricing_texts['choose_starter'], key="btn_basic", use_container_width=True, type="primary"):
            st.session_state.current_page = "register"
            st.session_state.selected_plan = "basic_solo"
            st.rerun()

    with col3:
        st.markdown(f"""
        <div class="pricing-card">
            <div class="pricing-subtitle">{pricing_texts['team_subtitle']}</div>
            <div class="pricing-title">{pricing_texts['team_title']}</div>
            <div class="pricing-price">{team_price}<span class="pricing-price-small">{team_period}</span></div>
            <div style="color: #22c55e; font-size: 13px; margin-bottom: 16px;">{team_note}</div>
            <div class="pricing-minutes">
                <div class="pricing-minutes-value">600</div>
                <div class="pricing-minutes-label">{pricing_texts['minutes_month']}</div>
            </div>
            <div class="pricing-features">
                <div class="pricing-feature"><span class="pricing-feature-icon">✓</span> {pricing_texts['audio_files']}</div>
                <div class="pricing-feature"><span class="pricing-feature-icon">✓</span> {pricing_texts['gdpr']}</div>
                <div class="pricing-feature"><span class="pricing-feature-icon">✓</span> {pricing_texts['dialect']}</div>
                <div class="pricing-feature"><span class="pricing-feature-icon">✓</span> 🇨🇭 {pricing_texts['swiss']}</div>
            </div>
            <div class="pricing-ideal">
                <div class="pricing-ideal-title">{ideal_for}</div>
                <div class="pricing-ideal-item">{team_ideals[0]}</div>
                <div class="pricing-ideal-item">{team_ideals[1]}</div>
                <div class="pricing-ideal-item">{team_ideals[2]}</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button(pricing_texts['choose_pro'], key="btn_team", use_container_width=True, type="primary"):
            st.session_state.current_page = "register"
            st.session_state.selected_plan = "team"
            st.rerun()

    # Datenschutz-Popup
    render_privacy_modal()

    return False


def is_email_already_registered(email: str) -> bool:
    """Prüft ob eine E-Mail-Adresse bereits registriert UND verifiziert wurde."""
    registrations_file = PROJECT_ROOT / "registrations.json"
    try:
        if registrations_file.exists():
            with open(registrations_file, "r", encoding="utf-8") as f:
                registrations = json.load(f)
            email_lower = email.lower().strip()
            for reg in registrations:
                if reg.get("email", "").lower().strip() == email_lower:
                    # Nur als registriert zählen wenn verifiziert
                    # Alte Einträge ohne "verified"-Feld gelten als verifiziert (Bestandsschutz)
                    return reg.get("verified", True)
        return False
    except Exception:
        return False


def is_email_pending_verification(email: str) -> bool:
    """Prüft ob eine E-Mail-Adresse registriert aber noch NICHT verifiziert ist."""
    registrations_file = PROJECT_ROOT / "registrations.json"
    try:
        if registrations_file.exists():
            with open(registrations_file, "r", encoding="utf-8") as f:
                registrations = json.load(f)
            email_lower = email.lower().strip()
            for reg in registrations:
                if reg.get("email", "").lower().strip() == email_lower:
                    return reg.get("verified") is False
        return False
    except Exception:
        return False


def generate_verification_token() -> str:
    """Generiert einen sicheren Verifizierungs-Token."""
    return secrets.token_urlsafe(32)


def save_registration(company: str, email: str) -> str:
    """Speichert die Registrierung mit Verifizierungs-Token. Gibt den Token zurück."""
    registrations_file = PROJECT_ROOT / "registrations.json"
    email_normalized = email.lower().strip()
    token = generate_verification_token()
    token_expires = (datetime.now() + timedelta(hours=48)).isoformat()

    try:
        if registrations_file.exists():
            with open(registrations_file, "r", encoding="utf-8") as f:
                registrations = json.load(f)
        else:
            registrations = []

        # Prüfen ob Email bereits existiert (Update statt Duplikat)
        existing_idx = None
        for idx, reg in enumerate(registrations):
            if reg.get("email", "").lower().strip() == email_normalized:
                existing_idx = idx
                break

        new_entry = {
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "company": company,
            "email": email_normalized,
            "verified": False,
            "verification_token": token,
            "token_expires": token_expires
        }

        if existing_idx is not None:
            # Bestehenden unverifizierten Eintrag aktualisieren (neuer Token)
            old_entry = registrations[existing_idx]
            if old_entry.get("verified", True):
                # Bereits verifiziert – nicht überschreiben
                return ""
            registrations[existing_idx] = new_entry
        else:
            registrations.append(new_entry)

        with open(registrations_file, "w", encoding="utf-8") as f:
            json.dump(registrations, f, ensure_ascii=False, indent=2)
        return token
    except Exception as e:
        print(f"Fehler beim Speichern der Registrierung: {e}")
        return ""


def verify_email_token(token: str) -> dict | None:
    """Verifiziert einen E-Mail-Token. Gibt die Registrierung zurück oder None."""
    registrations_file = PROJECT_ROOT / "registrations.json"
    try:
        if not registrations_file.exists():
            return None

        with open(registrations_file, "r", encoding="utf-8") as f:
            registrations = json.load(f)

        for idx, reg in enumerate(registrations):
            if reg.get("verification_token") == token:
                # Token gefunden – Ablauf prüfen
                token_expires = reg.get("token_expires", "")
                if token_expires:
                    expiry = datetime.fromisoformat(token_expires)
                    if datetime.now() > expiry:
                        return {"status": "expired", "email": reg.get("email", "")}

                # Token gültig – als verifiziert markieren
                registrations[idx]["verified"] = True
                registrations[idx]["verification_token"] = None  # Token löschen
                registrations[idx]["token_expires"] = None
                registrations[idx]["verified_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                with open(registrations_file, "w", encoding="utf-8") as f:
                    json.dump(registrations, f, ensure_ascii=False, indent=2)

                return {
                    "status": "success",
                    "email": reg.get("email", ""),
                    "company": reg.get("company", "")
                }

        return None  # Token nicht gefunden
    except Exception as e:
        print(f"Fehler bei Token-Verifizierung: {e}")
        return None


def get_registration_by_email(email: str) -> dict | None:
    """Holt die Registrierungsdaten für eine E-Mail-Adresse."""
    registrations_file = PROJECT_ROOT / "registrations.json"
    try:
        if registrations_file.exists():
            with open(registrations_file, "r", encoding="utf-8") as f:
                registrations = json.load(f)
            email_lower = email.lower().strip()
            for reg in registrations:
                if reg.get("email", "").lower().strip() == email_lower:
                    return reg
        return None
    except Exception:
        return None


# ============================================================================
# Feedback / Rating System
# ============================================================================

FEEDBACK_FILE = PROJECT_ROOT / "feedback_ratings.json"

def save_feedback_rating(rating: int):
    """Speichert ein Feedback-Rating in die JSON-Datei."""
    try:
        # Bestehende Daten laden
        if FEEDBACK_FILE.exists():
            with open(FEEDBACK_FILE, "r", encoding="utf-8") as f:
                feedback_data = json.load(f)
        else:
            feedback_data = {"ratings": [], "summary": {"1": 0, "2": 0, "3": 0, "4": 0, "5": 0}}

        # Neues Rating hinzufügen
        feedback_data["ratings"].append({
            "rating": rating,
            "timestamp": datetime.now().isoformat()
        })

        # Summary aktualisieren
        feedback_data["summary"][str(rating)] = feedback_data["summary"].get(str(rating), 0) + 1

        # Speichern
        with open(FEEDBACK_FILE, "w", encoding="utf-8") as f:
            json.dump(feedback_data, f, ensure_ascii=False, indent=2)

        return True
    except Exception as e:
        print(f"Fehler beim Speichern des Feedbacks: {e}")
        return False


def show_feedback_thank_you(rating: int):
    """Zeigt eine Danke-Seite nach dem Feedback."""
    lang = st.session_state.get("language", "en")

    thank_you_text = {
        "en": "Thank you for your feedback!",
        "de": "Vielen Dank für Ihr Feedback!",
        "fr": "Merci pour votre avis!",
        "it": "Grazie per il tuo feedback!"
    }

    subtitle_text = {
        "en": "Your rating helps us improve MINU Minutes AI.",
        "de": "Ihre Bewertung hilft uns, MINU Protokoll KI zu verbessern.",
        "fr": "Votre évaluation nous aide à améliorer MINU Minutes IA.",
        "it": "La tua valutazione ci aiuta a migliorare MINU Verbali IA."
    }

    st.markdown(f"""
        <div style="text-align: center; padding: 60px 20px;">
            <p style="font-size: 48px; margin-bottom: 20px;">{'⭐' * rating}</p>
            <h1 style="color: #1d1d1f; font-size: 2rem; margin-bottom: 10px;">{thank_you_text.get(lang, thank_you_text['en'])}</h1>
            <p style="color: #86868b; font-size: 1.1rem;">{subtitle_text.get(lang, subtitle_text['en'])}</p>
            <p style="margin-top: 40px;">
                <a href="{APP_URL}" style="background: #0071e3; color: white; padding: 12px 24px; border-radius: 980px; text-decoration: none; font-size: 17px;">
                    → MINU Minutes AI
                </a>
            </p>
        </div>
    """, unsafe_allow_html=True)


def show_settings_page():
    """Zeigt die Kontoeinstellungen-Seite mit Sidebar-Navigation."""
    lang = st.session_state.get("language", "de")
    user_email = st.session_state.get("user_email", "")
    user_company = st.session_state.get("user_company", "")

    # Aktuelle Settings-Sektion aus Session State
    if "settings_section" not in st.session_state:
        st.session_state.settings_section = "profile"
    settings_section = st.session_state.settings_section

    # Subscription laden
    user_subscription = get_user_subscription(user_email) if user_email else None
    user_plan = user_subscription.get("plan", "free_trial") if user_subscription else "free_trial"

    # Plan-Infos
    plan_info = {
        "free_trial": {"name": "Free Trial", "price": "CHF 0", "minutes": 60},
        "basic_solo": {"name": "Starter", "price": "CHF 19", "minutes": 180},
        "team": {"name": "PRO", "price": "CHF 79", "minutes": 600}
    }
    current_plan = plan_info.get(user_plan, plan_info["free_trial"])

    # Verbleibende Minuten
    remaining_mins = get_remaining_minutes(user_email) if user_email else 0
    used_mins = current_plan["minutes"] - int(remaining_mins)
    usage_percent = min(100, max(0, (used_mins / current_plan["minutes"]) * 100)) if current_plan["minutes"] > 0 else 0

    # Texte
    texts = {
        "de": {
            "settings": "Einstellungen",
            "profile": "Profil",
            "billing": "Abrechnung",
            "usage": "Nutzung",
            "back": "Zuruck",
            "email": "E-Mail",
            "company": "Firma / Organisation",
            "plan": "Aktueller Plan",
            "minutes_month": "Minuten pro Monat",
            "change_plan": "Plan andern",
            "payment": "Zahlungsmethode",
            "no_payment": "Keine Zahlungsmethode hinterlegt",
            "invoices": "Rechnungen",
            "no_invoices": "Keine Rechnungen vorhanden",
            "cancel": "Abo kundigen",
            "cancel_info": "Bei Kundigung bleibt der Zugang bis Periodenende bestehen.",
            "minutes_used": "Minuten verwendet",
            "minutes_left": "Minuten verbleibend",
            "renews": "Erneuert am 1. des Monats",
            "logout": "Abmelden",
            "per_month": "/ Monat"
        },
        "en": {
            "settings": "Settings",
            "profile": "Profile",
            "billing": "Billing",
            "usage": "Usage",
            "back": "Back",
            "email": "Email",
            "company": "Company / Organization",
            "plan": "Current Plan",
            "minutes_month": "Minutes per month",
            "change_plan": "Change Plan",
            "payment": "Payment Method",
            "no_payment": "No payment method on file",
            "invoices": "Invoices",
            "no_invoices": "No invoices yet",
            "cancel": "Cancel Subscription",
            "cancel_info": "Access remains until end of billing period.",
            "minutes_used": "Minutes used",
            "minutes_left": "Minutes remaining",
            "renews": "Renews on the 1st",
            "logout": "Logout",
            "per_month": "/ month"
        },
        "fr": {
            "settings": "Parametres",
            "profile": "Profil",
            "billing": "Facturation",
            "usage": "Utilisation",
            "back": "Retour",
            "email": "E-mail",
            "company": "Entreprise / Organisation",
            "plan": "Forfait actuel",
            "minutes_month": "Minutes par mois",
            "change_plan": "Changer de forfait",
            "payment": "Moyen de paiement",
            "no_payment": "Aucun moyen de paiement",
            "invoices": "Factures",
            "no_invoices": "Aucune facture",
            "cancel": "Annuler abonnement",
            "cancel_info": "Acces maintenu jusqu'a la fin de la periode.",
            "minutes_used": "Minutes utilisees",
            "minutes_left": "Minutes restantes",
            "renews": "Renouvellement le 1er",
            "logout": "Deconnexion",
            "per_month": "/ mois"
        },
        "it": {
            "settings": "Impostazioni",
            "profile": "Profilo",
            "billing": "Fatturazione",
            "usage": "Utilizzo",
            "back": "Indietro",
            "email": "E-mail",
            "company": "Azienda / Organizzazione",
            "plan": "Piano attuale",
            "minutes_month": "Minuti al mese",
            "change_plan": "Cambia piano",
            "payment": "Metodo di pagamento",
            "no_payment": "Nessun metodo di pagamento",
            "invoices": "Fatture",
            "no_invoices": "Nessuna fattura",
            "cancel": "Annulla abbonamento",
            "cancel_info": "Accesso mantenuto fino alla fine del periodo.",
            "minutes_used": "Minuti usati",
            "minutes_left": "Minuti rimanenti",
            "renews": "Rinnovo il 1",
            "logout": "Esci",
            "per_month": "/ mese"
        }
    }
    t = texts.get(lang, texts["de"])

    # Initialen
    initials = user_email[:2].upper() if user_email else "??"
    if "@" in user_email:
        name_part = user_email.split("@")[0]
        if "." in name_part:
            parts = name_part.split(".")
            initials = (parts[0][0] + parts[1][0]).upper()
        else:
            initials = name_part[:2].upper()

    # CSS fur die Settings-Seite
    st.markdown("""
    <style>
        /* Settings Page Layout */
        .settings-page {
            display: flex;
            min-height: 100vh;
            background: #f8f7f5;
        }
        .settings-sidebar {
            width: 220px;
            background: #f8f7f5;
            border-right: 1px solid #e5e5e5;
            padding: 20px 0;
            position: fixed;
            height: 100vh;
            overflow-y: auto;
        }
        .settings-main {
            margin-left: 240px;
            padding: 30px 40px;
            flex: 1;
            max-width: 600px;
        }
        .nav-section-title {
            font-size: 11px;
            font-weight: 600;
            color: #999;
            text-transform: uppercase;
            letter-spacing: 0.5px;
            padding: 8px 20px;
            margin-top: 15px;
        }
        .nav-item {
            display: block;
            padding: 8px 20px;
            color: #444;
            text-decoration: none;
            font-size: 14px;
            border-left: 3px solid transparent;
            transition: all 0.15s;
        }
        .nav-item:hover {
            background: #f0efeb;
            color: #1a1a1a;
        }
        .nav-item.active {
            background: #f0efeb;
            color: #1a1a1a;
            border-left-color: #7c3aed;
            font-weight: 500;
        }
        .section-title {
            font-size: 22px;
            font-weight: 600;
            color: #1a1a1a;
            margin-bottom: 25px;
        }
        .form-group {
            margin-bottom: 20px;
        }
        .form-label {
            font-size: 13px;
            font-weight: 500;
            color: #666;
            margin-bottom: 6px;
            display: block;
        }
        .form-value {
            font-size: 15px;
            color: #1a1a1a;
            padding: 10px 12px;
            background: #fff;
            border: 1px solid #e0e0e0;
            border-radius: 8px;
        }
        .form-input {
            width: 100%;
            padding: 10px 12px;
            font-size: 15px;
            border: 1px solid #e0e0e0;
            border-radius: 8px;
            background: #fff;
        }
        .avatar-circle {
            width: 60px;
            height: 60px;
            border-radius: 50%;
            background: linear-gradient(135deg, #7c3aed, #2dd4bf);
            color: white;
            display: flex;
            align-items: center;
            justify-content: center;
            font-weight: 700;
            font-size: 20px;
        }
        .plan-card {
            background: #fff;
            border: 1px solid #e0e0e0;
            border-radius: 10px;
            padding: 20px;
            margin-bottom: 20px;
        }
        .plan-name {
            font-size: 18px;
            font-weight: 600;
            color: #1a1a1a;
        }
        .plan-price {
            font-size: 24px;
            font-weight: 700;
            color: #7c3aed;
        }
        .usage-bar-bg {
            height: 8px;
            background: #e5e5e5;
            border-radius: 4px;
            overflow: hidden;
            margin: 10px 0;
        }
        .usage-bar-fill {
            height: 100%;
            background: #7c3aed;
            border-radius: 4px;
        }
        .invoice-row {
            display: flex;
            justify-content: space-between;
            padding: 12px 0;
            border-bottom: 1px solid #f0f0f0;
            font-size: 14px;
        }
        .invoice-row:last-child {
            border-bottom: none;
        }
        .badge-paid {
            background: #dcfce7;
            color: #16a34a;
            padding: 2px 8px;
            border-radius: 4px;
            font-size: 12px;
            font-weight: 500;
        }
        .btn-settings {
            padding: 10px 20px;
            border-radius: 8px;
            font-size: 14px;
            font-weight: 500;
            cursor: pointer;
            border: none;
            transition: all 0.15s;
        }
        .btn-primary {
            background: #7c3aed;
            color: white;
        }
        .btn-primary:hover {
            background: #6d28d9;
        }
        .btn-outline {
            background: white;
            border: 1px solid #e0e0e0;
            color: #444;
        }
        .btn-outline:hover {
            background: #f5f5f5;
        }
        .btn-danger {
            background: white;
            border: 1px solid #fecaca;
            color: #dc2626;
        }
        .btn-danger:hover {
            background: #fef2f2;
        }
        /* Hide Streamlit elements */
        .stApp > header { display: none; }
        section[data-testid="stSidebar"] { display: none; }
    </style>
    """, unsafe_allow_html=True)

    # Layout mit Columns fur Sidebar-Effekt
    col_nav, col_content = st.columns([1, 3])

    with col_nav:
        # Zuruck-Button
        if st.button(f"← {t['back']}", key="settings_back", use_container_width=True):
            st.session_state.show_settings_page = False
            st.rerun()

        st.markdown(f"<div style='font-size: 11px; font-weight: 600; color: #999; text-transform: uppercase; padding: 15px 0 8px; margin-top: 10px;'>{t['settings']}</div>", unsafe_allow_html=True)

        # Navigation Buttons
        if st.button(t['profile'], key="nav_profile", use_container_width=True, type="primary" if settings_section == "profile" else "secondary"):
            st.session_state.settings_section = "profile"
            st.rerun()

        if st.button(t['billing'], key="nav_billing", use_container_width=True, type="primary" if settings_section == "billing" else "secondary"):
            st.session_state.settings_section = "billing"
            st.rerun()

        if st.button(t['usage'], key="nav_usage", use_container_width=True, type="primary" if settings_section == "usage" else "secondary"):
            st.session_state.settings_section = "usage"
            st.rerun()

        st.markdown("<hr style='margin: 20px 0; border: none; border-top: 1px solid #e5e5e5;'>", unsafe_allow_html=True)

        if st.button(t['logout'], key="nav_logout", use_container_width=True):
            st.session_state.authenticated = False
            st.session_state.user_email = ""
            st.session_state.user_company = ""
            st.session_state.show_settings_page = False
            st.session_state.current_page = "landing"
            st.rerun()

    with col_content:
        # PROFIL SEKTION
        if settings_section == "profile":
            st.markdown(f"### {t['profile']}")
            st.markdown("")

            # Avatar
            st.markdown(f"""
            <div style="display: flex; align-items: center; gap: 20px; margin-bottom: 30px;">
                <div style="width: 70px; height: 70px; border-radius: 50%; background: linear-gradient(135deg, #7c3aed, #2dd4bf); color: white; display: flex; align-items: center; justify-content: center; font-weight: 700; font-size: 24px;">
                    {initials}
                </div>
                <div>
                    <div style="font-size: 18px; font-weight: 600; color: #1a1a1a;">{user_email}</div>
                    <div style="font-size: 14px; color: #666;">{user_company}</div>
                </div>
            </div>
            """, unsafe_allow_html=True)

            # E-Mail (read-only)
            st.markdown(f"**{t['email']}**")
            st.markdown(f"""
            <div style="padding: 12px 15px; background: #f5f5f5; border-radius: 8px; color: #666; font-size: 15px; margin-bottom: 20px;">
                {user_email}
            </div>
            """, unsafe_allow_html=True)

            # Firma (read-only)
            st.markdown(f"**{t['company']}**")
            st.markdown(f"""
            <div style="padding: 12px 15px; background: #f5f5f5; border-radius: 8px; color: #666; font-size: 15px; margin-bottom: 20px;">
                {user_company}
            </div>
            """, unsafe_allow_html=True)

            # Plan
            st.markdown(f"**{t['plan']}**")
            badge_color = {"free_trial": "#6b7280", "basic_solo": "#7c3aed", "team": "#0d9488"}.get(user_plan, "#6b7280")
            st.markdown(f"""
            <div style="display: flex; align-items: center; gap: 10px; padding: 12px 15px; background: #f5f5f5; border-radius: 8px; margin-bottom: 20px;">
                <span style="background: {badge_color}; color: white; padding: 4px 12px; border-radius: 6px; font-size: 13px; font-weight: 600;">{current_plan['name']}</span>
                <span style="color: #666; font-size: 14px;">{current_plan['minutes']} {t['minutes_month']}</span>
            </div>
            """, unsafe_allow_html=True)

        # ABRECHNUNG SEKTION
        elif settings_section == "billing":
            st.markdown(f"### {t['billing']}")
            st.markdown("")

            # Aktueller Plan
            st.markdown(f"**{t['plan']}**")
            st.markdown(f"""
            <div style="background: white; border: 1px solid #e0e0e0; border-radius: 10px; padding: 20px; margin-bottom: 25px;">
                <div style="display: flex; justify-content: space-between; align-items: center;">
                    <div>
                        <div style="font-size: 20px; font-weight: 600; color: #1a1a1a;">{current_plan['name']}</div>
                        <div style="font-size: 14px; color: #666; margin-top: 4px;">{current_plan['minutes']} {t['minutes_month']}</div>
                    </div>
                    <div style="text-align: right;">
                        <div style="font-size: 28px; font-weight: 700; color: #7c3aed;">{current_plan['price']}</div>
                        <div style="font-size: 13px; color: #999;">{t['per_month']}</div>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)

            if user_plan != "team":
                if st.button(t['change_plan'], type="primary", use_container_width=True):
                    st.query_params["upgrade"] = "1"
                    st.rerun()
                st.markdown("")

            # Zahlungsmethode
            st.markdown(f"**{t['payment']}**")
            if user_plan == "free_trial":
                st.markdown(f"""
                <div style="padding: 15px; background: #f5f5f5; border-radius: 8px; color: #666; font-size: 14px; margin-bottom: 25px;">
                    {t['no_payment']}
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <div style="background: white; border: 1px solid #e0e0e0; border-radius: 10px; padding: 15px; margin-bottom: 25px; display: flex; align-items: center; gap: 15px;">
                    <span style="background: #1a1a1a; color: white; padding: 6px 12px; border-radius: 4px; font-family: monospace; font-size: 12px; font-weight: 600;">VISA</span>
                    <div>
                        <div style="font-family: monospace; font-size: 14px;">**** **** **** 4242</div>
                        <div style="font-size: 12px; color: #999;">Exp: 12/27</div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

            # Rechnungen
            st.markdown(f"**{t['invoices']}**")
            if user_plan == "free_trial":
                st.markdown(f"""
                <div style="padding: 15px; background: #f5f5f5; border-radius: 8px; color: #666; font-size: 14px; margin-bottom: 25px;">
                    {t['no_invoices']}
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div style="background: white; border: 1px solid #e0e0e0; border-radius: 10px; overflow: hidden; margin-bottom: 25px;">
                    <div style="display: flex; padding: 12px 15px; border-bottom: 1px solid #f0f0f0;">
                        <span style="flex: 2;">13. Feb 2026</span>
                        <span style="flex: 1;">{current_plan['price']}</span>
                        <span style="flex: 1;"><span style="background: #dcfce7; color: #16a34a; padding: 2px 8px; border-radius: 4px; font-size: 12px;">Bezahlt</span></span>
                    </div>
                    <div style="display: flex; padding: 12px 15px;">
                        <span style="flex: 2;">13. Jan 2026</span>
                        <span style="flex: 1;">{current_plan['price']}</span>
                        <span style="flex: 1;"><span style="background: #dcfce7; color: #16a34a; padding: 2px 8px; border-radius: 4px; font-size: 12px;">Bezahlt</span></span>
                    </div>
                </div>
                """, unsafe_allow_html=True)

            # Abo kundigen
            if user_plan != "free_trial":
                st.markdown(f"**{t['cancel']}**")
                st.markdown(f"""
                <div style="padding: 15px; background: #fef2f2; border: 1px solid #fecaca; border-radius: 8px; margin-bottom: 15px;">
                    <div style="font-size: 14px; color: #666;">{t['cancel_info']}</div>
                </div>
                """, unsafe_allow_html=True)
                st.button(t['cancel'], type="secondary")

        # NUTZUNG SEKTION
        elif settings_section == "usage":
            st.markdown(f"### {t['usage']}")
            st.markdown("")

            # Usage Stats
            st.markdown(f"""
            <div style="background: white; border: 1px solid #e0e0e0; border-radius: 10px; padding: 25px; margin-bottom: 25px;">
                <div style="display: flex; justify-content: space-between; margin-bottom: 15px;">
                    <span style="font-size: 15px; color: #666;">{t['minutes_used']}</span>
                    <span style="font-size: 15px; font-weight: 600; color: #1a1a1a;">{used_mins} / {current_plan['minutes']}</span>
                </div>
                <div style="height: 12px; background: #e5e5e5; border-radius: 6px; overflow: hidden; margin-bottom: 15px;">
                    <div style="height: 100%; background: linear-gradient(90deg, #7c3aed, #2dd4bf); width: {usage_percent}%; border-radius: 6px;"></div>
                </div>
                <div style="display: flex; justify-content: space-between;">
                    <span style="font-size: 14px; color: #999;">{t['renews']}</span>
                    <span style="font-size: 14px; color: #7c3aed; font-weight: 600;">{int(remaining_mins)} {t['minutes_left']}</span>
                </div>
            </div>
            """, unsafe_allow_html=True)

            # Plan Info
            badge_color = {"free_trial": "#6b7280", "basic_solo": "#7c3aed", "team": "#0d9488"}.get(user_plan, "#6b7280")
            st.markdown(f"""
            <div style="background: white; border: 1px solid #e0e0e0; border-radius: 10px; padding: 20px;">
                <div style="display: flex; align-items: center; gap: 12px; margin-bottom: 15px;">
                    <span style="background: {badge_color}; color: white; padding: 5px 14px; border-radius: 6px; font-size: 14px; font-weight: 600;">{current_plan['name']}</span>
                    <span style="font-size: 20px; font-weight: 700; color: #1a1a1a;">{current_plan['price']}</span>
                    <span style="font-size: 14px; color: #999;">{t['per_month']}</span>
                </div>
                <div style="font-size: 14px; color: #666;">
                    {current_plan['minutes']} {t['minutes_month']}
                </div>
            </div>
            """, unsafe_allow_html=True)

            st.markdown("")
            if user_plan != "team":
                if st.button(t['change_plan'], type="primary", use_container_width=True):
                    st.query_params["upgrade"] = "1"
                    st.rerun()


@st.dialog("📱 App installieren / Install App", width="small")
def show_install_dialog():
    """Zeigt die Installationsanleitung für die PWA."""
    lang = st.session_state.get("language", "en")

    if lang == "de":
        st.markdown("""
### iPhone / iPad
1. Tippe unten auf das **Teilen-Symbol** ↑
2. Scrolle nach unten
3. Tippe auf **"Zum Home-Bildschirm"**
4. Tippe auf **"Hinzufügen"**

### Android
1. Tippe auf das **Menü** (⋮)
2. Wähle **"App installieren"** oder **"Zum Startbildschirm"**

### Chrome / Edge (Desktop)
1. Klicke auf das **Menü** (⋮) oben rechts
2. Wähle **"MINU Minutes AI installieren"**
        """)
    else:
        st.markdown("""
### iPhone / iPad
1. Tap the **Share button** ↑ at the bottom
2. Scroll down
3. Tap **"Add to Home Screen"**
4. Tap **"Add"**

### Android
1. Tap the **Menu** (⋮)
2. Select **"Install app"** or **"Add to Home Screen"**

### Chrome / Edge (Desktop)
1. Click the **Menu** (⋮) in the top right
2. Select **"Install MINU Minutes AI"**
        """)

    if st.button("OK", use_container_width=True):
        st.query_params.clear()
        st.rerun()


def _render_legal_topnav():
    """Rendert die vereinfachte Kopfzeile für AGB/Datenschutz-Seiten."""
    st.markdown('''
    <style>
        /* Streamlit-Standard-Padding minimieren für Legal-Seiten */
        [data-testid="stAppViewBlockContainer"] {
            padding-top: 0 !important;
        }
        section[data-testid="stMain"] > div:first-child {
            padding-top: 0 !important;
        }
        .block-container {
            padding-top: 0 !important;
        }
    </style>
    <div style="position:fixed;top:0;left:0;width:100%;background:#ffffff;border-bottom:1px solid #e4e3df;padding:0 1.5rem;height:52px;display:flex;align-items:center;justify-content:space-between;z-index:1000;box-sizing:border-box;">
        <a href="/" style="display:flex;align-items:center;text-decoration:none;">
            <span style="font-size:18px;font-weight:700;background:linear-gradient(90deg,#7c3aed,#2dd4bf);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;">MINU-AI</span>
        </a>
    </div>
    <div style="height:52px;"></div>
    ''', unsafe_allow_html=True)


def show_agb_fullpage():
    """Zeigt die AGB als Vollseiten-Ansicht mit professionellem Layout."""
    lang = st.session_state.get("language", "de")

    # Kopfzeile (Navigation)
    _render_legal_topnav()

    close_text = {"de": "Schliessen", "en": "Close", "fr": "Fermer", "it": "Chiudi"}

    # --- Professionelles CSS für AGB-Seite (gleicher Stil wie Datenschutz) ---
    st.markdown('''
    <style>
        /* Header-Banner */
        .agb-header {
            background: linear-gradient(135deg, #7c3aed 0%, #2dd4bf 100%);
            color: white;
            padding: 2rem 2rem;
            border-radius: 12px;
            margin-bottom: 1.5rem;
        }
        .agb-header__label {
            font-size: 0.82rem;
            letter-spacing: 0.12em;
            text-transform: uppercase;
            opacity: 0.75;
            margin-bottom: 0.4rem;
        }
        .agb-header h1 {
            font-size: 1.9rem;
            font-weight: 700;
            margin: 0 0 0.4rem 0;
            color: white;
        }
        .agb-header__meta {
            font-size: 0.88rem;
            opacity: 0.7;
        }
        .agb-header__close {
            float: right;
            margin-top: -2.8rem;
            padding: 8px 20px;
            background: rgba(255,255,255,0.2);
            color: white;
            text-decoration: none;
            border-radius: 8px;
            font-weight: 500;
            font-size: 0.9rem;
            border: 1px solid rgba(255,255,255,0.3);
            cursor: pointer;
            transition: background 0.2s;
        }
        .agb-header__close:hover {
            background: rgba(255,255,255,0.35);
            color: white;
        }

        /* Inhaltsverzeichnis */
        .agb-toc {
            background: #ffffff;
            border: 1px solid #e4e3df;
            border-radius: 12px;
            padding: 1.2rem 1.5rem;
            margin-bottom: 2rem;
        }
        .agb-toc__title {
            font-weight: 700;
            font-size: 0.82rem;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            color: #86868b;
            margin-bottom: 0.8rem;
        }
        .agb-toc ol {
            list-style: none;
            counter-reset: toc;
            columns: 2;
            column-gap: 2rem;
            padding: 0;
            margin: 0;
        }
        .agb-toc li {
            counter-increment: toc;
            break-inside: avoid;
            margin-bottom: 0.4rem;
        }
        .agb-toc li::before {
            content: counter(toc) ".";
            color: #7c3aed;
            font-weight: 700;
            margin-right: 0.4rem;
            font-size: 0.85rem;
        }
        .agb-toc a {
            color: #1a1a1a;
            text-decoration: none;
            font-size: 0.92rem;
            transition: color 0.2s;
        }
        .agb-toc a:hover {
            color: #7c3aed;
        }

        /* Sektionen */
        .agb-section {
            margin-bottom: 2.2rem;
            max-width: 820px;
        }
        .agb-section__heading {
            font-size: 1.25rem;
            font-weight: 700;
            color: #1a1a1a;
            margin-bottom: 0.8rem;
            padding-bottom: 0.5rem;
            border-bottom: 2px solid #e4e3df;
            display: flex;
            align-items: center;
            gap: 0.6rem;
        }
        .agb-num {
            background: #7c3aed;
            color: white;
            width: 28px;
            height: 28px;
            border-radius: 50%;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            font-size: 0.8rem;
            font-weight: 700;
            flex-shrink: 0;
        }
        .agb-section p {
            margin-bottom: 0.7rem;
            font-size: 1rem;
            line-height: 1.7;
            color: #333;
        }

        /* Kontakt-Karte */
        .agb-contact {
            background: #fff;
            border: 1px solid #e4e3df;
            border-radius: 8px;
            padding: 1rem 1.3rem;
            margin: 0.6rem 0;
        }
        .agb-contact p {
            margin-bottom: 0.25rem;
            font-size: 0.95rem;
        }
        .agb-contact a {
            color: #7c3aed;
        }

        /* Listen */
        .agb-list {
            list-style: none;
            padding-left: 0;
            margin-bottom: 1rem;
        }
        .agb-list li {
            position: relative;
            padding-left: 1.4rem;
            margin-bottom: 0.5rem;
            font-size: 0.98rem;
            line-height: 1.6;
            color: #333;
        }
        .agb-list li::before {
            content: "";
            position: absolute;
            left: 0;
            top: 0.6rem;
            width: 7px;
            height: 7px;
            background: #7c3aed;
            border-radius: 50%;
        }

        /* Highlight-Box */
        .agb-highlight {
            background: #f5f0ff;
            border-left: 4px solid #7c3aed;
            padding: 1rem 1.3rem;
            border-radius: 0 8px 8px 0;
            margin: 0.8rem 0;
            font-size: 0.95rem;
            color: #333;
            line-height: 1.6;
        }

        /* Datum unterhalb des Headers */
        .agb-meta-date {
            font-size: 0.88rem;
            color: #86868b;
            margin-bottom: 1.2rem;
        }

        /* Gesamtes Content Padding unten */
        .agb-content-wrap {
            padding-bottom: 80px;
        }

        /* Responsive */
        @media (max-width: 640px) {
            .agb-header h1 { font-size: 1.5rem; }
            .agb-toc ol { columns: 1; }
            .agb-header__close { float: none; margin-top: 0.8rem; display: inline-block; }
        }
    </style>
    ''', unsafe_allow_html=True)

    if lang == "de":
        st.markdown("""<div class="agb-content-wrap">
<div class="agb-header">
<h1>Allgemeine Geschäftsbedingungen</h1>
<a href="/" class="agb-header__close">✕ Schliessen</a>
</div>
<p class="agb-meta-date">SPEKTRUM Partner GmbH · Stand: Februar 2026</p>
<nav class="agb-toc">
<div class="agb-toc__title">Inhaltsverzeichnis</div>
<ol>
<li><a href="#a1">Geltungsbereich</a></li>
<li><a href="#a2">Leistungsbeschreibung</a></li>
<li><a href="#a3">Vertragsschluss und Registrierung</a></li>
<li><a href="#a4">Preise und Zahlung</a></li>
<li><a href="#a5">Kostenlose Testphase</a></li>
<li><a href="#a6">Kündigung</a></li>
<li><a href="#a7">Nutzungsbedingungen</a></li>
<li><a href="#a8">Datenschutz</a></li>
<li><a href="#a9">Haftung</a></li>
<li><a href="#a10">Verfügbarkeit</a></li>
<li><a href="#a11">Änderungen der AGB</a></li>
<li><a href="#a12">Anwendbares Recht und Gerichtsstand</a></li>
</ol>
</nav>
<div class="agb-section" id="a1">
<div class="agb-section__heading"><span class="agb-num">§1</span> Geltungsbereich</div>
<p>Diese Allgemeinen Geschäftsbedingungen gelten für die Nutzung der Web-Applikation MINU-AI (nachfolgend «Dienst»), bereitgestellt von der SPEKTRUM Partner GmbH, Josefstrasse 181, 8005 Zürich, Schweiz (nachfolgend «Anbieter»).</p>
<p>Mit der Registrierung oder Nutzung des Dienstes akzeptiert der Nutzer diese AGB.</p>
</div>
<div class="agb-section" id="a2">
<div class="agb-section__heading"><span class="agb-num">§2</span> Leistungsbeschreibung</div>
<p>MINU-AI ist eine webbasierte Anwendung zur automatisierten Erstellung von Meeting-Protokollen. Der Dienst umfasst:</p>
<ul class="agb-list">
<li>Transkription von Audio-Aufnahmen mittels OpenAI Whisper</li>
<li>Protokollerstellung mittels Künstlicher Intelligenz (Mistral AI)</li>
<li>Export als PDF und Word-Dokument</li>
<li>E-Mail-Versand der Protokolle</li>
</ul>
</div>
<div class="agb-section" id="a3">
<div class="agb-section__heading"><span class="agb-num">§3</span> Vertragsschluss und Registrierung</div>
<p>3.1 Der Vertrag kommt durch die Registrierung und Bestätigung der E-Mail-Adresse zustande.</p>
<p>3.2 Der Nutzer muss mindestens 18 Jahre alt sein oder die Zustimmung eines Erziehungsberechtigten haben.</p>
<p>3.3 Die Registrierung ist nur für geschäftliche oder berufliche Zwecke vorgesehen.</p>
</div>
<div class="agb-section" id="a4">
<div class="agb-section__heading"><span class="agb-num">§4</span> Preise und Zahlung</div>
<p>4.1 Die aktuellen Preise sind auf der Website ersichtlich. Alle Preise verstehen sich in Schweizer Franken (CHF) inklusive gesetzlicher Mehrwertsteuer.</p>
<p>4.2 Die Zahlung erfolgt im Voraus per Kreditkarte über den Zahlungsdienstleister Stripe.</p>
<p>4.3 Bei Abonnements verlängert sich der Vertrag automatisch um die jeweilige Laufzeit, sofern nicht rechtzeitig gekündigt wird.</p>
</div>
<div class="agb-section" id="a5">
<div class="agb-section__heading"><span class="agb-num">§5</span> Kostenlose Testphase</div>
<p>5.1 Neuen Nutzern wird eine kostenlose Testphase von 3 Wochen mit 60 Minuten Transkriptionsvolumen gewährt.</p>
<p>5.2 Nach Ablauf der Testphase endet der Zugang automatisch, sofern kein kostenpflichtiges Abonnement abgeschlossen wird.</p>
</div>
<div class="agb-section" id="a6">
<div class="agb-section__heading"><span class="agb-num">§6</span> Kündigung</div>
<p>6.1 Monatsabonnements können jederzeit zum Ende der laufenden Abrechnungsperiode gekündigt werden.</p>
<p>6.2 Jahresabonnements können jederzeit zum Ende der laufenden Vertragslaufzeit gekündigt werden.</p>
</div>
<div class="agb-section" id="a7">
<div class="agb-section__heading"><span class="agb-num">§7</span> Nutzungsbedingungen</div>
<p>7.1 Der Nutzer verpflichtet sich, den Dienst nur für rechtmässige Zwecke zu nutzen.</p>
<p>7.2 Das Hochladen von rechtswidrigen, beleidigenden oder urheberrechtlich geschützten Inhalten ohne entsprechende Berechtigung ist untersagt.</p>
<p>7.3 Der Nutzer ist für die Einholung aller erforderlichen Einwilligungen der in den Aufnahmen vorkommenden Personen verantwortlich.</p>
</div>
<div class="agb-section" id="a8">
<div class="agb-section__heading"><span class="agb-num">§8</span> Datenschutz</div>
<p>Die Verarbeitung personenbezogener Daten erfolgt gemäss unserer <a href="/?datenschutz=1" style="color:#7c3aed;">Datenschutzerklärung</a>. Der Anbieter hält die Vorgaben der EU-DSGVO und des Schweizer Datenschutzgesetzes (DSG) ein.</p>
</div>
<div class="agb-section" id="a9">
<div class="agb-section__heading"><span class="agb-num">§9</span> Haftung</div>
<p>9.1 Die Haftung des Anbieters ist – soweit gesetzlich zulässig – auf Vorsatz beschränkt. Die Haftung für Fahrlässigkeit ist ausgeschlossen.</p>
<p>9.2 Die Gesamthaftung des Anbieters ist in jedem Fall auf die Höhe der vom Nutzer in den letzten 12 Monaten bezahlten Beträge beschränkt.</p>
<p>9.3 Die Haftung für die Richtigkeit der KI-generierten Protokolle ist ausgeschlossen. Der Nutzer ist für die Überprüfung der Ergebnisse verantwortlich.</p>
<p>9.4 Der Anbieter haftet nicht für Schäden durch höhere Gewalt, Ausfall von Drittdiensten (OpenAI, Mistral AI) oder unvorhersehbare technische Störungen.</p>
<p>9.5 Die vorstehenden Haftungsbeschränkungen gelten nicht für Schäden aus der Verletzung des Lebens, des Körpers oder der Gesundheit.</p>
</div>
<div class="agb-section" id="a10">
<div class="agb-section__heading"><span class="agb-num">§10</span> Verfügbarkeit</div>
<p>Der Anbieter bemüht sich um eine hohe Verfügbarkeit des Dienstes, kann jedoch keine ununterbrochene Verfügbarkeit garantieren. Wartungsarbeiten werden nach Möglichkeit vorab angekündigt.</p>
</div>
<div class="agb-section" id="a11">
<div class="agb-section__heading"><span class="agb-num">§11</span> Änderungen der AGB</div>
<p>Der Anbieter behält sich vor, diese AGB jederzeit zu ändern. Änderungen werden dem Nutzer per E-Mail mitgeteilt. Widerspricht der Nutzer nicht innerhalb von 30 Tagen, gelten die Änderungen als akzeptiert.</p>
</div>
<div class="agb-section" id="a12">
<div class="agb-section__heading"><span class="agb-num">§12</span> Anwendbares Recht und Gerichtsstand</div>
<p>12.1 Es gilt ausschliesslich Schweizer Recht unter Ausschluss des UN-Kaufrechts.</p>
<p>12.2 Gerichtsstand ist Zürich, Schweiz. Für Verbraucher mit Wohnsitz in der EU gelten zusätzlich die zwingenden Verbraucherschutzvorschriften des Wohnsitzlandes.</p>
</div>
<div class="agb-section" id="a13" style="margin-top: 2.5rem; padding-top: 1.5rem; border-top: 2px solid #e4e3df;">
<div class="agb-section__heading"><span class="agb-num" style="background:#7c3aed;">B2B</span> Widerrufsrecht – Ausschluss</div>
<p>MINU-AI richtet sich ausschliesslich an Geschäftskunden (B2B). Ein gesetzliches Widerrufsrecht nach der EU-Verbraucherrechterichtlinie 2011/83/EU besteht daher nicht.</p>
<p>Vertragsbeendigungen sind ausschliesslich nach den Kündigungsregelungen in §6 möglich.</p>
</div>
</div>""", unsafe_allow_html=True)
    else:
        st.markdown("""<div class="agb-content-wrap">
<div class="agb-header">
<h1>General Terms and Conditions</h1>
<a href="/" class="agb-header__close">✕ Close</a>
</div>
<p class="agb-meta-date">SPEKTRUM Partner GmbH · Version: February 2026</p>
<nav class="agb-toc">
<div class="agb-toc__title">Table of Contents</div>
<ol>
<li><a href="#a1">Scope</a></li>
<li><a href="#a2">Service Description</a></li>
<li><a href="#a3">Contract Conclusion and Registration</a></li>
<li><a href="#a4">Prices and Payment</a></li>
<li><a href="#a5">Free Trial Period</a></li>
<li><a href="#a6">Cancellation</a></li>
<li><a href="#a7">Terms of Use</a></li>
<li><a href="#a8">Data Protection</a></li>
<li><a href="#a9">Liability</a></li>
<li><a href="#a10">Availability</a></li>
<li><a href="#a11">Amendments to the Terms</a></li>
<li><a href="#a12">Applicable Law and Jurisdiction</a></li>
</ol>
</nav>
<div class="agb-section" id="a1">
<div class="agb-section__heading"><span class="agb-num">§1</span> Scope</div>
<p>These General Terms and Conditions apply to the use of the web application MINU-AI (hereinafter «Service»), provided by SPEKTRUM Partner GmbH, Josefstrasse 181, 8005 Zurich, Switzerland (hereinafter «Provider»).</p>
<p>By registering or using the Service, the user accepts these Terms.</p>
</div>
<div class="agb-section" id="a2">
<div class="agb-section__heading"><span class="agb-num">§2</span> Service Description</div>
<p>MINU-AI is a web-based application for the automated creation of meeting protocols. The Service includes:</p>
<ul class="agb-list">
<li>Transcription of audio recordings using OpenAI Whisper</li>
<li>Protocol generation using Artificial Intelligence (Mistral AI)</li>
<li>Export as PDF and Word documents</li>
<li>Email delivery of protocols</li>
</ul>
</div>
<div class="agb-section" id="a3">
<div class="agb-section__heading"><span class="agb-num">§3</span> Contract Conclusion and Registration</div>
<p>3.1 The contract is concluded by registration and confirmation of the email address.</p>
<p>3.2 The user must be at least 18 years old or have the consent of a legal guardian.</p>
<p>3.3 Registration is intended for business or professional purposes only.</p>
</div>
<div class="agb-section" id="a4">
<div class="agb-section__heading"><span class="agb-num">§4</span> Prices and Payment</div>
<p>4.1 Current prices are listed on the website. All prices are in Swiss Francs (CHF) including statutory value-added tax.</p>
<p>4.2 Payment is made in advance by credit card through the payment service provider Stripe.</p>
<p>4.3 For subscriptions, the contract automatically renews for the respective term unless cancelled in time.</p>
</div>
<div class="agb-section" id="a5">
<div class="agb-section__heading"><span class="agb-num">§5</span> Free Trial Period</div>
<p>5.1 New users are granted a free trial period of 3 weeks with 60 minutes of transcription volume.</p>
<p>5.2 After the trial period expires, access ends automatically unless a paid subscription is purchased.</p>
</div>
<div class="agb-section" id="a6">
<div class="agb-section__heading"><span class="agb-num">§6</span> Cancellation</div>
<p>6.1 Monthly subscriptions can be cancelled at any time effective at the end of the current billing period.</p>
<p>6.2 Annual subscriptions can be cancelled at any time effective at the end of the current contract term.</p>
</div>
<div class="agb-section" id="a7">
<div class="agb-section__heading"><span class="agb-num">§7</span> Terms of Use</div>
<p>7.1 The user agrees to use the Service only for lawful purposes.</p>
<p>7.2 Uploading illegal, offensive, or copyright-protected content without appropriate authorization is prohibited.</p>
<p>7.3 The user is responsible for obtaining all necessary consents from persons appearing in the recordings.</p>
</div>
<div class="agb-section" id="a8">
<div class="agb-section__heading"><span class="agb-num">§8</span> Data Protection</div>
<p>The processing of personal data is carried out in accordance with our <a href="/?datenschutz=1" style="color:#7c3aed;">Privacy Policy</a>. The Provider complies with the requirements of the EU GDPR and the Swiss Data Protection Act (FADP).</p>
</div>
<div class="agb-section" id="a9">
<div class="agb-section__heading"><span class="agb-num">§9</span> Liability</div>
<p>9.1 The Provider's liability is – to the extent permitted by law – limited to intent (wilful misconduct). Liability for negligence is excluded.</p>
<p>9.2 In any event, the Provider's total liability is limited to the amounts paid by the user in the preceding 12 months.</p>
<p>9.3 Liability for the accuracy of AI-generated protocols is excluded. The user is responsible for reviewing the results.</p>
<p>9.4 The Provider is not liable for damages caused by force majeure, failure of third-party services (OpenAI, Mistral AI), or unforeseeable technical disruptions.</p>
<p>9.5 The above limitations of liability do not apply to damages arising from injury to life, body, or health.</p>
</div>
<div class="agb-section" id="a10">
<div class="agb-section__heading"><span class="agb-num">§10</span> Availability</div>
<p>The Provider strives for high availability of the Service but cannot guarantee uninterrupted availability. Maintenance work will be announced in advance where possible.</p>
</div>
<div class="agb-section" id="a11">
<div class="agb-section__heading"><span class="agb-num">§11</span> Amendments to the Terms</div>
<p>The Provider reserves the right to amend these Terms at any time. Amendments will be communicated to the user by email. If the user does not object within 30 days, the amendments are deemed accepted.</p>
</div>
<div class="agb-section" id="a12">
<div class="agb-section__heading"><span class="agb-num">§12</span> Applicable Law and Jurisdiction</div>
<p>12.1 Swiss law applies exclusively, to the exclusion of the UN Convention on Contracts for the International Sale of Goods.</p>
<p>12.2 The place of jurisdiction is Zurich, Switzerland. For consumers domiciled in the EU, the mandatory consumer protection provisions of the country of domicile additionally apply.</p>
</div>
<div class="agb-section" id="a13" style="margin-top: 2.5rem; padding-top: 1.5rem; border-top: 2px solid #e4e3df;">
<div class="agb-section__heading"><span class="agb-num" style="background:#7c3aed;">B2B</span> Right of Withdrawal – Exclusion</div>
<p>MINU-AI is intended exclusively for business customers (B2B). Therefore, no statutory right of withdrawal under EU Consumer Rights Directive 2011/83/EU applies.</p>
<p>Contract terminations are only possible in accordance with the cancellation provisions in §6.</p>
</div>
</div>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    if st.button(f"✕ {close_text.get(lang, 'Schliessen')}", use_container_width=True, key="close_bottom"):
        st.query_params.clear()
        st.rerun()


def show_datenschutz_fullpage():
    """Zeigt die Datenschutzerklärung als Vollseiten-Ansicht mit professionellem Layout."""
    lang = st.session_state.get("language", "de")

    # Kopfzeile (Navigation)
    _render_legal_topnav()

    close_text = {"de": "Schliessen", "en": "Close", "fr": "Fermer", "it": "Chiudi"}

    # --- Professionelles CSS für Datenschutzseite ---
    st.markdown('''
    <style>
        /* Header-Banner */
        .ds-header {
            background: linear-gradient(135deg, #7c3aed 0%, #2dd4bf 100%);
            color: white;
            padding: 2rem 2rem;
            border-radius: 12px;
            margin-bottom: 1.5rem;
        }
        .ds-header__label {
            font-size: 0.82rem;
            letter-spacing: 0.12em;
            text-transform: uppercase;
            opacity: 0.75;
            margin-bottom: 0.4rem;
        }
        .ds-header h1 {
            font-size: 1.9rem;
            font-weight: 700;
            margin: 0 0 0.4rem 0;
            color: white;
        }
        .ds-header__meta {
            font-size: 0.88rem;
            opacity: 0.7;
        }
        .ds-header__close {
            float: right;
            margin-top: -2.8rem;
            padding: 8px 20px;
            background: rgba(255,255,255,0.2);
            color: white;
            text-decoration: none;
            border-radius: 8px;
            font-weight: 500;
            font-size: 0.9rem;
            border: 1px solid rgba(255,255,255,0.3);
            cursor: pointer;
            transition: background 0.2s;
        }
        .ds-header__close:hover {
            background: rgba(255,255,255,0.35);
            color: white;
        }

        /* Inhaltsverzeichnis */
        .ds-toc {
            background: #ffffff;
            border: 1px solid #e4e3df;
            border-radius: 12px;
            padding: 1.2rem 1.5rem;
            margin-bottom: 2rem;
        }
        .ds-toc__title {
            font-weight: 700;
            font-size: 0.82rem;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            color: #86868b;
            margin-bottom: 0.8rem;
        }
        .ds-toc ol {
            list-style: none;
            counter-reset: toc;
            columns: 2;
            column-gap: 2rem;
            padding: 0;
            margin: 0;
        }
        .ds-toc li {
            counter-increment: toc;
            break-inside: avoid;
            margin-bottom: 0.4rem;
        }
        .ds-toc li::before {
            content: counter(toc) ".";
            color: #7c3aed;
            font-weight: 700;
            margin-right: 0.4rem;
            font-size: 0.85rem;
        }
        .ds-toc a {
            color: #1a1a1a;
            text-decoration: none;
            font-size: 0.92rem;
            transition: color 0.2s;
        }
        .ds-toc a:hover {
            color: #7c3aed;
        }

        /* Sektionen */
        .ds-section {
            margin-bottom: 2.2rem;
            max-width: 820px;
        }
        .ds-section__heading {
            font-size: 1.25rem;
            font-weight: 700;
            color: #1a1a1a;
            margin-bottom: 0.8rem;
            padding-bottom: 0.5rem;
            border-bottom: 2px solid #e4e3df;
            display: flex;
            align-items: center;
            gap: 0.6rem;
        }
        .ds-num {
            background: #7c3aed;
            color: white;
            width: 28px;
            height: 28px;
            border-radius: 50%;
            display: inline-flex;
            align-items: center;
            justify-content: center;
            font-size: 0.8rem;
            font-weight: 700;
            flex-shrink: 0;
        }
        .ds-section p {
            margin-bottom: 0.7rem;
            font-size: 1rem;
            line-height: 1.7;
            color: #333;
        }

        /* Kontakt-Karte */
        .ds-contact {
            background: #fff;
            border: 1px solid #e4e3df;
            border-radius: 8px;
            padding: 1rem 1.3rem;
            margin: 0.6rem 0;
        }
        .ds-contact p {
            margin-bottom: 0.25rem;
            font-size: 0.95rem;
        }
        .ds-contact a {
            color: #7c3aed;
        }

        /* Daten-Kategorie-Karten */
        .ds-data-cat {
            background: #fff;
            border: 1px solid #e4e3df;
            border-radius: 8px;
            padding: 1rem 1.3rem;
            margin: 0.6rem 0;
        }
        .ds-data-cat__title {
            font-weight: 700;
            color: #1a1a1a;
            margin-bottom: 0.25rem;
            font-size: 1rem;
        }
        .ds-data-cat__detail {
            font-size: 0.95rem;
            color: #333;
            margin-bottom: 0.15rem;
        }
        .ds-data-cat__legal {
            font-size: 0.85rem;
            color: #86868b;
            margin-top: 0.3rem;
        }

        /* Tabelle */
        .ds-table-wrap { overflow-x: auto; margin: 0.8rem 0; }
        .ds-table {
            width: 100%;
            border-collapse: collapse;
            font-size: 0.92rem;
        }
        .ds-table th {
            background: #7c3aed;
            color: white;
            padding: 0.65rem 0.8rem;
            text-align: left;
            font-weight: 600;
        }
        .ds-table td {
            padding: 0.65rem 0.8rem;
            border-bottom: 1px solid #e4e3df;
        }
        .ds-table tr:nth-child(even) td {
            background: #faf9fe;
        }

        /* Listen */
        .ds-list {
            list-style: none;
            padding-left: 0;
            margin-bottom: 1rem;
        }
        .ds-list li {
            position: relative;
            padding-left: 1.4rem;
            margin-bottom: 0.5rem;
            font-size: 0.98rem;
            line-height: 1.6;
            color: #333;
        }
        .ds-list li::before {
            content: "";
            position: absolute;
            left: 0;
            top: 0.6rem;
            width: 7px;
            height: 7px;
            background: #7c3aed;
            border-radius: 50%;
        }

        /* Highlight-Box */
        .ds-highlight {
            background: #f5f0ff;
            border-left: 4px solid #7c3aed;
            padding: 1rem 1.3rem;
            border-radius: 0 8px 8px 0;
            margin: 0.8rem 0;
            font-size: 0.95rem;
            color: #333;
            line-height: 1.6;
        }

        /* Nummerierte Schritte */
        .ds-steps {
            padding-left: 1.5rem;
            margin-bottom: 1rem;
        }
        .ds-steps li {
            margin-bottom: 0.4rem;
            font-size: 0.98rem;
            color: #333;
            line-height: 1.6;
        }

        /* Datum unterhalb des Headers */
        .ds-meta-date {
            font-size: 0.88rem;
            color: #86868b;
            margin-bottom: 1.2rem;
        }

        /* Gesamtes Legal-Content Padding unten */
        .ds-content-wrap {
            padding-bottom: 80px;
        }

        /* Responsive */
        @media (max-width: 640px) {
            .ds-header h1 { font-size: 1.5rem; }
            .ds-toc ol { columns: 1; }
            .ds-header__close { float: none; margin-top: 0.8rem; display: inline-block; }
        }
    </style>
    ''', unsafe_allow_html=True)

    if lang == "de":
        st.markdown("""<div class="ds-content-wrap">
<div class="ds-header">
<h1>Datenschutzerklärung</h1>
<a href="/" class="ds-header__close">✕ Schliessen</a>
</div>
<p class="ds-meta-date">SPEKTRUM Partner GmbH · Stand: Februar 2026</p>
<nav class="ds-toc">
<div class="ds-toc__title">Inhaltsverzeichnis</div>
<ol>
<li><a href="#s1">Verantwortlicher</a></li>
<li><a href="#s2">Datenschutzbeauftragter</a></li>
<li><a href="#s3">EU-Vertreter</a></li>
<li><a href="#s4">Anwendbares Recht</a></li>
<li><a href="#s5">Erhobene Daten und Zwecke</a></li>
<li><a href="#s6">Auftragsverarbeiter</a></li>
<li><a href="#s7">Datenfluss Audio-Verarbeitung</a></li>
<li><a href="#s8">Speicherdauer und Löschung</a></li>
<li><a href="#s9">Datenübermittlung in Drittländer</a></li>
<li><a href="#s10">Ihre Rechte</a></li>
<li><a href="#s11">Cookies und Tracking</a></li>
<li><a href="#s12">Automatisierte Entscheidungsfindung</a></li>
<li><a href="#s13">Datensicherheit</a></li>
<li><a href="#s14">Pflicht zur Bereitstellung</a></li>
<li><a href="#s15">Änderungen</a></li>
</ol>
</nav>
<div class="ds-section" id="s1">
<div class="ds-section__heading"><span class="ds-num">1</span> Verantwortlicher</div>
<div class="ds-contact">
<p><strong>SPEKTRUM Partner GmbH</strong></p>
<p>Josefstrasse 181, 8005 Zürich, Schweiz</p>
<p>E-Mail: <a href="mailto:minutes-ai@spekt.ch">minutes-ai@spekt.ch</a></p>
<p>Website: <a href="https://minu-ai.ch" target="_blank" rel="noopener">https://minu-ai.ch</a></p>
</div>
</div>
<div class="ds-section" id="s2">
<div class="ds-section__heading"><span class="ds-num">2</span> Datenschutzbeauftragter</div>
<p>Ein Datenschutzbeauftragter ist nach Art. 37 DSGVO nicht bestellt, da die Voraussetzungen nicht erfüllt sind. Bei Fragen zum Datenschutz wenden Sie sich bitte an: <a href="mailto:minutes-ai@spekt.ch" style="color:#7c3aed;">minutes-ai@spekt.ch</a></p>
</div>
<div class="ds-section" id="s3">
<div class="ds-section__heading"><span class="ds-num">3</span> EU-Vertreter (Art. 27 DSGVO)</div>
<p>Da die SPEKTRUM Partner GmbH ihren Sitz ausserhalb der EU hat und Dienstleistungen an Personen im EWR anbietet, haben wir gemäss Art. 27 DSGVO folgenden Vertreter in der Europäischen Union benannt:</p>
<div class="ds-contact" style="border-left: 4px solid #7c3aed;">
<p><strong>VGS Datenschutzpartner GmbH</strong></p>
<p>Am Kaiserkai 69, 20457 Hamburg, Deutschland</p>
<p>E-Mail: <a href="mailto:info@datenschutzpartner.eu" style="color:#7c3aed;">info@datenschutzpartner.eu</a></p>
</div>
<p>Der EU-Vertreter dient als Anlaufstelle für Aufsichtsbehörden und betroffene Personen in der EU bei sämtlichen Fragen im Zusammenhang mit der Datenverarbeitung.</p>
</div>
<div class="ds-section" id="s4">
<div class="ds-section__heading"><span class="ds-num">4</span> Anwendbares Recht und Geltungsbereich</div>
<p>Diese Datenschutzerklärung gilt für die Nutzung der App «MINU-AI» und der Website minu-ai.ch.</p>
<p>Die Verarbeitung personenbezogener Daten erfolgt in Übereinstimmung mit dem <strong>Schweizer Bundesgesetz über den Datenschutz (DSG/revDSG)</strong> sowie, soweit anwendbar, der <strong>Europäischen Datenschutz-Grundverordnung (DSGVO)</strong>.</p>
<p>Das Schweizer DSG findet Anwendung, da der Sitz unseres Unternehmens in der Schweiz liegt. Die DSGVO findet zusätzlich Anwendung, soweit wir Dienstleistungen an Personen im Europäischen Wirtschaftsraum (EWR) anbieten (Art. 3 Abs. 2 DSGVO).</p>
</div>
<div class="ds-section" id="s5">
<div class="ds-section__heading"><span class="ds-num">5</span> Erhobene Daten und Zwecke</div>
<p>Wir verarbeiten folgende personenbezogene Daten:</p>
<div class="ds-data-cat">
<div class="ds-data-cat__title">Registrierungsdaten</div>
<div class="ds-data-cat__detail">Firmenname, E-Mail-Adresse</div>
<div class="ds-data-cat__detail">Zweck: Vertragserfüllung, Kontozugang, Kommunikation</div>
<div class="ds-data-cat__legal">Rechtsgrundlage: Art. 6 Abs. 1 lit. b DSGVO (Vertragserfüllung); Art. 31 Abs. 1 DSG</div>
</div>
<div class="ds-data-cat">
<div class="ds-data-cat__title">Audiodaten</div>
<div class="ds-data-cat__detail">Hochgeladene Audio-Dateien zur Transkription</div>
<div class="ds-data-cat__detail">Zweck: Erbringung der Dienstleistung (Transkription und Protokollerstellung)</div>
<div class="ds-data-cat__legal">Rechtsgrundlage: Art. 6 Abs. 1 lit. b DSGVO (Vertragserfüllung); Art. 31 Abs. 1 DSG</div>
</div>
<div class="ds-data-cat">
<div class="ds-data-cat__title">Technische Daten</div>
<div class="ds-data-cat__detail">IP-Adresse, Browser-Typ, Betriebssystem, Zugriffszeiten</div>
<div class="ds-data-cat__detail">Zweck: Sicherheit, Fehleranalyse</div>
<div class="ds-data-cat__legal">Rechtsgrundlage: Art. 6 Abs. 1 lit. f DSGVO (berechtigtes Interesse); Art. 31 Abs. 1 DSG</div>
</div>
</div>
<div class="ds-section" id="s6">
<div class="ds-section__heading"><span class="ds-num">6</span> Auftragsverarbeiter (Subprozessoren)</div>
<p>Für die Erbringung unserer Dienstleistungen setzen wir sorgfältig ausgewählte Drittanbieter (Auftragsverarbeiter/Subprozessoren) ein. Mit jedem dieser Anbieter besteht ein Auftragsverarbeitungsvertrag (DPA) gemäss Art. 28 DSGVO bzw. Art. 9 DSG.</p>
<div class="ds-table-wrap">
<table class="ds-table">
<thead>
<tr><th>Anbieter</th><th>Standort</th><th>Zweck</th><th>Datenschutz</th></tr>
</thead>
<tbody>
<tr><td>OpenAI, Inc.</td><td>USA</td><td>Whisper API – Transkription</td><td>DPA, SCCs, EU-US DPF</td></tr>
<tr><td>Mistral AI SAS</td><td>Frankreich (EU)</td><td>KI-Protokollerstellung</td><td>DPA, DSGVO nativ</td></tr>
<tr><td>Stripe, Inc.</td><td>USA</td><td>Zahlungsabwicklung</td><td>DPA, PCI DSS, SCCs</td></tr>
<tr><td>Hetzner Online GmbH</td><td>Deutschland (EU)</td><td>Server-Hosting</td><td>DPA, ISO 27001</td></tr>
<tr><td>Umami (Self-hosted)</td><td>Deutschland (EU)</td><td>Datenschutzfreundliche Webanalyse</td><td>Self-hosted, keine PII</td></tr>
</tbody>
</table>
</div>
<p>Audio-Dateien, die zur Transkription hochgeladen werden, werden temporär an OpenAI (USA) übermittelt und nach der Verarbeitung sofort von unserem Server gelöscht. OpenAI speichert API-Daten standardmässig bis zu 30 Tage für Sicherheitsüberwachung (Abuse Monitoring). Daten werden von OpenAI nicht für Modelltraining verwendet.</p>
<p>Die Protokollerstellung erfolgt durch Mistral AI (Frankreich/EU). Mistral unterliegt direkt der DSGVO und verarbeitet ausschliesslich Text – keine Audio-Dateien.</p>
</div>
<div class="ds-section" id="s7">
<div class="ds-section__heading"><span class="ds-num">7</span> Datenfluss Audio-Verarbeitung</div>
<p>Die Verarbeitung Ihrer Audio-Dateien erfolgt in folgenden Schritten:</p>
<ol class="ds-steps">
<li>Sie laden eine Audio-Datei in MINU-AI hoch</li>
<li>Die Datei wird temporär auf unserem Server (Hetzner, Deutschland) zwischengespeichert</li>
<li>Die Audio-Datei wird verschlüsselt (TLS) an die OpenAI Whisper API gesendet</li>
<li>OpenAI gibt den transkribierten Text zurück</li>
<li><strong>Die Audio-Datei wird sofort von unserem Server gelöscht</strong></li>
<li>Der Transkriptionstext wird verschlüsselt an Mistral AI gesendet</li>
<li>Mistral AI erstellt ein strukturiertes Protokoll und gibt es zurück</li>
<li>Das Protokoll wird Ihnen angezeigt – keine dauerhafte serverseitige Speicherung</li>
</ol>
<div class="ds-highlight">
<strong>Wichtig:</strong> MINU-AI speichert weder Audio-Dateien noch Transkripte dauerhaft auf dem Server. Die Audio-Datei existiert nur während des aktiven Verarbeitungsprozesses.
</div>
</div>
<div class="ds-section" id="s8">
<div class="ds-section__heading"><span class="ds-num">8</span> Speicherdauer und Löschung</div>
<ul class="ds-list">
<li><strong>Audiodaten:</strong> Werden nach Verarbeitung sofort gelöscht, keine dauerhafte Speicherung</li>
<li><strong>Transkripte/Protokolle:</strong> Werden nicht serverseitig gespeichert, nur lokal beim Nutzer</li>
<li><strong>Kontodaten:</strong> Bis zur Löschung des Kontos oder 3 Jahre nach letzter Aktivität</li>
<li><strong>Rechnungsdaten:</strong> 10 Jahre (gesetzliche Aufbewahrungspflicht)</li>
<li><strong>Bei OpenAI (Standard):</strong> Bis zu 30 Tage für Abuse Monitoring (ZDR-Aktivierung in Prüfung)</li>
<li><strong>Bei Mistral AI:</strong> Keine Speicherung über die Verarbeitung hinaus bei API-Nutzung</li>
</ul>
</div>
<div class="ds-section" id="s9">
<div class="ds-section__heading"><span class="ds-num">9</span> Datenübermittlung in Drittländer</div>
<p>Bei Nutzung von US-Diensten (OpenAI, Stripe) erfolgt eine Datenübermittlung in die USA. Diese erfolgt auf Basis von:</p>
<ul class="ds-list">
<li>EU-US Data Privacy Framework (DPF)</li>
<li>Standardvertragsklauseln (SCCs)</li>
<li>Zusätzliche technische Schutzmassnahmen (Verschlüsselung in Transit)</li>
</ul>
<p>Die Verarbeitung durch Mistral AI (Frankreich) und Hetzner (Deutschland) erfolgt vollständig innerhalb der EU. Ein Drittlandtransfer findet hier nicht statt.</p>
</div>
<div class="ds-section" id="s10">
<div class="ds-section__heading"><span class="ds-num">10</span> Ihre Rechte</div>
<p>Sie haben folgende Rechte:</p>
<div class="ds-table-wrap">
<table class="ds-table">
<thead>
<tr><th>Recht</th><th>DSGVO</th><th>Schweizer DSG</th></tr>
</thead>
<tbody>
<tr><td>Auskunft</td><td>Art. 15 DSGVO</td><td>Art. 25 DSG</td></tr>
<tr><td>Berichtigung</td><td>Art. 16 DSGVO</td><td>Art. 32 Abs. 1 DSG</td></tr>
<tr><td>Löschung</td><td>Art. 17 DSGVO</td><td>Art. 32 Abs. 2 lit. c DSG</td></tr>
<tr><td>Einschränkung</td><td>Art. 18 DSGVO</td><td>–</td></tr>
<tr><td>Datenübertragbarkeit</td><td>Art. 20 DSGVO</td><td>Art. 28 DSG</td></tr>
<tr><td>Widerspruch</td><td>Art. 21 DSGVO</td><td>Art. 32 Abs. 2 lit. b DSG</td></tr>
<tr><td>Beschwerde</td><td>Zuständige Aufsichtsbehörde</td><td>EDÖB (s. unten)</td></tr>
</tbody>
</table>
</div>
<p>Kontakt für Anfragen: <a href="mailto:minutes-ai@spekt.ch" style="color:#7c3aed;">minutes-ai@spekt.ch</a></p>
<p><strong>Beschwerderecht:</strong> Sie haben das Recht, sich bei der zuständigen Aufsichtsbehörde zu beschweren:</p>
<ul class="ds-list">
<li><strong>Schweiz:</strong> Eidgenössischer Datenschutz- und Öffentlichkeitsbeauftragter (EDÖB), Feldeggweg 1, 3003 Bern – <a href="https://www.edoeb.admin.ch" target="_blank" rel="noopener" style="color:#7c3aed;">www.edoeb.admin.ch</a></li>
<li><strong>EU:</strong> Die für Sie zuständige Aufsichtsbehörde gemäss <a href="https://edpb.europa.eu/about-edpb/about-edpb/members_de" target="_blank" rel="noopener" style="color:#7c3aed;">EDPB-Mitgliederliste</a></li>
</ul>
</div>
<div class="ds-section" id="s11">
<div class="ds-section__heading"><span class="ds-num">11</span> Cookies und Tracking</div>
<ul class="ds-list">
<li>Wir verwenden ausschliesslich technisch notwendige Session-Cookies</li>
<li>Für die Webseitenanalyse nutzen wir <strong>Umami Analytics</strong> – ein datenschutzfreundliches Tool ohne Cookies, ohne persönliche Datenerfassung und ohne Tracking über Webseiten hinweg</li>
<li>Kein Google Analytics, keine Werbe-Cookies, keine Weitergabe an Dritte</li>
</ul>
<p>Rechtsgrundlage: Art. 6 Abs. 1 lit. f DSGVO (berechtigtes Interesse); Art. 31 Abs. 1 DSG</p>
</div>
<div class="ds-section" id="s12">
<div class="ds-section__heading"><span class="ds-num">12</span> Automatisierte Entscheidungsfindung</div>
<p>Die KI-gestützte Transkription und Protokollerstellung stellt keine automatisierte Entscheidungsfindung im Sinne von Art. 22 DSGVO dar, da keine rechtlichen oder ähnlich erheblichen Wirkungen für Sie entstehen. Die Ergebnisse dienen als Hilfsmittel und bedürfen stets der Überprüfung durch den Nutzer.</p>
</div>
<div class="ds-section" id="s13">
<div class="ds-section__heading"><span class="ds-num">13</span> Datensicherheit</div>
<p>Wir setzen dem Stand der Technik entsprechende technische und organisatorische Massnahmen ein, um Ihre Daten zu schützen. Dazu gehören insbesondere:</p>
<ul class="ds-list">
<li>Verschlüsselung aller Datenübertragungen (TLS 1.2+)</li>
<li>Zugriffskontrolle und Authentifizierung</li>
<li>Regelmässige Sicherheitsupdates</li>
<li>Sofortige Löschung von Audio-Dateien nach Verarbeitung</li>
<li>Keine dauerhafte serverseitige Speicherung von Transkripten</li>
</ul>
</div>
<div class="ds-section" id="s14">
<div class="ds-section__heading"><span class="ds-num">14</span> Pflicht zur Bereitstellung</div>
<p>Die Bereitstellung der Registrierungsdaten (Firmenname, E-Mail) ist für den Vertragsschluss erforderlich. Ohne diese Daten können wir den Dienst nicht erbringen. Die Bereitstellung von Audio-Dateien erfolgt freiwillig im Rahmen der jeweiligen Nutzung.</p>
</div>
<div class="ds-section" id="s15">
<div class="ds-section__heading"><span class="ds-num">15</span> Änderungen</div>
<p>Wir behalten uns vor, diese Datenschutzerklärung jederzeit anzupassen, insbesondere bei Änderungen der Rechtslage, neuen Funktionen oder neuen Subprozessoren.</p>
<p>Die aktuelle Fassung ist stets abrufbar. Bei wesentlichen Änderungen informieren wir registrierte Nutzer per E-Mail.</p>
</div>
</div>""", unsafe_allow_html=True)
    else:
        st.markdown("""<div class="ds-content-wrap">
<div class="ds-header">
<h1>Privacy Policy</h1>
<a href="/" class="ds-header__close">✕ Close</a>
</div>
<p class="ds-meta-date">SPEKTRUM Partner GmbH · Version: February 2026</p>
<nav class="ds-toc">
<div class="ds-toc__title">Table of Contents</div>
<ol>
<li><a href="#s1">Controller</a></li>
<li><a href="#s2">Data Protection Officer</a></li>
<li><a href="#s3">EU Representative</a></li>
<li><a href="#s4">Applicable Law</a></li>
<li><a href="#s5">Data Collected and Purposes</a></li>
<li><a href="#s6">Processors</a></li>
<li><a href="#s7">Audio Data Flow</a></li>
<li><a href="#s8">Storage and Deletion</a></li>
<li><a href="#s9">Third Country Transfer</a></li>
<li><a href="#s10">Your Rights</a></li>
<li><a href="#s11">Cookies and Tracking</a></li>
<li><a href="#s12">Automated Decision-Making</a></li>
<li><a href="#s13">Data Security</a></li>
<li><a href="#s14">Obligation to Provide Data</a></li>
<li><a href="#s15">Changes</a></li>
</ol>
</nav>
<div class="ds-section" id="s1">
<div class="ds-section__heading"><span class="ds-num">1</span> Controller</div>
<div class="ds-contact">
<p><strong>SPEKTRUM Partner GmbH</strong></p>
<p>Josefstrasse 181, 8005 Zurich, Switzerland</p>
<p>Email: <a href="mailto:minutes-ai@spekt.ch">minutes-ai@spekt.ch</a></p>
<p>Website: <a href="https://minu-ai.ch" target="_blank" rel="noopener">https://minu-ai.ch</a></p>
</div>
</div>
<div class="ds-section" id="s2">
<div class="ds-section__heading"><span class="ds-num">2</span> Data Protection Officer</div>
<p>A Data Protection Officer is not appointed as the requirements under Art. 37 GDPR are not met. For privacy questions, please contact: <a href="mailto:minutes-ai@spekt.ch" style="color:#7c3aed;">minutes-ai@spekt.ch</a></p>
</div>
<div class="ds-section" id="s3">
<div class="ds-section__heading"><span class="ds-num">3</span> EU Representative (Art. 27 GDPR)</div>
<p>As SPEKTRUM Partner GmbH is based outside the EU and offers services to individuals in the EEA, we have appointed a representative in the European Union in accordance with Art. 27 GDPR:</p>
<div class="ds-contact" style="border-left: 4px solid #7c3aed;">
<p><strong>VGS Datenschutzpartner GmbH</strong></p>
<p>Am Kaiserkai 69, 20457 Hamburg, Germany</p>
<p>Email: <a href="mailto:info@datenschutzpartner.eu" style="color:#7c3aed;">info@datenschutzpartner.eu</a></p>
<p>Website: <a href="https://www.datenschutzpartner.eu" target="_blank" rel="noopener" style="color:#7c3aed;">www.datenschutzpartner.eu</a></p>
</div>
<p>The EU representative serves as a point of contact for supervisory authorities and data subjects in the EU for all data processing matters.</p>
</div>
<div class="ds-section" id="s4">
<div class="ds-section__heading"><span class="ds-num">4</span> Applicable Law and Scope</div>
<p>This privacy policy applies to the use of the «MINU-AI» app and the website minu-ai.ch.</p>
<p>The processing of personal data is carried out in accordance with the <strong>Swiss Federal Act on Data Protection (FADP/revFADP)</strong> and, where applicable, the <strong>European General Data Protection Regulation (GDPR)</strong>.</p>
<p>The Swiss FADP applies as our company is based in Switzerland. The GDPR additionally applies insofar as we offer services to individuals in the European Economic Area (EEA) (Art. 3(2) GDPR).</p>
</div>
<div class="ds-section" id="s5">
<div class="ds-section__heading"><span class="ds-num">5</span> Data Collected and Purposes</div>
<p>We process the following personal data:</p>
<div class="ds-data-cat">
<div class="ds-data-cat__title">Registration Data</div>
<div class="ds-data-cat__detail">Company name, email address</div>
<div class="ds-data-cat__detail">Purpose: Contract fulfillment, account access, communication</div>
<div class="ds-data-cat__legal">Legal basis: Art. 6(1)(b) GDPR (contract performance); Art. 31(1) FADP</div>
</div>
<div class="ds-data-cat">
<div class="ds-data-cat__title">Audio Data</div>
<div class="ds-data-cat__detail">Uploaded audio files for transcription</div>
<div class="ds-data-cat__detail">Purpose: Service provision (transcription and protocol generation)</div>
<div class="ds-data-cat__legal">Legal basis: Art. 6(1)(b) GDPR (contract performance); Art. 31(1) FADP</div>
</div>
<div class="ds-data-cat">
<div class="ds-data-cat__title">Technical Data</div>
<div class="ds-data-cat__detail">IP address, browser type, operating system, access times</div>
<div class="ds-data-cat__detail">Purpose: Security, error analysis</div>
<div class="ds-data-cat__legal">Legal basis: Art. 6(1)(f) GDPR (legitimate interest); Art. 31(1) FADP</div>
</div>
</div>
<div class="ds-section" id="s6">
<div class="ds-section__heading"><span class="ds-num">6</span> Processors (Sub-processors)</div>
<p>We engage carefully selected third-party providers (processors/sub-processors) for our services. A data processing agreement (DPA) is in place with each provider in accordance with Art. 28 GDPR and Art. 9 FADP.</p>
<div class="ds-table-wrap">
<table class="ds-table">
<thead>
<tr><th>Provider</th><th>Location</th><th>Purpose</th><th>Data Protection</th></tr>
</thead>
<tbody>
<tr><td>OpenAI, Inc.</td><td>USA</td><td>Whisper API – Transcription</td><td>DPA, SCCs, EU-US DPF</td></tr>
<tr><td>Mistral AI SAS</td><td>France (EU)</td><td>AI protocol generation</td><td>DPA, GDPR native</td></tr>
<tr><td>Stripe, Inc.</td><td>USA</td><td>Payment processing</td><td>DPA, PCI DSS, SCCs</td></tr>
<tr><td>Hetzner Online GmbH</td><td>Germany (EU)</td><td>Server hosting</td><td>DPA, ISO 27001</td></tr>
<tr><td>Umami (Self-hosted)</td><td>Germany (EU)</td><td>Privacy-friendly web analytics</td><td>Self-hosted, no PII</td></tr>
</tbody>
</table>
</div>
<p>Audio files uploaded for transcription are temporarily transmitted to OpenAI (USA) and immediately deleted from our server after processing. OpenAI stores API data for up to 30 days by default for abuse monitoring. Data is not used by OpenAI for model training.</p>
<p>Protocol generation is performed by Mistral AI (France/EU). Mistral is directly subject to the GDPR and processes only text – no audio files.</p>
</div>
<div class="ds-section" id="s7">
<div class="ds-section__heading"><span class="ds-num">7</span> Audio Data Flow</div>
<p>Your audio files are processed in the following steps:</p>
<ol class="ds-steps">
<li>You upload an audio file to MINU-AI</li>
<li>The file is temporarily stored on our server (Hetzner, Germany)</li>
<li>The audio file is sent encrypted (TLS) to the OpenAI Whisper API</li>
<li>OpenAI returns the transcribed text</li>
<li><strong>The audio file is immediately deleted from our server</strong></li>
<li>The transcription text is sent encrypted to Mistral AI</li>
<li>Mistral AI creates a structured protocol and returns it</li>
<li>The protocol is displayed to you – no permanent server-side storage</li>
</ol>
<div class="ds-highlight">
<strong>Important:</strong> MINU-AI does not permanently store audio files or transcripts on the server. The audio file exists only during the active processing.
</div>
</div>
<div class="ds-section" id="s8">
<div class="ds-section__heading"><span class="ds-num">8</span> Storage Duration and Deletion</div>
<ul class="ds-list">
<li><strong>Audio data:</strong> Deleted immediately after processing, no permanent storage</li>
<li><strong>Transcripts/Protocols:</strong> Not stored server-side, only locally by the user</li>
<li><strong>Account data:</strong> Until account deletion or 3 years after last activity</li>
<li><strong>Billing data:</strong> 10 years (legal retention requirement)</li>
<li><strong>At OpenAI (default):</strong> Up to 30 days for abuse monitoring (ZDR activation under review)</li>
<li><strong>At Mistral AI:</strong> No storage beyond processing for API usage</li>
</ul>
</div>
<div class="ds-section" id="s9">
<div class="ds-section__heading"><span class="ds-num">9</span> Data Transfer to Third Countries</div>
<p>When using US services (OpenAI, Stripe), data is transferred to the USA. This is based on:</p>
<ul class="ds-list">
<li>EU-US Data Privacy Framework (DPF)</li>
<li>Standard Contractual Clauses (SCCs)</li>
<li>Additional technical safeguards (encryption in transit)</li>
</ul>
<p>Processing by Mistral AI (France) and Hetzner (Germany) takes place entirely within the EU. No third-country transfer occurs here.</p>
</div>
<div class="ds-section" id="s10">
<div class="ds-section__heading"><span class="ds-num">10</span> Your Rights</div>
<p>You have the following rights:</p>
<div class="ds-table-wrap">
<table class="ds-table">
<thead>
<tr><th>Right</th><th>GDPR</th><th>Swiss FADP</th></tr>
</thead>
<tbody>
<tr><td>Access</td><td>Art. 15 GDPR</td><td>Art. 25 FADP</td></tr>
<tr><td>Rectification</td><td>Art. 16 GDPR</td><td>Art. 32(1) FADP</td></tr>
<tr><td>Erasure</td><td>Art. 17 GDPR</td><td>Art. 32(2)(c) FADP</td></tr>
<tr><td>Restriction</td><td>Art. 18 GDPR</td><td>–</td></tr>
<tr><td>Data portability</td><td>Art. 20 GDPR</td><td>Art. 28 FADP</td></tr>
<tr><td>Objection</td><td>Art. 21 GDPR</td><td>Art. 32(2)(b) FADP</td></tr>
<tr><td>Complaint</td><td>Competent supervisory authority</td><td>FDPIC (see below)</td></tr>
</tbody>
</table>
</div>
<p>Contact for inquiries: <a href="mailto:minutes-ai@spekt.ch" style="color:#7c3aed;">minutes-ai@spekt.ch</a></p>
<p><strong>Right to complain:</strong> You have the right to lodge a complaint with the competent supervisory authority:</p>
<ul class="ds-list">
<li><strong>Switzerland:</strong> Federal Data Protection and Information Commissioner (FDPIC), Feldeggweg 1, 3003 Bern – <a href="https://www.edoeb.admin.ch" target="_blank" rel="noopener" style="color:#7c3aed;">www.edoeb.admin.ch</a></li>
<li><strong>EU:</strong> The supervisory authority responsible for you per <a href="https://edpb.europa.eu/about-edpb/about-edpb/members_en" target="_blank" rel="noopener" style="color:#7c3aed;">EDPB member list</a></li>
</ul>
</div>
<div class="ds-section" id="s11">
<div class="ds-section__heading"><span class="ds-num">11</span> Cookies and Tracking</div>
<ul class="ds-list">
<li>We use only technically necessary session cookies</li>
<li>For website analytics we use <strong>Umami Analytics</strong> – a privacy-friendly tool without cookies, without personal data collection, and without cross-site tracking</li>
<li>No Google Analytics, no advertising cookies, no sharing with third parties</li>
</ul>
<p>Legal basis: Art. 6(1)(f) GDPR (legitimate interest); Art. 31(1) FADP</p>
</div>
<div class="ds-section" id="s12">
<div class="ds-section__heading"><span class="ds-num">12</span> Automated Decision-Making</div>
<p>AI-powered transcription and protocol generation does not constitute automated decision-making within the meaning of Art. 22 GDPR, as no legal or similarly significant effects arise. The results serve as an aid and always require review by the user.</p>
</div>
<div class="ds-section" id="s13">
<div class="ds-section__heading"><span class="ds-num">13</span> Data Security</div>
<p>We implement state-of-the-art technical and organizational measures to protect your data, including:</p>
<ul class="ds-list">
<li>Encryption of all data transfers (TLS 1.2+)</li>
<li>Access control and authentication</li>
<li>Regular security updates</li>
<li>Immediate deletion of audio files after processing</li>
<li>No permanent server-side storage of transcripts</li>
</ul>
</div>
<div class="ds-section" id="s14">
<div class="ds-section__heading"><span class="ds-num">14</span> Obligation to Provide Data</div>
<p>Providing registration data (company name, email) is required for contract conclusion. Without this data, we cannot provide the service. Providing audio files is voluntary within the scope of respective use.</p>
</div>
<div class="ds-section" id="s15">
<div class="ds-section__heading"><span class="ds-num">15</span> Changes</div>
<p>We reserve the right to update this privacy policy at any time, in particular due to changes in legislation, new features, or new sub-processors.</p>
<p>The current version is always available. For material changes, we will notify registered users by email.</p>
</div>
</div>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    if st.button(f"✕ {close_text.get(lang, 'Schliessen')}", use_container_width=True, key="close_bottom_privacy"):
        st.query_params.clear()
        st.rerun()


@st.dialog("Privacy Policy / Datenschutz", width="large")
def show_privacy_dialog():
    """Zeigt den umfassenden Datenschutz-Dialog (DSGVO-konform wie Brevo)."""
    lang = st.session_state.get("language", "en")

    if lang == "de":
        st.markdown("""
### Datenschutzerklärung – MINU Minutes AI
*Stand: Februar 2025*

#### 1. Verantwortliche Stelle
**MINU-AI**
Schweiz
E-Mail: info@minu-ai.ch | Web: minu-ai.ch

#### 2. Welche Daten wir erheben
**Registrierungsdaten:** Firmenname, E-Mail-Adresse, Registrierungszeitpunkt
**Nutzungsdaten:** Hochgeladene Audio-Dateien (nur temporär), generierte Protokolle
**Technische Daten:** IP-Adresse, Browser-Typ, Geräteinformationen, Zugriffszeiten

#### 3. Zweck der Datenverarbeitung
- **Diensterbringung:** Transkription Ihrer Audio-Dateien und Protokollerstellung
- **Kontoverwaltung:** Bereitstellung des Zugangs zur Beta-Version
- **Kommunikation:** Benachrichtigungen über den Service, technische Updates
- **Marketing:** Informationen über neue Funktionen, Updates und Angebote (nur mit Ihrer Einwilligung)
- **Verbesserung:** Analyse zur Optimierung unseres Dienstes

#### 4. Rechtsgrundlage (DSGVO Art. 6)
- **Vertragserfüllung (Art. 6 Abs. 1 lit. b):** Für die Erbringung unseres Dienstes
- **Einwilligung (Art. 6 Abs. 1 lit. a):** Für Marketing-Kommunikation
- **Berechtigtes Interesse (Art. 6 Abs. 1 lit. f):** Für Sicherheit und technische Logs

#### 5. Datenübermittlung an Dritte

**Mistral AI (Frankreich/EU) 🇪🇺**
Für die Protokollerstellung nutzen wir Mistral AI, ein französisches Unternehmen mit Servern in der EU. Ihre Textdaten verbleiben innerhalb der Europäischen Union.
→ [Datenschutzrichtlinie von Mistral AI](https://mistral.ai/terms/#privacy-policy)

**OpenAI, LLC (USA) – nur für Transkription**
Für die Sprache-zu-Text-Umwandlung nutzen wir OpenAI Whisper. Ihre Audio-Daten werden zur Transkription an OpenAI übertragen.
→ [Datenschutzrichtlinie von OpenAI](https://openai.com/privacy)

Die Übermittlung in die USA erfolgt auf Basis von **Standardvertragsklauseln (SCCs)** gemäss Art. 46 Abs. 2 lit. c DSGVO.

**Hosting-Provider**
Unsere Server werden bei europäischen Anbietern mit Standort in der EU/Schweiz betrieben.

#### 6. Speicherdauer
- **Audio-Dateien:** Sofortige Löschung nach Verarbeitung (max. 1 Stunde)
- **Generierte Protokolle:** Verbleiben nur auf Ihrem Gerät
- **Registrierungsdaten:** Bis zur Löschung Ihres Kontos oder Widerruf
- **Technische Logs:** Maximal 30 Tage

#### 7. Ihre Rechte nach DSGVO
Sie haben jederzeit folgende Rechte:
- **Auskunft (Art. 15):** Welche Daten wir über Sie speichern
- **Berichtigung (Art. 16):** Korrektur unrichtiger Daten
- **Löschung (Art. 17):** "Recht auf Vergessenwerden"
- **Einschränkung (Art. 18):** Einschränkung der Verarbeitung
- **Datenübertragbarkeit (Art. 20):** Export Ihrer Daten
- **Widerspruch (Art. 21):** Gegen bestimmte Verarbeitungen
- **Widerruf der Einwilligung:** Jederzeit für Marketing-E-Mails

**Kontakt:** info@minu-ai.ch

#### 8. Marketing-Einwilligung
Mit der Registrierung und Aktivierung der Checkbox erklären Sie sich ausdrücklich einverstanden, dass wir Ihnen E-Mails mit Informationen über MINU Minutes AI senden dürfen. **Abmeldung jederzeit möglich** durch:
- Klick auf "Abmelden" in jeder E-Mail
- E-Mail an info@minu-ai.ch

#### 9. Datensicherheit
Wir setzen technische und organisatorische Massnahmen ein:
- Verschlüsselte Übertragung (HTTPS/TLS)
- Zugriffskontrolle und Authentifizierung
- Regelmässige Sicherheitsupdates
- Keine dauerhafte Speicherung von Audio-Daten

#### 10. Beschwerderecht
Sie haben das Recht, eine Beschwerde bei der zuständigen Datenschutzbehörde einzureichen (z.B. EDÖB in der Schweiz, oder Ihre lokale Aufsichtsbehörde).

#### 11. Änderungen
Wir können diese Datenschutzerklärung jederzeit aktualisieren. Bei wesentlichen Änderungen informieren wir Sie per E-Mail.
        """)
    else:  # English (default)
        st.markdown("""
### Privacy Policy – MINU Minutes AI
*Last updated: February 2025*

#### 1. Data Controller
**MINU-AI**
Switzerland
Email: info@minu-ai.ch | Web: minu-ai.ch

#### 2. Data We Collect
**Registration Data:** Company name, email address, registration timestamp
**Usage Data:** Uploaded audio files (temporary only), generated protocols
**Technical Data:** IP address, browser type, device information, access times

#### 3. Purpose of Processing
- **Service Delivery:** Transcription of your audio files and protocol generation
- **Account Management:** Providing access to the beta version
- **Communication:** Service notifications, technical updates
- **Marketing:** Information about new features, updates and offers (only with your consent)
- **Improvement:** Analytics to optimize our service

#### 4. Legal Basis (GDPR Art. 6)
- **Contract Performance (Art. 6.1.b):** For providing our service
- **Consent (Art. 6.1.a):** For marketing communications
- **Legitimate Interest (Art. 6.1.f):** For security and technical logs

#### 5. Data Transfer to Third Parties

**Mistral AI (France/EU) 🇪🇺**
For protocol generation, we use Mistral AI, a French company with servers in the EU. Your text data remains within the European Union.
→ [Mistral AI Privacy Policy](https://mistral.ai/terms/#privacy-policy)

**OpenAI, LLC (USA) – Transcription only**
For speech-to-text conversion, we use OpenAI Whisper. Your audio data is transferred to OpenAI for transcription.
→ [OpenAI Privacy Policy](https://openai.com/privacy)

Transfer to the USA is based on **Standard Contractual Clauses (SCCs)** pursuant to Art. 46.2.c GDPR.

**Hosting Provider**
Our servers are operated by European providers located in the EU/Switzerland.

#### 6. Data Retention
- **Audio Files:** Immediate deletion after processing (max. 1 hour)
- **Generated Protocols:** Remain only on your device
- **Registration Data:** Until account deletion or withdrawal
- **Technical Logs:** Maximum 30 days

#### 7. Your Rights under GDPR
You have the following rights at any time:
- **Access (Art. 15):** What data we store about you
- **Rectification (Art. 16):** Correction of inaccurate data
- **Erasure (Art. 17):** "Right to be forgotten"
- **Restriction (Art. 18):** Restriction of processing
- **Portability (Art. 20):** Export your data
- **Object (Art. 21):** Against certain processing
- **Withdraw Consent:** For marketing emails at any time

**Contact:** info@minu-ai.ch

#### 8. Marketing Consent
By registering and checking the consent box, you expressly agree that we may send you emails with information about MINU Minutes AI and MINU. **Unsubscribe anytime** by:
- Clicking "Unsubscribe" in any email
- Emailing info@minu-ai.ch

#### 9. Data Security
We implement technical and organizational measures:
- Encrypted transmission (HTTPS/TLS)
- Access control and authentication
- Regular security updates
- No permanent storage of audio data

#### 10. Right to Complain
You have the right to lodge a complaint with the competent data protection authority (e.g., FDPIC in Switzerland, or your local supervisory authority).

#### 11. Changes
We may update this privacy policy at any time. We will notify you of significant changes by email.
        """)

    if st.button("OK", use_container_width=True, type="primary"):
        st.query_params.clear()
        st.rerun()


@st.dialog("Impressum", width="small")
def show_impressum_dialog():
    """Zeigt das Impressum kompakt."""
    st.markdown("""
<div style="font-size: 14px; line-height: 1.6;">
    <div style="text-align: center; margin-bottom: 16px;">
        <span style="font-size: 24px; font-weight: 700; background: linear-gradient(90deg, #7c3aed, #2dd4bf); -webkit-background-clip: text; -webkit-text-fill-color: transparent;">MINU-AI</span>
    </div>
    <div style="background: #f8f8f8; border-radius: 8px; padding: 12px; margin-bottom: 12px;">
        <div style="font-weight: 600; color: #333;">SPEKTRUM Partner GmbH</div>
        <div style="color: #666;">Josefstrasse 181, 8005 Zürich, Schweiz</div>
    </div>
    <table style="width: 100%; font-size: 13px; color: #555;">
        <tr><td style="padding: 4px 0; color: #888;">E-Mail</td><td style="padding: 4px 0;">minutes-ai@spekt.ch</td></tr>
        <tr><td style="padding: 4px 0; color: #888;">Web</td><td style="padding: 4px 0;">minu-ai.ch</td></tr>
        <tr><td style="padding: 4px 0; color: #888;">UID</td><td style="padding: 4px 0;">CH-020.4.068.247-6</td></tr>
    </table>
</div>
    """, unsafe_allow_html=True)

    if st.button("OK", use_container_width=True, type="primary"):
        st.query_params.clear()
        st.rerun()


@st.dialog("Datenschutz", width="large")
def show_datenschutz_dialog():
    """Zeigt die Datenschutzerklärung als eigenen Dialog."""
    lang = st.session_state.get("language", "en")

    if lang == "de":
        st.markdown("""
<div style="font-size: 12px; line-height: 1.6; max-height: 500px; overflow-y: auto;">
<p style="color: #888; font-size: 11px; margin-bottom: 12px;">Datenschutzerklärung – Stand: Februar 2026</p>

<b>1. Verantwortlicher</b><br>
SPEKTRUM Partner GmbH<br>
Josefstrasse 181, 8005 Zürich, Schweiz<br>
E-Mail: minutes-ai@spekt.ch<br>
Website: https://minu-ai.ch

<b>2. Datenschutzbeauftragter</b><br>
Ein Datenschutzbeauftragter ist nach Art. 37 DSGVO nicht bestellt, da die Voraussetzungen nicht erfüllt sind. Bei Fragen zum Datenschutz wenden Sie sich bitte an: minutes-ai@spekt.ch

<b>3. Erhobene Daten und Zwecke</b><br>
Wir verarbeiten folgende personenbezogene Daten:<br><br>
<u>Registrierungsdaten:</u><br>
• Firmenname, E-Mail-Adresse<br>
• Zweck: Vertragserfüllung, Kontozugang, Kommunikation<br>
• Rechtsgrundlage: Art. 6 Abs. 1 lit. b DSGVO (Vertragserfüllung)<br><br>
<u>Audiodaten:</u><br>
• Hochgeladene Audio-Dateien zur Transkription<br>
• Zweck: Erbringung der Dienstleistung (Transkription und Protokollerstellung)<br>
• Rechtsgrundlage: Art. 6 Abs. 1 lit. b DSGVO (Vertragserfüllung)<br><br>
<u>Technische Daten:</u><br>
• IP-Adresse, Browser-Typ, Betriebssystem, Zugriffszeiten<br>
• Zweck: Sicherheit, Fehleranalyse, Missbrauchsprävention<br>
• Rechtsgrundlage: Art. 6 Abs. 1 lit. f DSGVO (berechtigtes Interesse)<br><br>
<u>Zahlungsdaten:</u><br>
• Kreditkartendaten werden ausschliesslich von Stripe verarbeitet<br>
• Wir speichern keine vollständigen Zahlungsdaten<br>
• Rechtsgrundlage: Art. 6 Abs. 1 lit. b DSGVO (Vertragserfüllung)

<b>4. Datenübermittlung an Dritte</b><br>
Zur Erbringung unserer Dienste arbeiten wir mit folgenden Auftragsverarbeitern:<br><br>
<u>🇪🇺 Mistral AI (Frankreich/EU)</u><br>
• Zweck: KI-gestützte Protokollerstellung<br>
• Datenverarbeitung innerhalb der EU<br>
• Auftragsverarbeitungsvertrag (AVV) abgeschlossen<br><br>
<u>🇺🇸 OpenAI, Inc. (USA)</u><br>
• Zweck: Transkription mittels Whisper API<br>
• Drittlandübermittlung auf Basis von EU-Standardvertragsklauseln (SCC) gemäss Art. 46 Abs. 2 lit. c DSGVO<br>
• Auftragsverarbeitungsvertrag (AVV/DPA) abgeschlossen<br>
• OpenAI speichert keine Daten zur Modellverbesserung (API-Nutzung)<br><br>
<u>🇺🇸 Stripe, Inc. (USA)</u><br>
• Zweck: Sichere Zahlungsabwicklung<br>
• Drittlandübermittlung auf Basis von EU-Standardvertragsklauseln (SCC)<br>
• PCI-DSS-zertifiziert<br><br>
<u>🇨🇭 Hosting-Provider (Schweiz/EU)</u><br>
• Serverstandort: Schweiz oder EU<br>
• Auftragsverarbeitungsvertrag (AVV) abgeschlossen

<b>5. Speicherfristen</b><br>
• <b>Audiodateien:</b> Sofort nach Verarbeitung unwiderruflich gelöscht (keine Speicherung)<br>
• <b>Protokolle:</b> Werden nur lokal beim Nutzer gespeichert, nicht auf unseren Servern<br>
• <b>Registrierungsdaten:</b> Bis zur Kontolöschung oder 3 Jahre nach letzter Aktivität<br>
• <b>Server-Logs:</b> Maximal 30 Tage<br>
• <b>Rechnungsdaten:</b> 10 Jahre (gesetzliche Aufbewahrungspflicht)

<b>6. Ihre Rechte nach DSGVO</b><br>
Sie haben folgende Rechte bezüglich Ihrer personenbezogenen Daten:<br><br>
• <b>Auskunftsrecht (Art. 15):</b> Welche Daten wir über Sie speichern<br>
• <b>Berichtigungsrecht (Art. 16):</b> Korrektur unrichtiger Daten<br>
• <b>Löschungsrecht (Art. 17):</b> Löschung Ihrer Daten ("Recht auf Vergessenwerden")<br>
• <b>Einschränkung (Art. 18):</b> Einschränkung der Verarbeitung<br>
• <b>Datenübertragbarkeit (Art. 20):</b> Export Ihrer Daten in maschinenlesbarem Format<br>
• <b>Widerspruchsrecht (Art. 21):</b> Widerspruch gegen Verarbeitung auf Basis berechtigter Interessen<br><br>
Zur Ausübung Ihrer Rechte kontaktieren Sie uns unter: minutes-ai@spekt.ch

<b>7. Beschwerderecht bei der Aufsichtsbehörde</b><br>
Sie haben das Recht, sich bei einer Datenschutz-Aufsichtsbehörde zu beschweren:<br><br>
<u>Schweiz:</u><br>
Eidgenössischer Datenschutz- und Öffentlichkeitsbeauftragter (EDÖB)<br>
Feldeggweg 1, 3003 Bern<br>
www.edoeb.admin.ch<br><br>
<u>EU (je nach Wohnsitz):</u><br>
Liste der EU-Aufsichtsbehörden: https://edpb.europa.eu/about-edpb/about-edpb/members_de

<b>8. Cookies und Tracking</b><br>
• Wir verwenden ausschliesslich technisch notwendige Session-Cookies<br>
• Für die Webseitenanalyse nutzen wir Umami Analytics – ein datenschutzfreundliches Tool ohne Cookies, ohne persönliche Datenerfassung und ohne Tracking über Webseiten hinweg<br>
• Kein Google Analytics, keine Werbe-Cookies, keine Weitergabe an Dritte<br>
• Rechtsgrundlage: Art. 6 Abs. 1 lit. f DSGVO (berechtigtes Interesse an der Verbesserung unseres Dienstes)

<b>9. Automatisierte Entscheidungsfindung</b><br>
Die KI-gestützte Transkription und Protokollerstellung stellt keine automatisierte Entscheidungsfindung im Sinne von Art. 22 DSGVO dar, da:<br>
• keine rechtlichen oder ähnlich erheblichen Wirkungen für Sie entstehen<br>
• alle Ergebnisse von Ihnen überprüft und bearbeitet werden können<br>
• die Verarbeitung der Vertragserfüllung dient

<b>10. Datensicherheit</b><br>
Wir setzen technische und organisatorische Massnahmen ein:<br>
• TLS-Verschlüsselung aller Datenübertragungen (HTTPS)<br>
• Verschlüsselte Speicherung sensibler Daten<br>
• Zugriffsbeschränkungen und Authentifizierung<br>
• Regelmässige Sicherheitsupdates

<b>11. Pflicht zur Bereitstellung von Daten</b><br>
Die Bereitstellung Ihrer E-Mail-Adresse und Firmendaten ist für den Vertragsschluss erforderlich. Ohne diese Daten können wir den Dienst nicht erbringen. Die Bereitstellung von Audiodaten erfolgt freiwillig zur Nutzung der Transkriptionsfunktion.

<b>12. Änderungen dieser Datenschutzerklärung</b><br>
Wir behalten uns vor, diese Datenschutzerklärung bei Bedarf anzupassen. Die aktuelle Version ist stets auf unserer Website verfügbar. Bei wesentlichen Änderungen informieren wir Sie per E-Mail.
</div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
<div style="font-size: 12px; line-height: 1.6; max-height: 500px; overflow-y: auto;">
<p style="color: #888; font-size: 11px; margin-bottom: 12px;">Privacy Policy – Last updated: February 2026</p>

<b>1. Data Controller</b><br>
SPEKTRUM Partner GmbH<br>
Josefstrasse 181, 8005 Zurich, Switzerland<br>
Email: minutes-ai@spekt.ch<br>
Website: https://minu-ai.ch

<b>2. Data Protection Officer</b><br>
A Data Protection Officer is not appointed pursuant to Art. 37 GDPR as the requirements are not met. For data protection inquiries, please contact: minutes-ai@spekt.ch

<b>3. Data Collected and Purposes</b><br>
We process the following personal data:<br><br>
<u>Registration Data:</u><br>
• Company name, email address<br>
• Purpose: Contract performance, account access, communication<br>
• Legal basis: Art. 6(1)(b) GDPR (contract performance)<br><br>
<u>Audio Data:</u><br>
• Uploaded audio files for transcription<br>
• Purpose: Service delivery (transcription and protocol generation)<br>
• Legal basis: Art. 6(1)(b) GDPR (contract performance)<br><br>
<u>Technical Data:</u><br>
• IP address, browser type, operating system, access times<br>
• Purpose: Security, error analysis, abuse prevention<br>
• Legal basis: Art. 6(1)(f) GDPR (legitimate interest)<br><br>
<u>Payment Data:</u><br>
• Credit card data is processed exclusively by Stripe<br>
• We do not store complete payment data<br>
• Legal basis: Art. 6(1)(b) GDPR (contract performance)

<b>4. Data Transfer to Third Parties</b><br>
We work with the following processors to provide our services:<br><br>
<u>🇪🇺 Mistral AI (France/EU)</u><br>
• Purpose: AI-powered protocol generation<br>
• Data processing within the EU<br>
• Data Processing Agreement (DPA) in place<br><br>
<u>🇺🇸 OpenAI, Inc. (USA)</u><br>
• Purpose: Transcription via Whisper API<br>
• Third-country transfer based on EU Standard Contractual Clauses (SCC) pursuant to Art. 46(2)(c) GDPR<br>
• Data Processing Agreement (DPA) in place<br>
• OpenAI does not store data for model improvement (API usage)<br><br>
<u>🇺🇸 Stripe, Inc. (USA)</u><br>
• Purpose: Secure payment processing<br>
• Third-country transfer based on EU Standard Contractual Clauses (SCC)<br>
• PCI-DSS certified<br><br>
<u>🇨🇭 Hosting Provider (Switzerland/EU)</u><br>
• Server location: Switzerland or EU<br>
• Data Processing Agreement (DPA) in place

<b>5. Retention Periods</b><br>
• <b>Audio files:</b> Immediately and irreversibly deleted after processing (no storage)<br>
• <b>Protocols:</b> Stored only locally with the user, not on our servers<br>
• <b>Registration data:</b> Until account deletion or 3 years after last activity<br>
• <b>Server logs:</b> Maximum 30 days<br>
• <b>Billing data:</b> 10 years (statutory retention requirement)

<b>6. Your Rights under GDPR</b><br>
You have the following rights regarding your personal data:<br><br>
• <b>Right of access (Art. 15):</b> Information about what data we store about you<br>
• <b>Right to rectification (Art. 16):</b> Correction of inaccurate data<br>
• <b>Right to erasure (Art. 17):</b> Deletion of your data ("right to be forgotten")<br>
• <b>Right to restriction (Art. 18):</b> Restriction of processing<br>
• <b>Right to data portability (Art. 20):</b> Export of your data in machine-readable format<br>
• <b>Right to object (Art. 21):</b> Object to processing based on legitimate interests<br><br>
To exercise your rights, contact us at: minutes-ai@spekt.ch

<b>7. Right to Lodge a Complaint</b><br>
You have the right to lodge a complaint with a data protection supervisory authority:<br><br>
<u>Switzerland:</u><br>
Federal Data Protection and Information Commissioner (FDPIC)<br>
Feldeggweg 1, 3003 Bern<br>
www.edoeb.admin.ch<br><br>
<u>EU (depending on residence):</u><br>
List of EU supervisory authorities: https://edpb.europa.eu/about-edpb/about-edpb/members_en

<b>8. Cookies and Tracking</b><br>
• We use only technically necessary session cookies<br>
• For website analytics we use Umami Analytics – a privacy-friendly tool without cookies, without personal data collection, and without cross-site tracking<br>
• No Google Analytics, no advertising cookies, no sharing with third parties<br>
• Legal basis: Art. 6(1)(f) GDPR (legitimate interest in improving our service)

<b>9. Automated Decision-Making</b><br>
AI-powered transcription and protocol generation does not constitute automated decision-making within the meaning of Art. 22 GDPR because:<br>
• No legal or similarly significant effects arise for you<br>
• All results can be reviewed and edited by you<br>
• Processing serves contract performance

<b>10. Data Security</b><br>
We implement technical and organizational measures:<br>
• TLS encryption for all data transfers (HTTPS)<br>
• Encrypted storage of sensitive data<br>
• Access restrictions and authentication<br>
• Regular security updates

<b>11. Obligation to Provide Data</b><br>
Providing your email address and company data is required for contract conclusion. Without this data, we cannot provide the service. Providing audio data is voluntary for use of the transcription function.

<b>12. Changes to This Privacy Policy</b><br>
We reserve the right to update this privacy policy as needed. The current version is always available on our website. We will notify you by email of material changes.
</div>
        """, unsafe_allow_html=True)

    if st.button("OK", use_container_width=True, type="primary"):
        st.query_params.clear()
        st.rerun()


@st.dialog("AGB", width="large")
def show_agb_dialog():
    """Zeigt die AGB."""
    lang = st.session_state.get("language", "en")

    if lang == "de":
        st.markdown("""
<div style="font-size: 12px; line-height: 1.6; max-height: 500px; overflow-y: auto;">
<p style="color: #888; font-size: 11px; margin-bottom: 12px;">Allgemeine Geschäftsbedingungen – Stand: Februar 2026</p>

<b>§1 Geltungsbereich und B2B-Nutzung</b><br>
1.1 Diese Allgemeinen Geschäftsbedingungen (AGB) gelten für alle Verträge zwischen der SPEKTRUM Partner GmbH, Josefstrasse 181, 8005 Zürich, Schweiz (nachfolgend "Anbieter") und dem Nutzer (nachfolgend "Kunde") über die Nutzung der Software-as-a-Service-Lösung MINU-AI.<br><br>
1.2 <b>MINU-AI richtet sich ausschliesslich an Unternehmer, Gewerbetreibende, Freiberufler und juristische Personen (B2B).</b> Mit der Registrierung bestätigt der Kunde ausdrücklich, dass er in Ausübung seiner gewerblichen, selbstständigen oder freiberuflichen Tätigkeit handelt und nicht als Verbraucher im Sinne des Konsumentenschutzrechts.<br><br>
1.3 Verbraucherschutzrechte, insbesondere das gesetzliche Widerrufsrecht für Fernabsatzverträge, finden auf Verträge mit MINU-AI keine Anwendung.<br><br>
1.4 Der Kunde garantiert die Richtigkeit seiner Angaben zur Unternehmereigenschaft. Bei falschen Angaben haftet der Kunde für alle daraus entstehenden Schäden.

<b>§2 Leistungsbeschreibung</b><br>
2.1 MINU-AI ist eine cloudbasierte Anwendung zur automatischen Transkription von Audioaufnahmen und KI-gestützten Erstellung von Sitzungsprotokollen.<br><br>
2.2 Die Leistung umfasst: (a) Transkription mittels OpenAI Whisper API, (b) Protokollerstellung mittels Mistral AI, (c) Export in PDF- und Word-Format, (d) Webbasierter Zugang zur Anwendung.<br><br>
2.3 Der Anbieter kann den Funktionsumfang jederzeit erweitern, anpassen oder einschränken, sofern dies für den Kunden zumutbar ist.

<b>§3 Vertragsschluss und Unternehmereigenschaft</b><br>
3.1 Der Vertrag kommt mit Abschluss der Registrierung und Bestätigung dieser AGB zustande.<br><br>
3.2 Das Mindestalter für die Nutzung beträgt 18 Jahre.<br><br>
3.3 Der Kunde garantiert und sichert zu, dass er als Unternehmer im Sinne von Art. 14 BGB / OR handelt. Diese Zusicherung ist wesentliche Vertragsgrundlage.

<b>§4 Preise, Zahlung und Steuern</b><br>
4.1 Es gelten die zum Zeitpunkt der Bestellung auf der Website angegebenen Preise in Schweizer Franken (CHF). Alle Preise verstehen sich als Nettopreise.<br><br>
4.2 Die Zahlung erfolgt im Voraus per Kreditkarte über den Zahlungsdienstleister Stripe.<br><br>
4.3 Abonnements verlängern sich automatisch um die jeweilige Vertragslaufzeit, sofern nicht fristgerecht gekündigt wird.<br><br>
4.4 <b>Der Kunde ist für alle anfallenden Steuern, Abgaben und Zölle in seinem Land allein verantwortlich.</b> Für EU-Geschäftskunden mit gültiger UID gilt das Reverse-Charge-Verfahren. Der Anbieter übernimmt keine Verantwortung für die korrekte steuerliche Behandlung im Land des Kunden.<br><br>
4.5 <b>Wertsicherung:</b> Die Entgelte sind wertgesichert. Als Massstab dient der vom Bundesamt für Statistik (BFS) veröffentlichte Landesindex der Konsumentenpreise (LIK). Basismonat ist der Monat des Vertragsschlusses. Der Anbieter ist berechtigt, zu Beginn eines jeden Kalenderjahres die Preise im Umfang der Indexveränderung anzupassen. Der Kunde wird über Preisanpassungen per E-Mail informiert. Unterlässt der Anbieter die Indexierung in einzelnen Jahren, bedeutet dies keinen Verzicht auf die eingetretene Wertsicherung.

<b>§5 Geografische Verfügbarkeit</b><br>
5.1 <b>MINU-AI ist ausschliesslich für Kunden mit Sitz in folgenden Regionen verfügbar:</b> Schweiz, Europäische Union (EU), Europäischer Wirtschaftsraum (EWR) und Vereinigtes Königreich (UK).<br><br>
5.2 <b>Kunden aus folgenden Ländern sind von der Nutzung ausgeschlossen:</b> USA, China, Russland, Indien, Brasilien sowie alle Länder, die Sanktionen unterliegen.<br><br>
5.3 Mit der Registrierung bestätigt der Kunde, dass sein Unternehmen seinen Sitz in einer der zugelassenen Regionen hat.<br><br>
5.4 Der Anbieter behält sich vor, Registrierungen aus ausgeschlossenen Ländern abzulehnen oder bestehende Konten zu kündigen.

<b>§6 Kostenlose Testphase</b><br>
6.1 Neukunden erhalten eine kostenlose Testphase von 21 Tagen mit maximal 60 Transkriptionsminuten.<br><br>
6.2 Die Testphase endet automatisch ohne weitere Verpflichtungen.<br><br>
6.3 Auch während der Testphase besteht kein Widerrufsrecht, da der Dienst sofort zur Verfügung gestellt wird.

<b>§7 Laufzeit und Kündigung</b><br>
7.1 Monatsabonnements können jederzeit zum Ende des laufenden Abrechnungszeitraums gekündigt werden.<br><br>
7.2 Jahresabonnements können jederzeit zum Ende der laufenden Vertragslaufzeit gekündigt werden.<br><br>
7.3 Bereits gezahlte Beträge werden bei Kündigung nicht erstattet (keine anteilige Rückerstattung).

<b>§8 Widerrufsrecht – B2B-Ausschluss</b><br>
8.1 Da MINU-AI ausschliesslich für Geschäftskunden bestimmt ist, besteht kein gesetzliches Widerrufsrecht.<br><br>
8.2 Das Verbraucher-Widerrufsrecht gemäss EU-Richtlinie 2011/83/EU und entsprechenden nationalen Regelungen findet keine Anwendung.<br><br>
8.3 Vertragsbeendigungen sind ausschliesslich nach den Kündigungsregelungen in §7 möglich.

<b>§9 Verfügbarkeit und Ausfälle – Kein SLA</b><br>
9.1 <b>Der Anbieter schuldet keine bestimmte Verfügbarkeit der Dienste.</b> MINU-AI wird "as is" und "as available" bereitgestellt.<br><br>
9.2 Der Anbieter ist berechtigt, den Dienst jederzeit für Wartungsarbeiten, Updates oder aus anderen Gründen vorübergehend oder dauerhaft einzuschränken oder einzustellen.<br><br>
9.3 <b>Es besteht kein Anspruch auf ununterbrochene Verfügbarkeit.</b> Der Anbieter haftet nicht für Ausfälle, Unterbrechungen, Verzögerungen oder Störungen, gleich aus welchem Grund.<br><br>
9.4 Bei dauerhafter Einstellung des Dienstes wird der Kunde informiert. Bereits gezahlte Beträge für nicht genutzte Zeiträume werden nicht erstattet.

<b>§10 Support – Kein Anspruch</b><br>
10.1 <b>Der Anbieter bietet keinen garantierten Kundensupport.</b> Support-Anfragen können per E-Mail gestellt werden, es besteht jedoch kein Anspruch auf Beantwortung oder Reaktionszeiten.<br><br>
10.2 Der Anbieter kann nach eigenem Ermessen Support-Leistungen anbieten, einschränken oder einstellen.<br><br>
10.3 Telefonsupport, Schulungen oder Vor-Ort-Service sind nicht Bestandteil des Vertrages.

<b>§11 Haftungsausschluss und KI-Disclaimer</b><br>
11.1 <b>KI-generierte Inhalte sind ausschliesslich als Arbeitshilfen zu verstehen.</b> Sie stellen keine rechtsverbindlichen, vollständigen oder fehlerfreien Dokumente dar.<br><br>
11.2 <b>Der Kunde ist als Unternehmer verpflichtet, alle KI-generierten Ergebnisse eigenverantwortlich zu prüfen, zu verifizieren und ggf. zu korrigieren.</b> Der Anbieter übernimmt keine Gewähr für Richtigkeit, Vollständigkeit oder Eignung für einen bestimmten Zweck.<br><br>
11.3 <b>Die Haftung für KI-Fehler, Halluzinationen, Fehlinterpretationen, falsche Transkriptionen oder unvollständige Protokolle ist vollständig ausgeschlossen.</b><br><br>
11.4 <b>Die Haftung für Ausfälle, Störungen oder Nichtverfügbarkeit von Drittdiensten (OpenAI, Mistral AI, Stripe, Hosting-Provider) ist vollständig ausgeschlossen.</b><br><br>
11.5 Die Gesamthaftung des Anbieters ist in jedem Fall auf die Höhe der vom Kunden in den letzten 12 Monaten gezahlten Entgelte begrenzt, maximal jedoch auf den Betrag einer Monatszahlung.<br><br>
11.6 Die vorstehenden Haftungsbeschränkungen gelten nicht für Schäden aus der Verletzung des Lebens, des Körpers oder der Gesundheit sowie für Vorsatz.<br><br>
11.7 <b>Verkürzte Verjährung:</b> Schadenersatzansprüche des Kunden gegen den Anbieter verjähren innert 6 Monaten ab Kenntnis des Schadens und des Schädigers, spätestens jedoch innert 3 Jahren ab dem schadensverursachenden Ereignis. Diese Frist gilt nicht für Ansprüche aus Personenschäden oder Vorsatz.

<b>§12 Pflichten und Verantwortung des Kunden</b><br>
12.1 Der Kunde ist allein verantwortlich für die rechtmässige Nutzung von MINU-AI.<br><br>
12.2 Der Kunde stellt sicher, dass er alle erforderlichen Einwilligungen der aufgenommenen Personen eingeholt hat.<br><br>
12.3 Der Kunde nutzt MINU-AI nicht für rechtswidrige, belästigende, verleumderische oder anderweitig unzulässige Zwecke.<br><br>
12.4 Der Kunde stellt den Anbieter von allen Ansprüchen Dritter frei, die aus einer rechtswidrigen Nutzung resultieren.<br><br>
12.5 <b>Verbot von Reverse Engineering:</b> Dem Kunden ist es untersagt, die Software zu dekompilieren, zu disassemblieren, zurückzuentwickeln (Reverse Engineering) oder abgeleitete Werke zu erstellen. Bei Zuwiderhandlung wird eine Konventionalstrafe von CHF 10'000.– pro Verstoss fällig, ohne dass der Anbieter einen Schaden nachweisen muss. Die Geltendmachung eines darüber hinausgehenden Schadens bleibt vorbehalten.

<b>§13 Datenschutz</b><br>
Es gilt die separate Datenschutzerklärung (siehe Link "Datenschutz" im Footer).

<b>§14 Änderungsvorbehalt</b><br>
14.1 Der Anbieter behält sich vor, diese AGB jederzeit zu ändern.<br><br>
14.2 Änderungen werden dem Kunden per E-Mail mitgeteilt. Widerspricht der Kunde nicht innerhalb von 30 Tagen, gelten die geänderten AGB als akzeptiert.<br><br>
14.3 Bei wesentlichen Änderungen steht dem Kunden ein Sonderkündigungsrecht zu.

<b>§15 Schlussbestimmungen</b><br>
15.1 Es gilt ausschliesslich Schweizer Recht unter Ausschluss des UN-Kaufrechts.<br><br>
15.2 Ausschliesslicher Gerichtsstand ist Zürich, Schweiz.<br><br>
15.3 Sollten einzelne Bestimmungen unwirksam sein, bleibt die Wirksamkeit der übrigen Bestimmungen unberührt.<br><br>
15.4 Mündliche Nebenabreden bestehen nicht.
</div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
<div style="font-size: 12px; line-height: 1.6; max-height: 500px; overflow-y: auto;">
<p style="color: #888; font-size: 11px; margin-bottom: 12px;">General Terms and Conditions – Last updated: February 2026</p>

<b>§1 Scope and B2B Use</b><br>
1.1 These General Terms and Conditions (GTC) apply to all contracts between SPEKTRUM Partner GmbH, Josefstrasse 181, 8005 Zurich, Switzerland (hereinafter "Provider") and the user (hereinafter "Customer") regarding the use of the software-as-a-service solution MINU-AI.<br><br>
1.2 <b>MINU-AI is intended exclusively for entrepreneurs, business professionals, freelancers and legal entities (B2B).</b> By registering, the Customer expressly confirms that they are acting in the course of their commercial, self-employed or professional activity and not as a consumer.<br><br>
1.3 Consumer protection rights, in particular the statutory right of withdrawal for distance contracts, do not apply to contracts with MINU-AI.<br><br>
1.4 The Customer guarantees the accuracy of their information regarding their business status. The Customer is liable for any damages resulting from false information.

<b>§2 Service Description</b><br>
2.1 MINU-AI is a cloud-based application for automatic transcription of audio recordings and AI-assisted creation of meeting protocols.<br><br>
2.2 The service includes: (a) Transcription via OpenAI Whisper API, (b) Protocol generation via Mistral AI, (c) Export to PDF and Word format, (d) Web-based access to the application.<br><br>
2.3 The Provider may expand, modify or limit the scope of features at any time, provided this is reasonable for the Customer.

<b>§3 Contract Formation and Business Status</b><br>
3.1 The contract is concluded upon completion of registration and acceptance of these GTC.<br><br>
3.2 The minimum age for use is 18 years.<br><br>
3.3 The Customer warrants and represents that they are acting as a business entity. This representation is a material basis of the contract.

<b>§4 Prices, Payment and Taxes</b><br>
4.1 The prices displayed on the website at the time of order apply, in Swiss Francs (CHF).<br><br>
4.2 Payment is made in advance by credit card via the payment service provider Stripe.<br><br>
4.3 Subscriptions renew automatically for the respective contract period unless cancelled in due time.<br><br>
4.4 For EU business customers with a valid VAT identification number, the reverse charge procedure applies. The Customer is responsible for correct taxation in their country.<br><br>
4.5 <b>The Customer is solely responsible for the correct payment of all applicable taxes, duties and fees in their country of residence or place of business.</b> The Provider does not assume any tax obligations abroad and is not liable for tax consequences on the Customer's side.<br><br>
4.6 <b>Price Indexation:</b> The fees are subject to value adjustment. The Swiss Consumer Price Index (CPI) published by the Federal Statistical Office (FSO) serves as the benchmark. The base month is the month of contract conclusion. The Provider is entitled to adjust prices at the beginning of each calendar year in accordance with the index change. The Customer will be notified of price adjustments by email. If the Provider does not apply the indexation in individual years, this does not constitute a waiver of the accrued value adjustment.

<b>§5 Geographic Availability</b><br>
5.1 <b>MINU-AI is exclusively available to customers from the following regions:</b> European Union (EU), European Economic Area (EEA), Switzerland and the United Kingdom (UK).<br><br>
5.2 <b>The service is expressly NOT available for customers from the following countries:</b> United States of America (USA), China (including Hong Kong and Macao), Russia, India, Brazil, as well as all other countries not mentioned in §5.1.<br><br>
5.3 The Customer warrants that their place of business is located in one of the permitted regions according to §5.1. False information may lead to immediate contract termination.<br><br>
5.4 The Provider reserves the right to expand or restrict the list of permitted regions at any time. Existing contracts are not affected by changes to newly permitted regions.

<b>§6 Free Trial Period</b><br>
6.1 New customers receive a free trial period of 21 days with a maximum of 60 transcription minutes.<br><br>
6.2 The trial period ends automatically without further obligations.<br><br>
6.3 No right of withdrawal exists during the trial period, as the service is made available immediately.

<b>§7 Term and Cancellation</b><br>
7.1 Monthly subscriptions can be cancelled at any time effective at the end of the current billing period.<br><br>
7.2 Annual subscriptions can be cancelled at any time effective at the end of the current contract term.<br><br>
7.3 Amounts already paid will not be refunded upon cancellation (no pro-rata refunds).

<b>§8 Right of Withdrawal – B2B Exclusion</b><br>
8.1 As MINU-AI is intended exclusively for business customers, no statutory right of withdrawal exists.<br><br>
8.2 The consumer right of withdrawal under EU Directive 2011/83/EU and corresponding national regulations does not apply.<br><br>
8.3 Contract terminations are only possible according to the cancellation provisions in §7.

<b>§9 Availability and Outages – No SLA</b><br>
9.1 <b>The Provider does not guarantee any specific availability of the services.</b> MINU-AI is provided "as is" and "as available".<br><br>
9.2 The Provider is entitled to restrict or discontinue the service at any time for maintenance, updates or other reasons, temporarily or permanently.<br><br>
9.3 <b>There is no entitlement to uninterrupted availability.</b> The Provider is not liable for outages, interruptions, delays or disruptions, regardless of the reason.<br><br>
9.4 In case of permanent discontinuation of the service, the Customer will be informed. Amounts already paid for unused periods will not be refunded.

<b>§10 Support – No Entitlement</b><br>
10.1 <b>The Provider does not offer guaranteed customer support.</b> Support requests may be submitted by email, but there is no entitlement to response or response times.<br><br>
10.2 The Provider may offer, limit or discontinue support services at its own discretion.<br><br>
10.3 Phone support, training or on-site service are not part of the contract.

<b>§11 Disclaimer and AI Disclaimer</b><br>
11.1 <b>AI-generated content is intended solely as a work aid.</b> It does not constitute legally binding, complete or error-free documents.<br><br>
11.2 <b>As a business user, the Customer is obligated to independently review, verify and correct all AI-generated results.</b> The Provider makes no warranty for accuracy, completeness or fitness for a particular purpose.<br><br>
11.3 <b>Liability for AI errors, hallucinations, misinterpretations, incorrect transcriptions or incomplete protocols is completely excluded.</b><br><br>
11.4 <b>Liability for outages, disruptions or unavailability of third-party services (OpenAI, Mistral AI, Stripe, hosting providers) is completely excluded.</b><br><br>
11.5 The Provider's total liability is in any case limited to the amount of fees paid by the Customer in the last 12 months, but no more than one monthly payment.<br><br>
11.6 The above liability limitations do not apply to damages from injury to life, body or health, or for intentional misconduct.<br><br>
11.7 <b>Shortened Statute of Limitations:</b> Claims for damages by the Customer against the Provider shall become time-barred within 6 months of knowledge of the damage and the party causing the damage, but in any event no later than 3 years from the event causing the damage. This limitation period does not apply to claims arising from personal injury or intentional misconduct.

<b>§12 Customer Obligations and Responsibilities</b><br>
12.1 The Customer is solely responsible for the lawful use of MINU-AI.<br><br>
12.2 The Customer ensures that they have obtained all necessary consents from recorded persons.<br><br>
12.3 The Customer does not use MINU-AI for unlawful, harassing, defamatory or otherwise impermissible purposes.<br><br>
12.4 The Customer indemnifies the Provider against all third-party claims resulting from unlawful use.<br><br>
12.5 <b>Prohibition of Reverse Engineering:</b> The Customer is prohibited from decompiling, disassembling, reverse engineering or creating derivative works of the software. In case of violation, a contractual penalty of CHF 10,000 per infringement is due, without the Provider having to prove any damage. The right to claim damages exceeding this amount is reserved.

<b>§13 Data Protection</b><br>
The separate Privacy Policy applies (see "Privacy" link in the footer).

<b>§14 Amendment Clause</b><br>
14.1 The Provider reserves the right to amend these GTC at any time.<br><br>
14.2 Amendments will be communicated to the Customer by email. If the Customer does not object within 30 days, the amended GTC are deemed accepted.<br><br>
14.3 In case of material changes, the Customer has a special right of termination.

<b>§15 Final Provisions</b><br>
15.1 Swiss law applies exclusively, excluding the UN Convention on Contracts for the International Sale of Goods.<br><br>
15.2 Exclusive place of jurisdiction is Zurich, Switzerland.<br><br>
15.3 Should individual provisions be invalid, the validity of the remaining provisions remains unaffected.<br><br>
15.4 No verbal collateral agreements exist.
</div>
        """, unsafe_allow_html=True)

    if st.button("OK", use_container_width=True, type="primary"):
        st.query_params.clear()
        st.rerun()


def render_privacy_modal():
    """Rendert den Datenschutz-Hinweis."""
    privacy_text = {
        "en": """
**Privacy Policy - MINU Minutes AI (Beta)**

**Data Processing**
MINU Minutes AI uses OpenAI's services (Whisper API for transcription, GPT-4o for protocol generation) to process your audio files. Your audio data is sent to OpenAI servers for processing.

**Data Storage**
- Your registration data (company name, email address) is stored to provide you access to the beta version
- Audio files are processed temporarily and not permanently stored on our servers

**Marketing Use**
By registering, you consent to receiving occasional updates and information about MINU Minutes AI and MINU via email. You can unsubscribe at any time.

**Your Rights**
You have the right to access, correct, or delete your personal data at any time. Contact us at: info@minu-ai.ch

**Responsible Party**
MINU-AI
minu-ai.ch
        """,
        "de": """
**Datenschutzerklärung - MINU Minutes AI (Beta)**

**Datenverarbeitung**
MINU Minutes AI nutzt die Dienste von OpenAI (Whisper API für Transkription, GPT-4o für Protokollerstellung) zur Verarbeitung Ihrer Audio-Dateien. Ihre Audiodaten werden zur Verarbeitung an OpenAI-Server übermittelt.

**Datenspeicherung**
- Ihre Registrierungsdaten (Firmenname, E-Mail-Adresse) werden gespeichert, um Ihnen Zugang zur Beta-Version zu gewähren
- Audio-Dateien werden temporär verarbeitet und nicht dauerhaft auf unseren Servern gespeichert

**Marketing-Nutzung**
Mit der Registrierung stimmen Sie zu, gelegentlich Updates und Informationen über MINU Minutes AI per E-Mail zu erhalten. Sie können sich jederzeit abmelden.

**Ihre Rechte**
Sie haben jederzeit das Recht auf Auskunft, Berichtigung oder Löschung Ihrer personenbezogenen Daten. Kontaktieren Sie uns unter: info@minu-ai.ch

**Verantwortliche Stelle**
MINU-AI
minu-ai.ch
        """,
        "fr": """
**Politique de confidentialité - MINU Minutes AI (Bêta)**

**Traitement des données**
MINU Minutes AI utilise les services d'OpenAI (API Whisper pour la transcription, GPT-4o pour la génération de procès-verbaux) pour traiter vos fichiers audio. Vos données audio sont envoyées aux serveurs OpenAI pour traitement.

**Stockage des données**
- Vos données d'inscription (nom de l'entreprise, adresse e-mail) sont stockées pour vous donner accès à la version bêta
- Les fichiers audio sont traités temporairement et ne sont pas stockés de façon permanente sur nos serveurs

**Utilisation marketing**
En vous inscrivant, vous acceptez de recevoir occasionnellement des mises à jour et des informations sur MINU Minutes AI par e-mail. Vous pouvez vous désinscrire à tout moment.

**Vos droits**
Vous avez le droit d'accéder, de corriger ou de supprimer vos données personnelles à tout moment. Contactez-nous à: info@minu-ai.ch

**Responsable**
MINU-AI
minu-ai.ch
        """,
        "it": """
**Informativa sulla privacy - MINU Minutes AI (Beta)**

**Trattamento dei dati**
MINU Minutes AI utilizza i servizi di OpenAI (API Whisper per la trascrizione, GPT-4o per la generazione di verbali) per elaborare i tuoi file audio. I tuoi dati audio vengono inviati ai server OpenAI per l'elaborazione.

**Archiviazione dei dati**
- I tuoi dati di registrazione (nome azienda, indirizzo e-mail) vengono memorizzati per darti accesso alla versione beta
- I file audio vengono elaborati temporaneamente e non vengono archiviati permanentemente sui nostri server

**Utilizzo marketing**
Registrandoti, acconsenti a ricevere occasionalmente aggiornamenti e informazioni su MINU Minutes AI via e-mail. Puoi annullare l'iscrizione in qualsiasi momento.

**I tuoi diritti**
Hai il diritto di accedere, correggere o eliminare i tuoi dati personali in qualsiasi momento. Contattaci a: info@minu-ai.ch

**Responsabile**
MINU-AI
minu-ai.ch
        """
    }

    lang = st.session_state.get("language", "en")
    return privacy_text.get(lang, privacy_text["en"])


def main():
    st.set_page_config(
        page_title="MINU Minutes AI",
        page_icon="favicon.png",
        layout="centered",
        initial_sidebar_state="collapsed",
        menu_items={
            'Get Help': None,
            'Report a bug': None,
            'About': None
        }
    )

    # PWA Meta Tags und Service Worker
    st.markdown("""
        <link rel="manifest" href="/static/manifest.json">
        <meta name="theme-color" content="#7c3aed">
        <meta name="apple-mobile-web-app-capable" content="yes">
        <meta name="apple-mobile-web-app-status-bar-style" content="default">
        <meta name="apple-mobile-web-app-title" content="Minutes AI">
        <link rel="apple-touch-icon" href="/static/icons/icon-192x192.png">
        <meta name="mobile-web-app-capable" content="yes">
        <meta name="application-name" content="MINU Minutes AI">
        <meta name="msapplication-TileColor" content="#7c3aed">
        <meta name="msapplication-TileImage" content="/static/icons/icon-192x192.png">
        <script>
            if ('serviceWorker' in navigator) {
                window.addEventListener('load', function() {
                    navigator.serviceWorker.register('/static/sw.js')
                        .then(function(registration) {
                            console.log('ServiceWorker registered:', registration.scope);
                        })
                        .catch(function(error) {
                            console.log('ServiceWorker registration failed:', error);
                        });
                });
            }
        </script>
    """, unsafe_allow_html=True)

    # Umami Analytics (cookieless, DSGVO-konform)
    import streamlit.components.v1 as umami_components
    if ANALYTICS_SCRIPT:
        umami_components.html(f'<script defer src="{ANALYTICS_SCRIPT}" data-website-id="{ANALYTICS_ID}"></script>', height=0)

    # Custom CSS laden
    st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

    # Custom Footer mit Drop-ups für alle Links
    current_lang = st.session_state.get("language", "de")
    impressum_label = {"de": "Impressum", "en": "Legal Notice", "fr": "Mentions légales", "it": "Impronta"}
    agb_label = {"de": "AGB", "en": "Terms", "fr": "CGV", "it": "CGC"}
    datenschutz_label = {"de": "Datenschutz", "en": "Privacy", "fr": "Confidentialité", "it": "Privacy"}
    install_label = {"de": "Als App installieren", "en": "Install as App", "fr": "Installer l'app", "it": "Installa l'app"}

    # Install-Anleitung Inhalt
    install_content = {
        "de": '''<div style="font-weight: 600; margin-bottom: 10px;">📱 Als App installieren</div>
            <div style="font-size: 12px; color: #666; line-height: 1.6;">
                <b>iPhone/iPad:</b><br>Safari öffnen → Teilen-Button → "Zum Home-Bildschirm"<br><br>
                <b>Android:</b><br>Chrome öffnen → Menü (⋮) → "App installieren"<br><br>
                <b>Desktop:</b><br>Chrome/Edge → Adressleiste → Installieren-Icon
            </div>''',
        "en": '''<div style="font-weight: 600; margin-bottom: 10px;">📱 Install as App</div>
            <div style="font-size: 12px; color: #666; line-height: 1.6;">
                <b>iPhone/iPad:</b><br>Open Safari → Share button → "Add to Home Screen"<br><br>
                <b>Android:</b><br>Open Chrome → Menu (⋮) → "Install app"<br><br>
                <b>Desktop:</b><br>Chrome/Edge → Address bar → Install icon
            </div>''',
        "fr": '''<div style="font-weight: 600; margin-bottom: 10px;">📱 Installer l'application</div>
            <div style="font-size: 12px; color: #666; line-height: 1.6;">
                <b>iPhone/iPad:</b><br>Ouvrir Safari → Bouton partager → "Sur l'écran d'accueil"<br><br>
                <b>Android:</b><br>Ouvrir Chrome → Menu (⋮) → "Installer l'appli"<br><br>
                <b>Bureau:</b><br>Chrome/Edge → Barre d'adresse → Icône d'installation
            </div>''',
        "it": '''<div style="font-weight: 600; margin-bottom: 10px;">📱 Installa come App</div>
            <div style="font-size: 12px; color: #666; line-height: 1.6;">
                <b>iPhone/iPad:</b><br>Apri Safari → Pulsante condividi → "Aggiungi a Home"<br><br>
                <b>Android:</b><br>Apri Chrome → Menu (⋮) → "Installa app"<br><br>
                <b>Desktop:</b><br>Chrome/Edge → Barra indirizzi → Icona installa
            </div>'''
    }

    # AGB Kurzfassung
    agb_content = {
        "de": '''<div style="font-weight: 600; margin-bottom: 10px;">📋 AGB – Kurzfassung</div>
            <div style="font-size: 11px; color: #666; line-height: 1.5; max-height: 250px; overflow-y: auto;">
                <b>B2B-Dienst:</b> Nur für Unternehmen, kein Widerrufsrecht.<br><br>
                <b>Verfügbarkeit:</b> CH, EU, EWR, UK. Ausgeschlossen: USA, CN, RU, IN, BR.<br><br>
                <b>Zahlung:</b> Monatlich/jährlich via Stripe, automatische Verlängerung.<br><br>
                <b>Kündigung:</b> Jederzeit zum Ende der Laufzeit im Kundenportal.<br><br>
                <b>KI-Disclaimer:</b> Ergebnisse sind Arbeitshilfen, keine rechtssichere Dokumente. Prüfpflicht beim Kunden.<br><br>
                <b>Haftung:</b> Begrenzt auf gezahlte Entgelte der letzten 12 Monate.<br><br>
                <b>Recht:</b> Schweizer Recht, Gerichtsstand Zürich.
            </div>''',
        "en": '''<div style="font-weight: 600; margin-bottom: 10px;">📋 Terms – Summary</div>
            <div style="font-size: 11px; color: #666; line-height: 1.5; max-height: 250px; overflow-y: auto;">
                <b>B2B Service:</b> Business customers only, no withdrawal right.<br><br>
                <b>Availability:</b> CH, EU, EEA, UK. Excluded: USA, CN, RU, IN, BR.<br><br>
                <b>Payment:</b> Monthly/yearly via Stripe, auto-renewal.<br><br>
                <b>Cancellation:</b> Anytime effective end of period via customer portal.<br><br>
                <b>AI Disclaimer:</b> Results are work aids, not legally binding documents. Customer must verify.<br><br>
                <b>Liability:</b> Limited to fees paid in last 12 months.<br><br>
                <b>Law:</b> Swiss law, jurisdiction Zurich.
            </div>''',
        "fr": '''<div style="font-weight: 600; margin-bottom: 10px;">📋 CGV – Résumé</div>
            <div style="font-size: 11px; color: #666; line-height: 1.5; max-height: 250px; overflow-y: auto;">
                <b>Service B2B:</b> Entreprises uniquement, pas de droit de rétractation.<br><br>
                <b>Disponibilité:</b> CH, UE, EEE, UK. Exclus: USA, CN, RU, IN, BR.<br><br>
                <b>Paiement:</b> Mensuel/annuel via Stripe, renouvellement automatique.<br><br>
                <b>Résiliation:</b> À tout moment via le portail client.<br><br>
                <b>Disclaimer IA:</b> Résultats = aides au travail, pas de documents juridiques. Vérification par le client.<br><br>
                <b>Responsabilité:</b> Limitée aux frais des 12 derniers mois.<br><br>
                <b>Droit:</b> Droit suisse, juridiction Zurich.
            </div>''',
        "it": '''<div style="font-weight: 600; margin-bottom: 10px;">📋 CGC – Riepilogo</div>
            <div style="font-size: 11px; color: #666; line-height: 1.5; max-height: 250px; overflow-y: auto;">
                <b>Servizio B2B:</b> Solo aziende, nessun diritto di recesso.<br><br>
                <b>Disponibilità:</b> CH, UE, SEE, UK. Esclusi: USA, CN, RU, IN, BR.<br><br>
                <b>Pagamento:</b> Mensile/annuale via Stripe, rinnovo automatico.<br><br>
                <b>Disdetta:</b> In qualsiasi momento tramite il portale clienti.<br><br>
                <b>Disclaimer IA:</b> Risultati = ausili di lavoro, non documenti legali. Verifica da parte del cliente.<br><br>
                <b>Responsabilità:</b> Limitata ai costi degli ultimi 12 mesi.<br><br>
                <b>Diritto:</b> Diritto svizzero, foro Zurigo.
            </div>'''
    }

    # Datenschutz Kurzfassung
    datenschutz_content = {
        "de": '''<div style="font-weight: 600; margin-bottom: 10px;">🔒 Datenschutz – Kurzfassung</div>
            <div style="font-size: 11px; color: #666; line-height: 1.5; max-height: 250px; overflow-y: auto;">
                <b>Verantwortlicher:</b> SPEKTRUM Partner GmbH, Zürich<br><br>
                <b>Daten:</b> E-Mail, Firma, Audiodaten (für Transkription)<br><br>
                <b>Speicherung:</b> Audiodaten werden nach Verarbeitung gelöscht. Keine dauerhafte Speicherung von Inhalten.<br><br>
                <b>Drittanbieter:</b> OpenAI (Transkription), Mistral AI (Protokoll), Stripe (Zahlung) – alle DSGVO-konform.<br><br>
                <b>Analytics:</b> Umami (cookieless, keine persönlichen Daten)<br><br>
                <b>Ihre Rechte:</b> Auskunft, Berichtigung, Löschung, Widerspruch<br><br>
                <b>Kontakt:</b> minutes-ai@spekt.ch
            </div>''',
        "en": '''<div style="font-weight: 600; margin-bottom: 10px;">🔒 Privacy – Summary</div>
            <div style="font-size: 11px; color: #666; line-height: 1.5; max-height: 250px; overflow-y: auto;">
                <b>Controller:</b> SPEKTRUM Partner GmbH, Zurich<br><br>
                <b>Data:</b> Email, company, audio data (for transcription)<br><br>
                <b>Storage:</b> Audio deleted after processing. No permanent content storage.<br><br>
                <b>Third parties:</b> OpenAI (transcription), Mistral AI (protocol), Stripe (payment) – all GDPR compliant.<br><br>
                <b>Analytics:</b> Umami (cookieless, no personal data)<br><br>
                <b>Your rights:</b> Access, rectification, erasure, objection<br><br>
                <b>Contact:</b> minutes-ai@spekt.ch
            </div>''',
        "fr": '''<div style="font-weight: 600; margin-bottom: 10px;">🔒 Confidentialité – Résumé</div>
            <div style="font-size: 11px; color: #666; line-height: 1.5; max-height: 250px; overflow-y: auto;">
                <b>Responsable:</b> SPEKTRUM Partner GmbH, Zurich<br><br>
                <b>Données:</b> E-mail, entreprise, audio (pour transcription)<br><br>
                <b>Stockage:</b> Audio supprimé après traitement. Pas de stockage permanent.<br><br>
                <b>Tiers:</b> OpenAI, Mistral AI, Stripe – tous conformes RGPD.<br><br>
                <b>Analytics:</b> Umami (sans cookies, sans données personnelles)<br><br>
                <b>Vos droits:</b> Accès, rectification, suppression, opposition<br><br>
                <b>Contact:</b> minutes-ai@spekt.ch
            </div>''',
        "it": '''<div style="font-weight: 600; margin-bottom: 10px;">🔒 Privacy – Riepilogo</div>
            <div style="font-size: 11px; color: #666; line-height: 1.5; max-height: 250px; overflow-y: auto;">
                <b>Titolare:</b> SPEKTRUM Partner GmbH, Zurigo<br><br>
                <b>Dati:</b> E-mail, azienda, audio (per trascrizione)<br><br>
                <b>Archiviazione:</b> Audio eliminato dopo elaborazione. Nessun storage permanente.<br><br>
                <b>Terze parti:</b> OpenAI, Mistral AI, Stripe – tutti conformi GDPR.<br><br>
                <b>Analytics:</b> Umami (senza cookie, senza dati personali)<br><br>
                <b>I tuoi diritti:</b> Accesso, rettifica, cancellazione, opposizione<br><br>
                <b>Contatto:</b> minutes-ai@spekt.ch
            </div>'''
    }

    # Footer HTML mit String-Konkatenation (vermeidet f-string Probleme)
    footer_css = '''
    <style>
        .footer-dropup {
            position: relative;
            display: inline-block;
        }
        .footer-dropup-content {
            display: none;
            position: absolute;
            bottom: 100%;
            left: 50%;
            transform: translateX(-50%);
            background: #ffffff;
            border: 1px solid #e4e3df;
            border-radius: 12px;
            box-shadow: 0 -8px 24px rgba(0,0,0,0.15);
            padding: 16px 20px;
            min-width: 300px;
            max-width: 340px;
            margin-bottom: 12px;
            z-index: 1000;
        }
        .footer-dropup-content::after {
            content: "";
            position: absolute;
            bottom: -8px;
            left: 50%;
            transform: translateX(-50%);
            border-width: 8px 8px 0 8px;
            border-style: solid;
            border-color: #ffffff transparent transparent transparent;
        }
        .footer-dropup:hover .footer-dropup-content {
            display: block;
        }
        .footer-link {
            color: #86868b;
            text-decoration: none !important;
            font-size: 11px;
            cursor: pointer;
        }
        .footer-link:hover {
            color: #7c3aed;
            text-decoration: none !important;
        }
        a.footer-link, a.footer-link:visited, a.footer-link:active {
            color: #86868b;
            text-decoration: none !important;
        }
        a.footer-link:hover {
            color: #7c3aed;
            text-decoration: none !important;
        }
    </style>
    '''

    impressum_html = '''<div style="text-align: center; margin-bottom: 12px;">
        <span style="font-size: 18px; font-weight: 700; background: linear-gradient(90deg, #7c3aed, #2dd4bf); -webkit-background-clip: text; -webkit-text-fill-color: transparent;">MINU-AI</span>
    </div>
    <div style="background: #f8f8f8; border-radius: 8px; padding: 10px; margin-bottom: 10px;">
        <div style="font-weight: 600; color: #333; font-size: 13px;">SPEKTRUM Partner GmbH</div>
        <div style="color: #666; font-size: 12px;">Josefstrasse 181, 8005 Zürich</div>
    </div>
    <div style="font-size: 11px; color: #666; line-height: 1.6;">
        <div>📧 minutes-ai@spekt.ch</div>
        <div>🌐 minu-ai.ch</div>
        <div>🏢 UID: CH-020.4.068.247-6</div>
    </div>'''

    footer_html = (
        footer_css
        + '<div class="custom-footer">'
        + '<span style="color: #86868b; font-size: 11px;">🇨🇭 Swiss Developer · 🇪🇺 EU-DSGVO konform</span>'
        + '<span style="color: #d2d2d7; margin: 0 10px;">|</span>'
        + '<div class="footer-dropup">'
        + '<span class="footer-link">📱 ' + install_label.get(current_lang, "Install") + '</span>'
        + '<div class="footer-dropup-content">' + install_content.get(current_lang, install_content["en"]) + '</div>'
        + '</div>'
        + '<span style="color: #d2d2d7; margin: 0 10px;">|</span>'
        + '<div class="footer-dropup">'
        + '<span class="footer-link">' + impressum_label.get(current_lang, "Impressum") + '</span>'
        + '<div class="footer-dropup-content">' + impressum_html + '</div>'
        + '</div>'
        + '<span style="color: #d2d2d7; margin: 0 10px;">|</span>'
        + '<a href="?agb=1" class="footer-link">' + agb_label.get(current_lang, "AGB") + '</a>'
        + '<span style="color: #d2d2d7; margin: 0 10px;">|</span>'
        + '<a href="?datenschutz=1" class="footer-link">' + datenschutz_label.get(current_lang, "Datenschutz") + '</a>'
        + '</div>'
    )

    st.markdown(footer_html, unsafe_allow_html=True)

    # AGB und Datenschutz als Vollseiten-Ansicht (nach Footer, damit Kopf-/Fusszeile sichtbar bleibt)
    if "agb" in st.query_params:
        show_agb_fullpage()
        st.stop()

    if "datenschutz" in st.query_params:
        show_datenschutz_fullpage()
        st.stop()

    # Free Plan Info Dialog anzeigen wenn ?freeplaninfo=1 oder session state
    if "freeplaninfo" in st.query_params or st.session_state.get("show_free_plan_info", False):
        st.session_state.show_free_plan_info = False
        show_free_plan_info_dialog()

    # Feedback-Rating verarbeiten wenn ?feedback=X
    if "feedback" in st.query_params:
        try:
            rating = int(st.query_params["feedback"])
            if 1 <= rating <= 5:
                save_feedback_rating(rating)
                show_feedback_thank_you(rating)
                st.stop()
        except:
            pass

    # Logout behandeln
    if "logout" in st.query_params:
        st.session_state.authenticated = False
        st.session_state.user_email = ""
        st.session_state.user_company = ""
        st.session_state.current_page = "landing"
        st.query_params.clear()
        st.rerun()

    # Upgrade behandeln - zur Register-Seite weiterleiten für Abo-Upgrade (direkt zu PRO)
    if "upgrade" in st.query_params:
        st.session_state.current_page = "register"
        st.session_state.selected_plan = "team"  # Direkt zu PRO
        st.query_params.clear()
        st.rerun()

    # Stripe Customer Portal für Abo-Verwaltung
    if "manage_billing" in st.query_params:
        st.query_params.clear()
        user_email = st.session_state.get("user_email", "")
        if user_email:
            user_sub = get_user_subscription(user_email)
            customer_id = user_sub.get("stripe_customer_id") if user_sub else None
            if customer_id:
                # Stripe Portal URL erstellen
                return_url = "{APP_URL}/"
                portal_url = create_stripe_portal_session(customer_id, return_url)
                if portal_url:
                    st.markdown(f'<meta http-equiv="refresh" content="0;url={portal_url}">', unsafe_allow_html=True)
                    st.info("Du wirst zum Stripe Kundenportal weitergeleitet...")
                    st.stop()
            # Fallback wenn keine Stripe Customer ID vorhanden
            st.warning("Bitte kontaktiere support@minu-ai.ch für Abo-Änderungen.")
        st.rerun()

    # Stripe Payment Success Handler
    if "payment" in st.query_params:
        payment_status = st.query_params.get("payment")

        if payment_status == "success" and "session_id" in st.query_params:
            session_id = st.query_params.get("session_id")
            payment_info = verify_stripe_payment(session_id)

            if payment_info.get("success"):
                user_email = payment_info["user_email"]

                # Subscription upgraden
                upgrade_user_subscription(
                    user_email,
                    payment_info["plan"],
                    payment_info["billing_cycle"],
                    payment_info.get("subscription_id")
                )

                # Benutzer automatisch einloggen (Email normalisieren)
                email_normalized = user_email.lower().strip()
                st.session_state.authenticated = True
                st.session_state.is_admin = False
                st.session_state.user_email = email_normalized

                # Company aus Subscription laden
                user_sub = get_user_subscription(email_normalized)
                if user_sub:
                    st.session_state.user_company = user_sub.get("company", "")

                # Erfolgsmeldung speichern für Anzeige nach Redirect
                plan_name = "Starter" if payment_info["plan"] == "basic_solo" else "PRO"
                st.session_state.payment_success_message = f"🎉 Zahlung erfolgreich! Dein {plan_name}-Plan ist jetzt aktiv."
                log_activity("Payment Success", f"{user_email} - {plan_name}")

            st.query_params.clear()
            st.rerun()

        elif payment_status == "cancelled":
            st.warning("Zahlung abgebrochen. Du kannst es jederzeit erneut versuchen.")
            st.query_params.clear()

    # Passwortschutz
    if not check_password():
        st.stop()

    # Kopfzeile für eingeloggte Benutzer
    current_lang = st.session_state.get("language", "de")
    user_company = st.session_state.get("user_company", "")
    user_email = st.session_state.get("user_email", "")
    logout_text = {"de": "Abmelden", "en": "Logout", "fr": "Déconnexion", "it": "Disconnetti"}
    minutes_text = {"de": "Min. übrig", "en": "min. left", "fr": "min. restantes", "it": "min. rimanenti"}

    # Plan des Benutzers laden
    user_subscription = get_user_subscription(user_email) if user_email else None
    user_plan = user_subscription.get("plan", "free_trial") if user_subscription else "free_trial"

    # DEBUG: Zeige was geladen wird
    print(f"[DEBUG] user_email: '{user_email}' | user_plan: '{user_plan}' | subscription: {user_subscription}")
    plan_display_names = {
        "free_trial": "Free Trial",
        "basic_solo": "Starter",
        "team": "PRO"
    }
    plan_display = plan_display_names.get(user_plan, "Free Trial")
    plan_colors = {
        "free_trial": "#6b7280",  # Grau
        "basic_solo": "#7c3aed",  # Lila
        "team": "#2dd4bf"  # Türkis
    }
    plan_color = plan_colors.get(user_plan, "#6b7280")

    # Verbleibende Minuten laden
    remaining_mins = get_remaining_minutes(user_email) if user_email else 0
    remaining_mins_int = int(remaining_mins)

    # Farbe je nach Minuten-Stand
    if remaining_mins_int <= 0:
        mins_color = "#ef4444"  # Rot
    elif remaining_mins_int <= 10:
        mins_color = "#f59e0b"  # Orange
    else:
        mins_color = "#10b981"  # Grün

    # Upgrade-Text nur für Free Trial und Starter
    show_upgrade = user_plan in ["free_trial", "basic_solo"]

    # Abrechnung und Abo-Verwaltung nur für zahlende Nutzer (Starter/Pro)
    is_paid_user = user_plan != "free_trial"

    # Texte für Dropdown
    texts = {
        "de": {
            "upgrade": "Plan upgraden",
            "logout": "Abmelden",
            "settings": "Einstellungen",
            "minutes_this_month": "Minuten diesen Monat",
            "renews": "Erneuert am 1. des Monats",
            "help": "Hilfe & Support",
            "per_month": "/ Monat",
            "manage_billing": "Abo verwalten",
            "cancel_hint": "Zahlungsmethode, Rechnungen, Kundigung"
        },
        "en": {
            "upgrade": "Upgrade plan",
            "logout": "Logout",
            "settings": "Settings",
            "minutes_this_month": "Minutes this month",
            "renews": "Renews on the 1st",
            "help": "Help & Support",
            "per_month": "/ month",
            "manage_billing": "Manage subscription",
            "cancel_hint": "Payment method, invoices, cancellation"
        },
        "fr": {
            "upgrade": "Changer de forfait",
            "logout": "Deconnexion",
            "settings": "Parametres",
            "minutes_this_month": "Minutes ce mois",
            "renews": "Renouvellement le 1er",
            "help": "Aide & Support",
            "per_month": "/ mois",
            "manage_billing": "Gerer abonnement",
            "cancel_hint": "Paiement, factures, resiliation"
        },
        "it": {
            "upgrade": "Cambia piano",
            "logout": "Esci",
            "settings": "Impostazioni",
            "minutes_this_month": "Minuti questo mese",
            "renews": "Rinnovo il 1",
            "help": "Aiuto & Supporto",
            "per_month": "/ mese",
            "manage_billing": "Gestisci abbonamento",
            "cancel_hint": "Pagamento, fatture, disdetta"
        }
    }
    t = texts.get(current_lang, texts["de"])

    # Initialen aus E-Mail erstellen
    initials = user_email[:2].upper() if user_email else "??"
    if "@" in user_email:
        name_part = user_email.split("@")[0]
        if "." in name_part:
            parts = name_part.split(".")
            initials = (parts[0][0] + parts[1][0]).upper()
        else:
            initials = name_part[:2].upper()

    # Max Minuten je nach Plan
    max_minutes = {"free_trial": 60, "basic_solo": 180, "team": 600}.get(user_plan, 60)
    used_minutes = max_minutes - remaining_mins_int
    usage_percent = min(100, max(0, (used_minutes / max_minutes) * 100)) if max_minutes > 0 else 0

    # Farbe für Usage-Bar
    if usage_percent >= 80:
        bar_color = "#f59e0b"  # Orange
    else:
        bar_color = "#7c3aed"  # Lila

    # Plan Badge Farben
    badge_colors = {
        "free_trial": ("background: #f3f4f6; color: #6b7280;", "Gratis"),
        "basic_solo": ("background: #f0ecff; color: #7c3aed;", f"CHF 19.– {t['per_month']}"),
        "team": ("background: #e0f7f4; color: #0d9488;", f"CHF 79.– {t['per_month']}")
    }
    badge_style, price_text = badge_colors.get(user_plan, badge_colors["free_trial"])

    # Bedingte Sektionen für das Dropdown (nur für zahlende Nutzer)
    billing_html = ''
    manage_billing_html = ''
    if is_paid_user:
        billing_html = '<div style="margin-bottom: 20px;"><div style="font-size: 12px; color: #999; text-transform: uppercase; margin-bottom: 8px;">Abrechnung</div><div style="background: #f8f8f8; border-radius: 8px; padding: 12px;"><div style="display: flex; align-items: center; gap: 10px;"><span style="background: #1a1a1a; color: white; padding: 4px 8px; border-radius: 4px; font-family: monospace; font-size: 11px;">VISA</span><span style="font-size: 13px; color: #666;">**** 4242</span></div></div></div>'
        manage_billing_html = f'<a href="?manage_billing=1" style="display: block; width: 100%; padding: 10px; background: #f5f5f5; color: #444; text-align: center; border-radius: 8px; text-decoration: none; font-weight: 500; margin-bottom: 10px; border: 1px solid #e0e0e0;"><span style="display: block;">{t["manage_billing"]}</span><span style="display: block; font-size: 11px; color: #888; font-weight: 400; margin-top: 2px;">{t["cancel_hint"]}</span></a>'

    # Buttons als separate Variablen (vermeidet f-string Probleme)
    upgrade_btn_html = '<a href="?upgrade=1" style="display: block; width: 100%; padding: 10px; background: #7c3aed; color: white; text-align: center; border-radius: 8px; text-decoration: none; font-weight: 500; margin-bottom: 10px;">' + t["upgrade"] + '</a>'
    logout_btn_html = '<a href="?logout=1" style="display: block; width: 100%; padding: 10px; background: #fee2e2; color: #dc2626; text-align: center; border-radius: 8px; text-decoration: none; font-weight: 500;">' + t["logout"] + '</a>'

    # Gesamter Dropdown-Inhalt als separate Variable (String-Konkatenation für Stabilität)
    dropdown_content_html = (
        '<div style="padding: 16px;">'
        '<div style="margin-bottom: 20px;">'
        '<div style="font-size: 12px; color: #999; text-transform: uppercase; margin-bottom: 8px;">PROFIL</div>'
        '<div style="display: flex; align-items: center; gap: 12px; margin-bottom: 12px;">'
        '<div style="width: 50px; height: 50px; border-radius: 50%; background: linear-gradient(135deg, #7c3aed, #2dd4bf); color: white; display: flex; align-items: center; justify-content: center; font-weight: 700; font-size: 18px;">' + initials + '</div>'
        '<div>'
        '<div style="font-weight: 500; color: #1a1a1a;">' + user_email + '</div>'
        '<div style="font-size: 13px; color: #666;">' + user_company + '</div>'
        '</div>'
        '</div>'
        '</div>'
        '<div style="margin-bottom: 20px;">'
        '<div style="font-size: 12px; color: #999; text-transform: uppercase; margin-bottom: 8px;">Plan &amp; Nutzung</div>'
        '<div style="background: #f8f8f8; border-radius: 8px; padding: 12px;">'
        '<div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 8px;">'
        '<span style="background: ' + plan_color + '; color: white; padding: 3px 10px; border-radius: 5px; font-size: 12px; font-weight: 600;">' + plan_display + '</span>'
        '<span style="font-weight: 600; color: #1a1a1a;">' + price_text + '</span>'
        '</div>'
        '<div style="display: flex; justify-content: space-between; font-size: 13px; color: #666; margin-bottom: 6px;">'
        '<span>' + t['minutes_this_month'] + '</span>'
        '<span>' + str(used_minutes) + ' / ' + str(max_minutes) + '</span>'
        '</div>'
        '<div style="height: 6px; background: #e0e0e0; border-radius: 3px; overflow: hidden;">'
        '<div style="height: 100%; background: linear-gradient(90deg, #7c3aed, #2dd4bf); width: ' + str(int(usage_percent)) + '%;"></div>'
        '</div>'
        '<div style="font-size: 12px; color: #999; margin-top: 6px;">' + t['renews'] + '</div>'
        '</div>'
        '</div>'
        + billing_html
        + upgrade_btn_html
        + manage_billing_html
        + logout_btn_html
        + '</div>'
    )

    st.markdown(f'''
    <style>
        .topnav {{
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            background: #ffffff;
            border-bottom: 1px solid #e4e3df;
            padding: 0 1.5rem;
            height: 52px;
            display: flex;
            align-items: center;
            justify-content: space-between;
            z-index: 1000;
            box-sizing: border-box;
        }}
        .topnav-brand {{
            display: flex;
            align-items: center;
            gap: 9px;
            text-decoration: none;
            color: #1a1a1a;
        }}
        .topnav-logo {{
            width: 26px;
            height: 26px;
            background: linear-gradient(135deg, #7c3aed, #2dd4bf);
            border-radius: 6px;
            display: flex;
            align-items: center;
            justify-content: center;
            color: white;
            font-weight: 700;
            font-size: 13px;
        }}
        .topnav-title {{
            font-weight: 600;
            font-size: 14.5px;
            letter-spacing: -0.01em;
        }}
        .topnav-title span {{
            color: #9a9a9a;
            font-weight: 400;
        }}
        .dropdown-wrapper {{
            position: relative;
            padding-bottom: 10px;
            margin-bottom: -10px;
        }}
        .nav-pill {{
            display: flex;
            align-items: center;
            gap: 8px;
            padding: 5px 12px;
            border-radius: 20px;
            background: #f5f5f7;
            border: 1px solid #e4e3df;
            font-size: 12.5px;
            font-weight: 500;
            color: #5c5c5c;
            cursor: pointer;
            transition: all 0.15s;
        }}
        .nav-pill:hover {{
            background: #ebebeb;
            border-color: #9a9a9a;
        }}
        .pill-minutes {{
            color: {mins_color};
            font-weight: 600;
            font-family: monospace;
            font-size: 12px;
        }}
        .pill-divider {{
            width: 1px;
            height: 14px;
            background: #d2d2d7;
        }}
        .pill-org {{
            display: flex;
            align-items: center;
            gap: 4px;
        }}
        .pill-org-icon {{
            width: 14px;
            height: 14px;
            background: linear-gradient(135deg, #7c3aed, #2dd4bf);
            border-radius: 3px;
            display: flex;
            align-items: center;
            justify-content: center;
            color: white;
            font-size: 8px;
            font-weight: 700;
        }}
        .chevron {{
            margin-left: 4px;
            width: 0;
            height: 0;
            border-left: 4px solid transparent;
            border-right: 4px solid transparent;
            border-top: 5px solid #9a9a9a;
        }}
        /* Settings Toggle Checkbox - versteckt */
        .settings-checkbox {{
            position: absolute;
            opacity: 0;
            pointer-events: none;
        }}
        .account-dropdown {{
            position: absolute;
            top: 100%;
            right: 0;
            width: 340px;
            background: #ffffff;
            border: 1px solid #e4e3df;
            border-radius: 14px;
            box-shadow: 0 12px 40px rgba(0,0,0,0.12);
            z-index: 1001;
            opacity: 0;
            pointer-events: none;
            transform: translateY(-4px);
            transition: all 0.2s ease;
            overflow: visible;
            margin-top: 8px;
        }}
        /* Dropdown Toggle - Checkbox versteckt */
        .dropdown-toggle {{
            display: none;
        }}
        /* Dropdown aktiv via: Checkbox checked oder Hover */
        .dropdown-toggle:checked ~ .dropdown-content .account-dropdown,
        .dropdown-wrapper:hover .account-dropdown {{
            opacity: 1;
            pointer-events: all;
            transform: translateY(0);
        }}
        /* Overlay zum Schliessen bei Klick ausserhalb */
        .dropdown-overlay {{
            display: none;
            position: fixed;
            top: 0;
            left: 0;
            right: 0;
            bottom: 0;
            z-index: 1000;
            background: transparent;
            cursor: default;
        }}
        .dropdown-toggle:checked ~ .dropdown-content .dropdown-overlay {{
            display: block;
        }}
        .nav-pill {{
            position: relative;
            z-index: 1002;
        }}
        /* Detaillierte Settings-Ansicht zeigen (mit Abo & Rechnung) */
        .dd-standard {{
            display: none;
        }}
        .dd-settings {{
            display: block;
        }}
        .dd-user {{
            padding: 1rem 1.25rem;
            border-bottom: 1px solid #eeedea;
            display: flex;
            align-items: center;
            gap: 12px;
        }}
        .dd-avatar {{
            width: 40px;
            height: 40px;
            border-radius: 50%;
            background: linear-gradient(135deg, #7c3aed, #2dd4bf);
            color: white;
            display: flex;
            align-items: center;
            justify-content: center;
            font-weight: 700;
            font-size: 14px;
            flex-shrink: 0;
        }}
        .dd-user-info {{
            min-width: 0;
        }}
        .dd-user-email {{
            font-weight: 500;
            font-size: 0.9rem;
            color: #1a1a1a;
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
        }}
        .dd-user-org {{
            font-size: 0.78rem;
            color: #5c5c5c;
            margin-top: 1px;
        }}
        .dd-plan {{
            padding: 1rem 1.25rem;
            border-bottom: 1px solid #eeedea;
        }}
        .dd-plan-header {{
            display: flex;
            align-items: center;
            justify-content: space-between;
            margin-bottom: 10px;
        }}
        .dd-plan-badge {{
            display: inline-flex;
            align-items: center;
            padding: 4px 10px;
            border-radius: 5px;
            font-size: 0.72rem;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 0.04em;
            {badge_style}
        }}
        .dd-plan-price {{
            font-family: monospace;
            font-size: 0.8rem;
            color: #5c5c5c;
        }}
        .dd-usage-header {{
            display: flex;
            justify-content: space-between;
            margin-bottom: 6px;
            font-size: 0.8rem;
        }}
        .dd-usage-label {{
            color: #5c5c5c;
            font-weight: 500;
        }}
        .dd-usage-value {{
            font-family: monospace;
            font-size: 0.78rem;
            color: #1a1a1a;
            font-weight: 500;
        }}
        .dd-bar {{
            width: 100%;
            height: 6px;
            background: #f0efeb;
            border-radius: 3px;
            overflow: hidden;
        }}
        .dd-bar-fill {{
            height: 100%;
            border-radius: 3px;
            background: {bar_color};
            width: {usage_percent}%;
            transition: width 0.8s ease;
        }}
        .dd-usage-hint {{
            font-size: 0.7rem;
            color: #9a9a9a;
            margin-top: 5px;
        }}
        .dd-menu {{
            padding: 6px 0;
            border-bottom: 1px solid #eeedea;
        }}
        .dd-menu-item {{
            display: flex;
            align-items: center;
            gap: 10px;
            padding: 10px 1.25rem;
            font-size: 0.88rem;
            color: #1a1a1a;
            cursor: pointer;
            transition: background 0.12s;
            text-decoration: none;
            border: none;
            background: none;
            width: 100%;
            text-align: left;
        }}
        .dd-menu-item:hover {{
            background: #f5f5f7;
        }}
        .dd-menu-item .item-icon {{
            width: 20px;
            text-align: center;
            font-size: 14px;
        }}
        .dd-menu-item .item-badge {{
            font-size: 0.68rem;
            padding: 2px 6px;
            border-radius: 4px;
            font-weight: 600;
            background: #e0f7f4;
            color: #0d9488;
            margin-left: auto;
        }}
        .dd-menu-item.upgrade {{
            color: #7c3aed;
            font-weight: 600;
        }}
        .dd-menu-item.danger {{
            color: #ef4444;
        }}
        .dd-footer {{
            padding: 6px 0;
        }}
        .dd-footer .dd-menu-item {{
            color: #9a9a9a;
            font-size: 0.82rem;
        }}
        .dd-footer .dd-menu-item:hover {{
            color: #1a1a1a;
        }}
        /* lang-selector entfernt - Sprachauswahl nur auf der Hauptseite */
    </style>
    <nav class="topnav">
        <a href="/" style="text-decoration:none;">
            <span style="font-size:18px;font-weight:700;background:linear-gradient(90deg,#7c3aed,#2dd4bf);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;">MINU-AI</span>
        </a>
        <div style="display: flex; align-items: center;">
            <div class="dropdown-wrapper" id="profile-dropdown">
                <input type="checkbox" id="dropdown-toggle-cb" class="dropdown-toggle">
                <label for="dropdown-toggle-cb" class="nav-pill">
                    <span class="pill-minutes">{remaining_mins_int} min</span>
                    <span class="pill-divider"></span>
                    <span class="pill-org">
                        <span class="pill-org-icon">{user_company[0].upper() if user_company else 'M'}</span>
                        {user_company[:15]}{'...' if len(user_company) > 15 else ''}
                    </span>
                    <span class="chevron"></span>
                </label>
                <div class="dropdown-content">
                    <label for="dropdown-toggle-cb" class="dropdown-overlay"></label>
                    <div class="account-dropdown">{dropdown_content_html}</div>
                </div>
            </div>
        </div>
    </nav>
    <div style="height: 52px;"></div>
    ''', unsafe_allow_html=True)

    # Erfolgsmeldung nach Zahlung anzeigen (einmalig)
    if "payment_success_message" in st.session_state:
        st.success(st.session_state.payment_success_message)
        del st.session_state.payment_success_message

    # Upgrade-Dialog wenn keine Minuten mehr übrig
    if remaining_mins_int <= 0:
        upgrade_title = {"de": "Plan upgraden!", "en": "Upgrade your plan!", "fr": "Mettez à niveau!", "it": "Aggiorna il piano!"}
        upgrade_text = {"de": f"Du hast noch {remaining_mins_int} Minuten verbleibend. Benötigst du mehr?",
                        "en": f"You have {remaining_mins_int} minutes remaining. Need more?",
                        "fr": f"Il vous reste {remaining_mins_int} minutes. Besoin de plus?",
                        "it": f"Hai {remaining_mins_int} minuti rimanenti. Ne hai bisogno di più?"}
        upgrade_btn = {"de": "Jetzt upgraden", "en": "Upgrade now", "fr": "Mettre à niveau", "it": "Aggiorna ora"}

        st.markdown(f"""
        <div style="
            background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
            border-radius: 16px;
            padding: 24px;
            margin: 20px 0;
            text-align: center;
            box-shadow: 0 4px 20px rgba(0,0,0,0.3);
        ">
            <h3 style="color: white; margin: 0 0 12px 0; font-size: 20px;">{upgrade_title.get(current_lang, 'Upgrade!')}</h3>
            <p style="color: #a0a0a0; margin: 0 0 20px 0; font-size: 15px;">{upgrade_text.get(current_lang, 'Need more minutes?')}</p>
            <a href="?upgrade=1" style="
                display: inline-block;
                background: linear-gradient(135deg, #7c3aed, #2dd4bf);
                color: white;
                padding: 12px 28px;
                border-radius: 25px;
                text-decoration: none;
                font-weight: 600;
                font-size: 15px;
                box-shadow: 0 4px 15px rgba(124, 58, 237, 0.4);
            ">{upgrade_btn.get(current_lang, 'Upgrade now')}</a>
        </div>
        """, unsafe_allow_html=True)

    # PWA deaktiviert - verursacht Probleme auf iOS
    # Das apple-touch-icon für den Home-Bildschirm wird über nginx serviert
    # st.markdown(PWA_META_TAGS, unsafe_allow_html=True)
    # st.markdown(PWA_SERVICE_WORKER, unsafe_allow_html=True)

    # Sidebar rendern
    render_sidebar()

    # API-Keys prüfen
    openai_api_key = get_secret("OPENAI_API_KEY")
    mistral_api_key = get_secret("MISTRAL_API_KEY")

    if not openai_api_key:
        st.error("OPENAI_API_KEY nicht gefunden! Wird für Transkription benötigt.")
        st.stop()

    if not mistral_api_key:
        st.error("MISTRAL_API_KEY nicht gefunden! Wird für Protokollerstellung benötigt.")
        st.stop()

    # OpenAI Client für Whisper Transkription
    openai_client = OpenAI(api_key=openai_api_key)

    # Mistral API Key für Protokollerstellung (EU-basiert, via HTTP)
    # (kein SDK nötig - verwende requests)

    # Session State initialisieren
    if "transcript" not in st.session_state:
        st.session_state.transcript = None
    if "protocol" not in st.session_state:
        st.session_state.protocol = None
    if "pdf_bytes" not in st.session_state:
        st.session_state.pdf_bytes = None
    if "docx_bytes" not in st.session_state:
        st.session_state.docx_bytes = None
    if "processing" not in st.session_state:
        st.session_state.processing = False
    if "error" not in st.session_state:
        st.session_state.error = None
    if "language" not in st.session_state:
        # GeoIP-basierte Spracherkennung beim ersten Besuch
        if "language_auto_detected" not in st.session_state:
            st.session_state.language = detect_language_from_ip()
            st.session_state.language_auto_detected = True
        else:
            st.session_state.language = "en"

    # Hero Header - Apple Style mit Logo mittig
    st.markdown("")

    # Titel im neuen Layout-Stil
    current_lang = st.session_state.get("language", "en")
    subtitle_text = {"en": "Minutes AI", "de": "Protokoll KI", "fr": "Minutes IA", "it": "Verbali IA"}.get(current_lang, "Minutes AI")
    st.markdown(f"""
        <h1 style="text-align: center; font-size: 2.8rem; font-weight: 700; letter-spacing: -0.02em; color: #1d1d1f; margin-bottom: 0; line-height: 1.2;">
            <span style="font-size: 5.5rem; display: block; margin-bottom: -10px; background: linear-gradient(90deg, #7c3aed, #2dd4bf); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; font-weight: 800; font-style: italic;">MINU-AI</span>
            <span style="background: linear-gradient(90deg, #7c3aed, #2dd4bf); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text;">
                {get_text('slogan_gradient')}
            </span><br>
            {subtitle_text}
        </h1>
        <p style="text-align: center; font-size: 14px; color: #86868b; font-weight: 400; margin-top: 0.5rem; margin-bottom: 0.5rem;">
            Beta v1
        </p>
    """, unsafe_allow_html=True)

    # Free Plan Info Button deaktiviert - testet ob Profil-Dropdown funktioniert
    # if user_plan == "free_trial":
    #     render_free_plan_info_button()

    st.markdown(f"""
        <p style="text-align: center; font-size: 1.1rem; color: #6b7280; margin-top: 0.5rem; margin-bottom: 0.5rem;">
            {get_text('audio_conversion_text')}
        </p>
    """, unsafe_allow_html=True)

    # =========================================================================
    # FERTIG - Dokumente bereit
    # =========================================================================
    if st.session_state.pdf_bytes:
        st.markdown(f"<p style='text-align:center; font-size:17px; color:#34c759;'>{get_text('your_protocol_ready')}</p>", unsafe_allow_html=True)
        st.markdown("")

        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
        filename_pdf = f"Protokoll_{timestamp}.pdf"
        filename_docx = f"Protokoll_{timestamp}.docx"

        # Download Buttons
        col1, col2 = st.columns(2)
        with col1:
            st.download_button(
                label=get_text("download_pdf"),
                data=st.session_state.pdf_bytes,
                file_name=filename_pdf,
                mime="application/pdf",
                use_container_width=True
            )
        with col2:
            st.download_button(
                label=get_text("download_word"),
                data=st.session_state.docx_bytes,
                file_name=filename_docx,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

        st.markdown("")
        st.markdown("---")
        st.markdown("")

        # E-Mail Versand
        st.markdown(f"<p style='text-align:center; color:#1d1d1f; font-weight:600;'>{get_text('send_email')}</p>", unsafe_allow_html=True)
        st.markdown("")

        recipient = st.text_input("", placeholder=get_text("email_placeholder"), label_visibility="collapsed")

        # Format-Auswahl mit Toggle-Buttons
        st.markdown(f"<p style='text-align:center; color:#86868b; font-size:14px; margin-top:10px;'>{get_text('email_format_label')}</p>", unsafe_allow_html=True)

        # Session State für Format-Auswahl initialisieren
        if "email_send_pdf" not in st.session_state:
            st.session_state.email_send_pdf = True
        if "email_send_word" not in st.session_state:
            st.session_state.email_send_word = True

        col_pdf, col_word = st.columns(2)
        with col_pdf:
            pdf_label = "✓ PDF" if st.session_state.email_send_pdf else "PDF"
            pdf_style = "primary" if st.session_state.email_send_pdf else "secondary"
            if st.button(pdf_label, key="toggle_pdf", use_container_width=True, type=pdf_style):
                st.session_state.email_send_pdf = not st.session_state.email_send_pdf
                st.rerun()
        with col_word:
            word_label = "✓ Word" if st.session_state.email_send_word else "Word"
            word_style = "primary" if st.session_state.email_send_word else "secondary"
            if st.button(word_label, key="toggle_word", use_container_width=True, type=word_style):
                st.session_state.email_send_word = not st.session_state.email_send_word
                st.rerun()

        send_pdf = st.session_state.email_send_pdf
        send_word = st.session_state.email_send_word

        if st.button(get_text("send_button"), use_container_width=True, type="primary"):
            if not recipient:
                st.warning(get_text("email_error"))
            elif not send_pdf and not send_word:
                st.warning(get_text("email_format_error"))
            else:
                with st.spinner(get_text("processing")):
                    success, message = send_email_with_protocol(
                        st.session_state.pdf_bytes,
                        st.session_state.docx_bytes,
                        recipient,
                        f"Protokoll_{timestamp}",
                        send_pdf=send_pdf,
                        send_word=send_word
                    )
                    if success:
                        st.success(get_text("email_sent"))
                        formats = []
                        if send_pdf:
                            formats.append("PDF")
                        if send_word:
                            formats.append("Word")
                        log_activity("E-Mail gesendet", f"An: {recipient} ({', '.join(formats)})")
                    else:
                        st.error(message)
                        log_activity("E-Mail fehlgeschlagen", message)

        st.markdown("")

        # Neu starten
        if st.button(get_text("new_protocol"), use_container_width=True):
            st.session_state.transcript = None
            st.session_state.protocol = None
            st.session_state.pdf_bytes = None
            st.session_state.docx_bytes = None
            st.session_state.processing = False
            st.session_state.error = None
            st.rerun()

    # =========================================================================
    # UPLOAD - Warte auf Datei
    # =========================================================================
    else:
        # Recording Component (für alle Browser mit Mikrofon-Unterstützung)
        mic_select_label = {"de": "Mikrofon wählen", "en": "Select microphone", "fr": "Choisir microphone", "it": "Seleziona microfono"}
        mobile_recording_html = f'''
        <div id="mobile-recorder" style="display: block; text-align: center; padding: 10px 0; margin-bottom: 10px;">

            <!-- Mikrofon-Auswahl (versteckt bis Start geklickt) -->
            <div id="mic-selector-container" style="display: none; margin-bottom: 20px;">
                <label for="mic-select" style="display: block; color: #6b7280; font-size: 14px; margin-bottom: 8px;">
                    {mic_select_label.get(current_lang, 'Select microphone')}
                </label>
                <select id="mic-select" style="
                    padding: 12px 16px;
                    border-radius: 12px;
                    border: 2px solid #d1d1d6;
                    background: #f5f5f7;
                    font-size: 15px;
                    color: #1d1d1f;
                    min-width: 280px;
                    cursor: pointer;
                ">
                    <option value="">-- Mikrofon wählen --</option>
                </select>
                <br><br>
                <button id="confirm-mic-btn" onclick="confirmMicAndRecord()" style="
                    background: linear-gradient(90deg, #7c3aed, #2dd4bf);
                    color: white;
                    border: none;
                    padding: 14px 28px;
                    border-radius: 25px;
                    font-size: 16px;
                    font-weight: 600;
                    cursor: pointer;
                    box-shadow: 0 4px 15px rgba(124,58,237,0.3);
                ">
                    ✓ {get_text('start_recording').replace('🎙️ ', '')}
                </button>
                <button onclick="cancelMicSelection()" style="
                    background: transparent;
                    color: #6b7280;
                    border: none;
                    padding: 14px 20px;
                    font-size: 14px;
                    cursor: pointer;
                    margin-left: 10px;
                ">
                    ✕ Abbrechen
                </button>
            </div>

            <!-- Start Recording Button -->
            <button id="start-rec-btn" onclick="showMicSelector()" style="
                background: linear-gradient(90deg, #7c3aed, #2dd4bf);
                color: white;
                border: none;
                padding: 16px 32px;
                border-radius: 25px;
                font-size: 17px;
                font-weight: 600;
                cursor: pointer;
                display: inline-flex;
                align-items: center;
                justify-content: center;
                gap: 10px;
                box-shadow: 0 4px 15px rgba(124,58,237,0.3);
                width: 280px;
            ">
{get_text('start_recording')}
            </button>

            <!-- Recording in Progress -->
            <div id="recording-ui" style="display: none;">
                <div style="margin-bottom: 15px;">
                    <span id="rec-indicator" style="
                        display: inline-block;
                        width: 12px;
                        height: 12px;
                        background: #ef4444;
                        border-radius: 50%;
                        margin-right: 8px;
                        animation: pulse 1s infinite;
                    "></span>
                    <span id="rec-timer" style="font-size: 24px; font-weight: 600; color: #1d1d1f; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;">00:00</span>
                </div>
                <p style="color: #6b7280; margin-bottom: 15px;">{get_text('recording_in_progress')}</p>
                <button id="stop-rec-btn" onclick="stopRecording()" style="
                    background: #ef4444;
                    color: white;
                    border: none;
                    padding: 14px 28px;
                    border-radius: 25px;
                    font-size: 16px;
                    font-weight: 600;
                    cursor: pointer;
                ">
                    {get_text('stop_recording')}
                </button>
            </div>

            <style>
                @keyframes pulse {{
                    0%, 100% {{ opacity: 1; }}
                    50% {{ opacity: 0.5; }}
                }}
            </style>
        </div>

        <script>
            let mediaRecorder = null;
            let audioChunks = [];
            let recordingStartTime = null;
            let timerInterval = null;
            let warningShown = false;

            // Limits
            const WARNING_TIME = 45 * 60; // 45 Minuten in Sekunden
            const MAX_TIME = 60 * 60;     // 60 Minuten in Sekunden

            // Auf allen Geräten mit Mikrofon-Unterstützung anzeigen
            let selectedDeviceId = null;

            // Mikrofon-Auswahl anzeigen (erst bei Klick auf Start)
            async function showMicSelector() {{
                try {{
                    // Berechtigung anfragen
                    await navigator.mediaDevices.getUserMedia({{ audio: true }});

                    // Geräte laden
                    const devices = await navigator.mediaDevices.enumerateDevices();
                    const audioInputs = devices.filter(device => device.kind === 'audioinput');

                    const select = document.getElementById('mic-select');
                    select.innerHTML = '';

                    audioInputs.forEach((device, index) => {{
                        const option = document.createElement('option');
                        option.value = device.deviceId;
                        option.text = device.label || ('Mikrofon ' + (index + 1));
                        select.appendChild(option);
                    }});

                    if (audioInputs.length > 0) {{
                        selectedDeviceId = audioInputs[0].deviceId;
                    }}

                    select.addEventListener('change', (e) => {{
                        selectedDeviceId = e.target.value;
                    }});

                    // UI aktualisieren: Start-Button verstecken, Mic-Auswahl zeigen
                    document.getElementById('start-rec-btn').style.display = 'none';
                    document.getElementById('mic-selector-container').style.display = 'block';

                }} catch (err) {{
                    alert('{get_text("mic_permission_needed")}');
                    console.error('Mic permission error:', err);
                }}
            }}

            function cancelMicSelection() {{
                document.getElementById('mic-selector-container').style.display = 'none';
                document.getElementById('start-rec-btn').style.display = 'inline-flex';
            }}

            async function confirmMicAndRecord() {{
                await startRecording();
            }}

            // Prüfen ob Mikrofon verfügbar - Button nur zeigen wenn unterstützt
            if (!navigator.mediaDevices || !navigator.mediaDevices.getUserMedia) {{
                document.getElementById('mobile-recorder').style.display = 'none';
            }}

            function updateTimer() {{
                if (!recordingStartTime) return;
                const elapsed = Math.floor((Date.now() - recordingStartTime) / 1000);
                const mins = Math.floor(elapsed / 60).toString().padStart(2, '0');
                const secs = (elapsed % 60).toString().padStart(2, '0');
                document.getElementById('rec-timer').textContent = mins + ':' + secs;

                // Warnung bei 45 Minuten
                if (elapsed >= WARNING_TIME && !warningShown) {{
                    warningShown = true;
                    alert('{get_text("recording_warning_45min")}');
                    // Timer rot färben als visuelle Warnung
                    document.getElementById('rec-timer').style.color = '#ef4444';
                }}

                // Auto-Stopp bei 60 Minuten
                if (elapsed >= MAX_TIME) {{
                    alert('{get_text("recording_stopped_limit")}');
                    stopRecording();
                }}
            }}

            async function startRecording() {{
                try {{
                    // Reset warning flag
                    warningShown = false;

                    // Mikrofon-Constraints mit ausgewähltem Gerät
                    const audioConstraints = selectedDeviceId
                        ? {{ audio: {{ deviceId: {{ exact: selectedDeviceId }} }} }}
                        : {{ audio: true }};

                    const stream = await navigator.mediaDevices.getUserMedia(audioConstraints);

                    // Mikrofon-Auswahl verstecken
                    document.getElementById('mic-selector-container').style.display = 'none';
                    document.getElementById('start-rec-btn').style.display = 'none';

                    // Determine best audio format
                    let mimeType = 'audio/webm;codecs=opus';
                    if (!MediaRecorder.isTypeSupported(mimeType)) {{
                        mimeType = 'audio/mp4';
                        if (!MediaRecorder.isTypeSupported(mimeType)) {{
                            mimeType = 'audio/wav';
                        }}
                    }}

                    mediaRecorder = new MediaRecorder(stream, {{ mimeType: mimeType }});
                    audioChunks = [];

                    mediaRecorder.ondataavailable = (event) => {{
                        if (event.data.size > 0) {{
                            audioChunks.push(event.data);
                        }}
                    }};

                    mediaRecorder.onstop = () => {{
                        const audioBlob = new Blob(audioChunks, {{ type: mimeType }});
                        uploadRecording(audioBlob, mimeType);
                    }};

                    mediaRecorder.start(1000); // Collect data every second
                    recordingStartTime = Date.now();
                    timerInterval = setInterval(updateTimer, 1000);

                    // Update UI
                    document.getElementById('start-rec-btn').style.display = 'none';
                    document.getElementById('recording-ui').style.display = 'block';
                    document.getElementById('rec-divider').style.display = 'none';
                    document.getElementById('mobile-upload-btn').style.display = 'none';
                    document.getElementById('mobile-tips-btn').style.display = 'none';
                    closeTipsModal(); // Modal schliessen falls offen
                    document.getElementById('rec-timer').style.color = '#1d1d1f'; // Reset timer color
                    document.getElementById('rec-timer').textContent = '00:00';

                }} catch (err) {{
                    alert('{get_text("mic_permission_needed")}');
                    console.error('Recording error:', err);
                }}
            }}

            function stopRecording() {{
                if (mediaRecorder && mediaRecorder.state !== 'inactive') {{
                    mediaRecorder.stop();
                    mediaRecorder.stream.getTracks().forEach(track => track.stop());
                    clearInterval(timerInterval);

                    // UI zurücksetzen - nur Start-Button zeigen
                    document.getElementById('start-rec-btn').style.display = 'inline-flex';
                    document.getElementById('recording-ui').style.display = 'none';
                }}
            }}

            function uploadRecording(blob, mimeType) {{
                // Create a File object from the blob
                const extension = mimeType.includes('webm') ? 'webm' : (mimeType.includes('mp4') ? 'm4a' : 'wav');
                const fileName = 'recording_' + new Date().toISOString().slice(0,19).replace(/:/g,'-') + '.' + extension;

                // Create FormData and upload via hidden file input
                const dataTransfer = new DataTransfer();
                const file = new File([blob], fileName, {{ type: mimeType }});
                dataTransfer.items.add(file);

                // Find Streamlit's file uploader input and set the file
                const fileInputs = parent.document.querySelectorAll('input[type="file"]');
                if (fileInputs.length > 0) {{
                    const fileInput = fileInputs[0];
                    fileInput.files = dataTransfer.files;

                    // Trigger change event
                    const event = new Event('change', {{ bubbles: true }});
                    fileInput.dispatchEvent(event);
                }}
            }}

            function handleMobileUpload(input) {{
                if (input.files && input.files[0]) {{
                    const file = input.files[0];
                    const dataTransfer = new DataTransfer();
                    dataTransfer.items.add(file);

                    // Find Streamlit's file uploader input and set the file
                    const fileInputs = parent.document.querySelectorAll('input[type="file"]');
                    if (fileInputs.length > 0) {{
                        const fileInput = fileInputs[0];
                        fileInput.files = dataTransfer.files;

                        // Trigger change event
                        const event = new Event('change', {{ bubbles: true }});
                        fileInput.dispatchEvent(event);
                    }}
                }}
            }}

            function showTipsModal() {{
                const modal = document.getElementById('tips-modal-overlay');
                modal.style.display = 'flex';
                // Auch im parent document anzeigen für korrektes Overlay
                try {{
                    const parentModal = modal.cloneNode(true);
                    parentModal.id = 'tips-modal-parent';
                    parentModal.querySelector('button[onclick="closeTipsModal()"]').onclick = function() {{
                        parent.document.getElementById('tips-modal-parent').remove();
                    }};
                    parentModal.onclick = function() {{
                        parent.document.getElementById('tips-modal-parent').remove();
                    }};
                    parentModal.querySelector('div').onclick = function(e) {{ e.stopPropagation(); }};
                    parent.document.body.appendChild(parentModal);
                    modal.style.display = 'none';
                }} catch(e) {{
                    // Fallback: im iframe anzeigen
                }}
            }}

            function closeTipsModal() {{
                document.getElementById('tips-modal-overlay').style.display = 'none';
                try {{
                    const parentModal = parent.document.getElementById('tips-modal-parent');
                    if (parentModal) parentModal.remove();
                }} catch(e) {{}}
            }}

            // Hide Streamlit's file uploader and expander on mobile
            if (isMobile()) {{
                const style = document.createElement('style');
                style.textContent = `
                    [data-testid="stFileUploader"] {{
                        display: none !important;
                    }}
                    [data-testid="stExpander"] {{
                        display: none !important;
                    }}
                `;
                parent.document.head.appendChild(style);
            }}
        </script>
        '''

        # Recording Component einbetten (für alle Geräte)
        import streamlit.components.v1 as components
        components.html(mobile_recording_html, height=100)

        # File Uploader (immer anzeigen, auf Mobile unter dem Recorder)
        uploaded_file = st.file_uploader(
            get_text("upload_label"),
            type=["mp3", "wav", "m4a", "ogg", "webm", "mp4"],
            help=get_text("upload_help"),
            label_visibility="collapsed"
        )

        # Info-Box mit Tipps entfernt - zu viel UI-Clutter

        # Fehler anzeigen falls vorhanden
        if st.session_state.error:
            st.error(st.session_state.error)
            if st.button(get_text("retry"), use_container_width=True):
                st.session_state.error = None
                st.rerun()

        # =====================================================================
        # AUTOMATISCHER WORKFLOW nach Upload
        # =====================================================================
        if uploaded_file and not st.session_state.processing and not st.session_state.error:
            st.session_state.processing = True

            file_size = len(uploaded_file.getvalue())

            if file_size > MAX_FILE_SIZE:
                st.session_state.error = f"{get_text('file_too_large')} ({file_size // (1024*1024)} MB). {get_text('maximum')}: {MAX_FILE_SIZE_MB} MB"
                st.session_state.processing = False
                st.rerun()

            # Audiodauer prüfen (max 120 Minuten)
            audio_duration = 0.0  # Default-Wert für Minuten-Abrechnung
            try:
                audio_duration = get_audio_duration_minutes(uploaded_file)
                uploaded_file.seek(0)  # Datei zurücksetzen nach dem Lesen
                if audio_duration > MAX_AUDIO_DURATION_MINUTES:
                    st.session_state.error = f"Audiodatei zu lang ({int(audio_duration)} Min.). Maximum: {MAX_AUDIO_DURATION_MINUTES} Min."
                    st.session_state.processing = False
                    st.rerun()
            except Exception:
                audio_duration = 1.0  # Mindestens 1 Minute bei Fehler

            # Protokoll-Limits prüfen (Beta-Phase)
            user_email = st.session_state.get("user_email", "unknown")
            can_create, limit_error, remaining_today, total_remaining = check_protocol_limits(user_email)

            if not can_create:
                st.session_state.error = limit_error
                st.session_state.processing = False
                st.rerun()

            # Prüfen ob Datei zu gross für Whisper und ffmpeg benötigt wird
            if file_size > WHISPER_CHUNK_SIZE:
                global FFMPEG_PATH, FFMPEG_AVAILABLE
                # ffmpeg suchen
                FFMPEG_PATH = get_ffmpeg_path()
                FFMPEG_AVAILABLE = FFMPEG_PATH is not None

                if not FFMPEG_AVAILABLE:
                    # Versuche ffmpeg zu installieren
                    install_status = st.empty()
                    install_status.info("🔧 Installing ffmpeg...")

                    if install_ffmpeg_brew():
                        FFMPEG_PATH = get_ffmpeg_path()
                        FFMPEG_AVAILABLE = FFMPEG_PATH is not None
                        install_status.empty()

                if not FFMPEG_AVAILABLE:
                    st.session_state.error = get_text("ffmpeg_needed")
                    st.session_state.processing = False
                    st.rerun()

            # Datei-Info für Anzeige
            file_name = uploaded_file.name
            file_size_mb = round(file_size / (1024 * 1024), 1)

            # Animierte Verarbeitungsanzeige mit kreisförmigem Fortschrittsring
            progress_container = st.empty()

            def update_circular_progress(percent, status_msg, show_upload_success=True):
                # Berechne stroke-dashoffset (390 = Umfang des Kreises mit r=62)
                offset = 390 - (390 * percent / 100)
                # Spinner nur anzeigen wenn noch nicht fertig
                spinner_html = '<div class="working-spinner"></div>' if percent < 100 else ''
                # Upload-Erfolg nur am Anfang anzeigen
                upload_success_html = f'<div class="upload-success">✓ {get_text("file_uploaded_success")}</div>' if show_upload_success else ''
                progress_container.markdown(f"""
                <div class="processing-animation" id="progress-container">
                    {upload_success_html}
                    <div class="circular-progress">
                        <svg viewBox="0 0 140 140">
                            <defs>
                                <linearGradient id="progressGradient" x1="0%" y1="0%" x2="100%" y2="0%">
                                    <stop offset="0%" style="stop-color:#7c3aed;stop-opacity:1" />
                                    <stop offset="100%" style="stop-color:#2dd4bf;stop-opacity:1" />
                                </linearGradient>
                            </defs>
                            <circle class="bg-circle" cx="70" cy="70" r="62"></circle>
                            <circle class="progress-circle" cx="70" cy="70" r="62" style="stroke-dashoffset: {offset};"></circle>
                        </svg>
                        <div class="percentage">{percent}%</div>
                    </div>
                    <div class="file-info">{file_name} {file_size_mb}MB</div>
                    <div class="status-text">{status_msg}</div>
                    {spinner_html}
                </div>
                """, unsafe_allow_html=True)

            try:
                # Schritt 1: Transkription
                update_circular_progress(10, f"🎙️ {get_text('transcribing')}")

                # Log-Funktion (ohne UI-Anzeige)
                def log_status(msg):
                    pass  # Logging entfernt - UI zu überladen

                uploaded_file.seek(0)
                transcript = transcribe_audio(uploaded_file, openai_client, status_callback=log_status)

                # Transkript-Länge prüfen
                word_count = len(transcript.split())
                char_count = len(transcript)

                # Mindest-Wörter-Prüfung (verhindert Halluzinationen bei zu wenig Input)
                MIN_WORDS_REQUIRED = 50
                if word_count < MIN_WORDS_REQUIRED:
                    raise Exception(f"Zu wenig Inhalt transkribiert ({word_count} Wörter). Mindestens {MIN_WORDS_REQUIRED} Wörter erforderlich. Bitte prüfen Sie, ob die Audio-Datei hörbaren Sprachinhalt enthält.")

                st.session_state.transcript = transcript

                # Schritt 2: Protokoll erstellen
                update_circular_progress(50, f"📝 {get_text('generating')}")
                log_status(f"📝 {get_text('sending_to_mistral')}")

                protocol = generate_protocol_text(transcript, mistral_api_key)
                st.session_state.protocol = protocol

                # Debug: Protokoll-Länge
                protocol_words = len(protocol.split())
                log_status(f"📄 {get_text('protocol_generated')}: {protocol_words} {get_text('words')}")
                if protocol_words < 1500:
                    log_status(f"{get_text('warning_too_short')} ({protocol_words} < 1500)")

                # Schritt 3: PDF erstellen
                update_circular_progress(75, f"📄 {get_text('creating_pdf')}")

                pdf_bytes = parse_markdown_to_pdf(protocol)
                st.session_state.pdf_bytes = pdf_bytes

                # Schritt 4: Word erstellen
                update_circular_progress(90, f"📃 {get_text('creating_word')}")

                docx_bytes = parse_markdown_to_docx(protocol)
                st.session_state.docx_bytes = docx_bytes

                # Fertig
                update_circular_progress(100, f"✓ {get_text('file_uploaded_success')}")

                # Aktivität loggen
                log_activity("Protokoll erstellt", f"{protocol_words} Wörter, {word_count} Wörter Transkript, {audio_duration:.1f} Min. Audio")

                # Protokoll-Zähler erhöhen (Beta-Limits)
                increment_protocol_count(user_email)

                # Transkriptionsminuten vom Benutzerkonto abziehen
                if audio_duration > 0 and user_email and user_email != "unknown":
                    add_transcription_minutes(user_email, audio_duration)

                st.session_state.processing = False
                st.rerun()

            except Exception as e:
                st.session_state.error = f"{get_text('error')}: {str(e)}"
                st.session_state.processing = False
                st.rerun()


if __name__ == "__main__":
    main()
